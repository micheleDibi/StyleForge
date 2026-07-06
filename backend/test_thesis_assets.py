"""
Test unitari per thesis_assets (nessun DB/rete/API key).

Esecuzione: python3 -m pytest test_thesis_assets.py -q
"""

import os

os.environ.setdefault("ANTHROPIC_API_KEY", "dummy")

import pytest

from thesis_assets import (
    ChartAsset,
    ChartRenderError,
    HintAsset,
    TableAsset,
    assign_asset_numbers,
    build_figures_index,
    build_tables_index,
    count_words_excluding_assets,
    format_caption,
    parse_segments,
    protect_asset_blocks,
    render_chart_png,
    restore_asset_blocks,
    sanitize_generated_assets,
    table_to_markdown,
    table_to_plain_lines,
    validate_chart_spec,
    wrap_text_to_width,
)

CHAPTERS_STRUCTURE = {
    "chapters": [
        {"chapter_index": 0, "chapter_title": "Introduzione", "is_special": True},
        {"chapter_index": 1, "chapter_title": "Il quadro normativo", "sections": []},
        {"chapter_index": 2, "chapter_title": "Analisi dei dati", "sections": []},
        {"chapter_index": 0, "chapter_title": "Conclusione", "is_special": True},
        {"chapter_index": 0, "chapter_title": "Bibliografia", "is_special": True},
    ]
}

SYNTHETIC_CONTENT = """# Introduzione

Testo introduttivo con una citazione [1] e una nota {{nota: Rossi, M. (2020). *Titolo*. Roma: Editore. p.12}}.

# Il quadro normativo

## La riforma

Il testo della sezione discute la riforma in dettaglio.

[TABELLA: Confronto tra gli approcci normativi (fonte: ISTAT, 2022)]
| Criterio | Approccio A | Approccio B |
|---|---|---|
| Costo | Basso | Alto |
| Tempi | 6 mesi | 18 mesi |
Fonte: ISTAT, Rapporto annuale 2022
[/TABELLA]

Come mostra la tabella, gli approcci divergono.

HINT: "Inserire qui una fotografia dell'edificio scolastico oggetto del caso di studio."

[GRAFICO: Andamento degli alunni con BES 2015-2022]
{"type": "bar", "x_label": "Anno", "y_label": "Alunni (migliaia)",
 "labels": ["2015", "2018", "2022"],
 "series": [{"name": "BES", "values": [180, 245, 310]}],
 "source": "MIUR, 2023"}
[/GRAFICO]

# Capitolo rinominato dall'utente

## Sezione due

Altro testo con citazione [2].

[TABELLA: Righe irregolari]
| A | B | C |
| 1 | 2 |
| 1 | 2 | 3 | 4 |
[/TABELLA]

HINT: Inserire uno schema del flusso procedurale senza virgolette

[GRAFICO: Grafico con JSON rotto]
{"type": "bar", "labels": ["a"], "series": [{"values": [1,]}
[/GRAFICO]

# Conclusione

Testo conclusivo.

# Bibliografia

[1] Rossi, M. (2020). *Titolo*. Roma: Editore.
"""


def _segments():
    segs = parse_segments(SYNTHETIC_CONTENT)
    assign_asset_numbers(segs, CHAPTERS_STRUCTURE)
    return segs


# ═══════════════════════════════ PARSING ═══════════════════════════════
def test_segmentation_kinds_and_order():
    kinds = [s.kind for s in _segments()]
    assert kinds == [
        'text', 'table', 'text', 'hint', 'chart',
        'text', 'table', 'hint', 'chart', 'text',
    ]


def test_table_parsed_with_separator_and_source():
    table = _segments()[1].asset
    assert isinstance(table, TableAsset)
    assert table.header == ['Criterio', 'Approccio A', 'Approccio B']
    assert table.rows == [['Costo', 'Basso', 'Alto'], ['Tempi', '6 mesi', '18 mesi']]
    assert table.source == 'ISTAT, Rapporto annuale 2022'
    assert 'ISTAT' in table.caption


def test_ragged_rows_normalized_to_header_length():
    table = _segments()[6].asset
    assert table.header == ['A', 'B', 'C']
    assert table.rows == [['1', '2', ''], ['1', '2', '3']]


def test_hint_with_and_without_quotes():
    segs = _segments()
    assert segs[3].asset.text.startswith('Inserire qui una fotografia')
    assert '"' not in segs[3].asset.text
    assert segs[7].asset.text == 'Inserire uno schema del flusso procedurale senza virgolette'


def test_chart_valid_and_broken():
    segs = _segments()
    ok = segs[4].asset
    assert ok.error is None and ok.spec['type'] == 'bar'
    broken = segs[8].asset
    assert broken.error is not None


def test_unclosed_block_treated_as_text():
    content = "Testo prima.\n\n[TABELLA: orfana]\n| A | B |\n| 1 | 2 |\n\nTesto dopo senza chiusura."
    segs = parse_segments(content)
    assert all(s.kind == 'text' for s in segs)
    joined = '\n'.join(s.text for s in segs)
    assert '[TABELLA' not in joined
    assert '| A | B |' in joined  # il corpo resta testo


def test_text_roundtrip_preserves_headings():
    joined = '\n'.join(s.text for s in _segments() if s.kind == 'text')
    assert '# Il quadro normativo' in joined
    assert '## La riforma' in joined
    assert '{{nota: Rossi' in joined


# ═══════════════════════════ NUMERAZIONE ═══════════════════════════
def test_numbering_per_chapter_with_reset_and_fallback():
    segs = _segments()
    assert segs[1].asset.label == 'Tabella 1.1'
    assert segs[4].asset.label == 'Figura 1.1'
    # capitolo rinominato → contatore di fallback (dopo il 2 noto: 3)
    assert segs[6].asset.label == 'Tabella 3.1'
    assert segs[8].asset.label == 'Figura 3.1'


def test_known_second_chapter_number_used():
    content = "# Analisi dei dati\n\n[TABELLA: t]\n| A |\n| 1 |\n[/TABELLA]\n"
    segs = parse_segments(content)
    assign_asset_numbers(segs, CHAPTERS_STRUCTURE)
    assert segs[1].asset.label == 'Tabella 2.1'


def test_asset_before_any_chapter_gets_global_label():
    content = "[TABELLA: t]\n| A |\n| 1 |\n[/TABELLA]\n"
    segs = parse_segments(content)
    assign_asset_numbers(segs, CHAPTERS_STRUCTURE)
    assert segs[0].asset.label == 'Tabella 1'


def test_indices():
    segs = _segments()
    tables = build_tables_index(segs)
    figures = build_figures_index(segs)
    assert [t[0] for t in tables] == ['Tabella 1.1', 'Tabella 3.1']
    assert [f[0] for f in figures] == ['Figura 1.1', 'Figura 3.1']
    assert build_tables_index(parse_segments("solo testo")) == []


def test_format_caption():
    assert format_caption(TableAsset(caption='X', header=['a'], rows=[], label='Tabella 1.2')) == 'Tabella 1.2 – X'
    assert format_caption(TableAsset(caption='X', header=['a'], rows=[])) == 'X'


# ═══════════════════════════ SENTINELLE ═══════════════════════════
def test_protect_restore_identity():
    protected, mapping = protect_asset_blocks(SYNTHETIC_CONTENT)
    assert len(mapping) == 6  # 2 tabelle + 2 grafici + 2 hint
    assert '[TABELLA' not in protected and 'HINT:' not in protected
    restored = restore_asset_blocks(protected, mapping)
    segs = parse_segments(restored)
    kinds = [s.kind for s in segs]
    assert kinds.count('table') == 2 and kinds.count('chart') == 2 and kinds.count('hint') == 2
    # il testo attorno sopravvive
    assert 'Come mostra la tabella' in restored


def test_restore_lost_sentinel_reappends_block():
    protected, mapping = protect_asset_blocks(SYNTHETIC_CONTENT)
    # simula un LLM che ha cancellato la prima sentinella
    first = next(iter(mapping))
    mangled = protected.replace(first, '')
    restored = restore_asset_blocks(mangled, mapping)
    assert mapping[first] in restored  # mai perso: ri-appeso in fondo


def test_restore_duplicated_sentinel_deduplicated():
    protected, mapping = protect_asset_blocks("HINT: \"uno\"\n\ntesto")
    sent = next(iter(mapping))
    duplicated = protected + f"\n\n{sent}\n"
    restored = restore_asset_blocks(duplicated, mapping)
    assert restored.count('HINT:') == 1


def test_sentinels_survive_algorithmic_anti_ai_pass():
    from anti_ai_processor import humanize_text_post_processing
    protected, mapping = protect_asset_blocks(SYNTHETIC_CONTENT)
    out = humanize_text_post_processing(protected, profile='academic', seed=42)
    for sentinel in mapping:
        assert sentinel in out


def test_count_words_excluding_assets():
    with_assets = count_words_excluding_assets(SYNTHETIC_CONTENT)
    only_text = count_words_excluding_assets(
        '\n'.join(s.text for s in parse_segments(SYNTHETIC_CONTENT) if s.kind == 'text')
    )
    assert with_assets == only_text
    assert with_assets < len(SYNTHETIC_CONTENT.split())


# ═══════════════════════════ SANITIZE ═══════════════════════════
def test_sanitize_degrades_broken_chart_to_hint():
    out = sanitize_generated_assets(SYNTHETIC_CONTENT)
    segs = parse_segments(out)
    kinds = [s.kind for s in segs]
    assert kinds.count('chart') == 1  # quello rotto è degradato
    assert kinds.count('hint') == 3
    hints = [s.asset.text for s in segs if s.kind == 'hint']
    assert any('Grafico con JSON rotto' in h for h in hints)


def test_sanitize_strips_orphan_markers():
    out = sanitize_generated_assets("testo\n\n[TABELLA: orfana]\naltro testo con più di una parola")
    assert '[TABELLA' not in out
    assert 'altro testo' in out


def test_sanitize_inlines_note_tokens_in_cells():
    content = ('[TABELLA: t]\n| A | B |\n| x {{nota: Rossi, 2020}} | y |\n[/TABELLA]')
    out = sanitize_generated_assets(content)
    assert '{{nota' not in out
    assert '(Rossi, 2020)' in out


def test_sanitize_noop_on_plain_text():
    plain = "Solo testo senza asset.\n\nCon due paragrafi."
    assert sanitize_generated_assets(plain) == plain


# ═══════════════════════════ GRAFICI ═══════════════════════════
def _spec(ctype, **kw):
    spec = {
        "type": ctype,
        "x_label": "X", "y_label": "Y",
        "labels": ["a", "b", "c"],
        "series": [{"name": "S1", "values": [1, 2, 3]},
                   {"name": "S2", "values": [3, 1, 2]}],
    }
    spec.update(kw)
    return spec


@pytest.mark.parametrize("ctype", ["bar", "line", "pie", "scatter"])
def test_render_chart_png_all_types(ctype):
    spec = _spec(ctype)
    if ctype == "pie":
        spec["series"] = spec["series"][:1]
    png = render_chart_png(ChartAsset(caption="test", spec=spec))
    assert png[:8] == b'\x89PNG\r\n\x1a\n'
    assert len(png) > 1000


def test_render_chart_invalid_spec_raises():
    with pytest.raises(ChartRenderError):
        render_chart_png(ChartAsset(caption="rotto", error="JSON non valido"))
    with pytest.raises(ChartRenderError):
        render_chart_png(ChartAsset(caption="vuoto", spec=None))


def test_validate_chart_spec():
    assert validate_chart_spec(_spec("bar")) is None
    assert validate_chart_spec({"type": "radar"}) is not None
    assert validate_chart_spec(_spec("bar", labels=[])) is not None
    assert validate_chart_spec(_spec("bar", series=[{"values": ["x"]}])) is not None
    assert validate_chart_spec(_spec("pie", series=[{"values": [-1, 2, 3]}])) is not None


def test_charts_disabled_flag(monkeypatch):
    import config
    monkeypatch.setattr(config, 'THESIS_CHARTS_ENABLED', False)
    with pytest.raises(ChartRenderError):
        render_chart_png(ChartAsset(caption="x", spec=_spec("bar")))


# ═══════════════════════════ CONVERSIONI ═══════════════════════════
def test_table_to_markdown_escapes_pipes():
    t = TableAsset(caption='c', header=['A|B', 'C'], rows=[['1', '2']])
    md = table_to_markdown(t)
    assert '| A/B | C |' in md
    assert '--- ' in md


def test_table_to_plain_lines_aligned():
    t = TableAsset(caption='c', header=['Colonna', 'B'],
                   rows=[['x', 'y'], ['valore molto molto molto lungo davvero', 'z']])
    lines = table_to_plain_lines(t, max_col_width=12)
    assert lines[0].startswith('Colonna')
    assert set(lines[1]) <= {'-', ' '}
    assert all(len(l) < 40 for l in lines)
    assert any('…' in l for l in lines)


def test_wrap_text_to_width():
    measure = len  # 1 unità per carattere
    assert wrap_text_to_width('', 10, measure) == ['']
    lines = wrap_text_to_width('una frase di prova con parole', 10, measure)
    assert all(measure(l) <= 10 for l in lines)
    assert ' '.join(lines) == 'una frase di prova con parole'
    # parola singola più larga della colonna → spezzata
    lines = wrap_text_to_width('supercalifragilisti', 6, measure)
    assert all(measure(l) <= 6 for l in lines)
    assert ''.join(lines) == 'supercalifragilisti'


# ═══════════════════════════ BUILDER DOCX ═══════════════════════════
def test_docx_builders_smoke(tmp_path):
    from docx import Document
    from thesis_assets import add_docx_chart, add_docx_hint, add_docx_table

    doc = Document()
    ds = {"font_name": "Times New Roman", "font_size": 12}
    segs = _segments()
    for seg in segs:
        if seg.kind == 'table':
            add_docx_table(doc, seg.asset, ds)
        elif seg.kind == 'chart':
            add_docx_chart(doc, seg.asset, ds)
        elif seg.kind == 'hint':
            add_docx_hint(doc, seg.asset, ds)

    out = tmp_path / "assets.docx"
    doc.save(str(out))

    reread = Document(str(out))
    assert len(reread.tables) == 2
    xml = reread.element.xml
    assert 'FFF3CD' in xml            # shading dei riquadri HINT
    assert 'SUGGERIMENTO' in xml
    # un grafico valido → un'immagine inline; quello rotto → hint box
    from docx.oxml.ns import qn
    drawings = reread.element.findall('.//' + qn('w:drawing'))
    assert len(drawings) == 1


# ═══════════════════════════ FORMULE MATH ═══════════════════════════
MATH_CONTENT = """# Introduzione

Testo introduttivo senza formule.

# Il quadro normativo

Il rapporto $\\beta = \\omega / \\omega_n$ governa la risposta dinamica.

$$D = \\frac{1}{\\sqrt{(1 - \\beta^2)^2 + (2\\xi\\beta)^2}}$$

Segue il commento all'equazione, che ne discute il comportamento.

$$
u_0 = u_{st} \\cdot D
$$

# Capitolo rinominato dall'utente

$$E = mc^2$$

# Conclusione

Testo conclusivo.
"""


def _math_segments():
    segs = parse_segments(MATH_CONTENT)
    assign_asset_numbers(segs, CHAPTERS_STRUCTURE)
    return segs


def test_math_segmentation_kinds():
    kinds = [s.kind for s in _math_segments()]
    assert kinds == ['text', 'math', 'text', 'math', 'text', 'math', 'text']


def test_math_display_and_legacy_fence_parsed():
    segs = _math_segments()
    assert segs[1].asset.latex == r"D = \frac{1}{\sqrt{(1 - \beta^2)^2 + (2\xi\beta)^2}}"
    assert segs[3].asset.latex == r"u_0 = u_{st} \cdot D"  # blocco multi-riga legacy


def test_math_inline_stays_in_text_segment():
    segs = _math_segments()
    assert r"$\beta = \omega / \omega_n$" in segs[0].text


def test_math_numbering_per_chapter_and_fallback():
    segs = _math_segments()
    labels = [s.asset.label for s in segs if s.kind == 'math']
    # "Il quadro normativo" è il capitolo 1; il rinominato prende il fallback 3
    assert labels == ["(1.1)", "(1.2)", "(3.1)"]


def test_math_orphan_fence_removed_body_stays_text():
    segs = parse_segments("prima\n$$\ncorpo rimasto testo\ne basta")
    assert [s.kind for s in segs] == ['text']
    assert 'corpo rimasto testo' in segs[0].text
    assert '$$' not in segs[0].text


def test_count_words_excludes_math():
    text = "Due parole.\n\n$$E = mc^2$$\n\nAltre due parole $\\beta$ qui."
    # display esclusa e inline esclusa: contano solo le parole di prosa
    assert count_words_excluding_assets(text) == 6


def test_add_docx_math_omml_native_and_fallback(tmp_path):
    from docx import Document
    from thesis_assets import add_docx_math
    from thesis_math import MathAsset

    doc = Document()
    ds = {"font_name": "Times New Roman", "font_size": 12}

    add_docx_math(doc, MathAsset(latex=r"E = mc^2", label="(2.1)"), ds)
    p_omml = doc.paragraphs[-1]
    assert 'oMath' in p_omml._p.xml          # equazione nativa Word
    assert '(2.1)' in p_omml.text            # numero a destra via tab

    add_docx_math(doc, MathAsset(latex=r"\left( rotta", label="(2.2)"), ds)
    p_fallback = doc.paragraphs[-1]
    assert 'oMath' not in p_fallback._p.xml  # degradata (PNG o Unicode), mai crash

    out = tmp_path / "math.docx"
    doc.save(str(out))
    assert 'oMath' in Document(str(out)).element.xml
