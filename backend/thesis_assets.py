"""
Elementi visivi della tesi: tabelle, grafici e HINT — unica fonte di verità.

Il contenuto generato (`Thesis.generated_content`) resta un'unica stringa piatta;
questo modulo definisce la sintassi in-band degli asset e tutto ciò che serve a:
  - riconoscerla (parse_segments) e normalizzarla a generazione
    (sanitize_generated_assets);
  - proteggerla attraverso le riscritture LLM/anti-AI con sentinelle
    ZZASSET{k}ZZ (protect_asset_blocks / restore_asset_blocks);
  - numerarla in modo deterministico per capitolo (assign_asset_numbers) e
    costruire gli indici di tabelle/figure;
  - renderizzarla: grafici PNG accademici via matplotlib (render_chart_png),
    builder python-docx (add_docx_table/chart/hint), conversioni md/txt.

Sintassi (marcatori sempre su righe isolate):

    [TABELLA: <didascalia>]
    | Intestazione 1 | Intestazione 2 |
    | valore | valore |
    Fonte: ...            (riga opzionale)
    [/TABELLA]

    [GRAFICO: <didascalia>]
    {"type": "bar|line|pie|scatter", "x_label": "...", "y_label": "...",
     "labels": [...], "series": [{"name": "...", "values": [...]}],
     "source": "..."}
    [/GRAFICO]

    HINT: "<cosa inserire, perché, dove reperirlo>"

Il frontend replica parsing e numerazione in
frontend/src/utils/thesisAssets.js: le due implementazioni devono restare
allineate.
"""

import json
import logging
import re
from dataclasses import dataclass, field
from io import BytesIO
from typing import Callable, Optional

import config

logger = logging.getLogger(__name__)


class ChartRenderError(Exception):
    """Il grafico non è renderizzabile (spec invalida, matplotlib assente, flag off)."""


# ═══════════════════════════════════════════════════════════════════════════
# SINTASSI / REGEX
# ═══════════════════════════════════════════════════════════════════════════
TABLE_OPEN_RE = re.compile(r'^\s*\[TABELLA\s*:?\s*(.*?)\]\s*$')
TABLE_CLOSE_RE = re.compile(r'^\s*\[/TABELLA\]\s*$')
CHART_OPEN_RE = re.compile(r'^\s*\[GRAFICO\s*:?\s*(.*?)\]\s*$')
CHART_CLOSE_RE = re.compile(r'^\s*\[/GRAFICO\]\s*$')
HINT_RE = re.compile(r'^\s*HINT\s*:\s*(.+?)\s*$')
SOURCE_LINE_RE = re.compile(r'^\s*Fonte\s*:\s*(.+?)\s*$', re.IGNORECASE)
# Riga separatrice markdown |---|:---| ecc. (ignorata)
MD_SEPARATOR_RE = re.compile(r'^\s*\|?[\s:|-]+\|?\s*$')
SENTINEL_RE = re.compile(r'ZZASSET(\d+)ZZ')

# Oltre questo lookahead un marcatore senza chiusura è considerato testo rotto.
MAX_BLOCK_LINES = 80

CHART_TYPES = {'bar', 'line', 'pie', 'scatter'}

# Titoli dei capitoli speciali (non contano nella numerazione N.M)
SPECIAL_TITLES = {'introduzione', 'conclusione', 'conclusioni', 'bibliografia'}


# ═══════════════════════════════════════════════════════════════════════════
# MODELLO DATI
# ═══════════════════════════════════════════════════════════════════════════
@dataclass
class TableAsset:
    caption: str
    header: list
    rows: list
    source: Optional[str] = None
    label: Optional[str] = None  # "Tabella 1.2" (assegnata da assign_asset_numbers)


@dataclass
class ChartAsset:
    caption: str
    spec: Optional[dict] = None
    error: Optional[str] = None
    label: Optional[str] = None  # "Figura 2.1"


@dataclass
class HintAsset:
    text: str


@dataclass
class Segment:
    kind: str  # 'text' | 'table' | 'chart' | 'hint'
    text: Optional[str] = None
    asset: object = None


# ═══════════════════════════════════════════════════════════════════════════
# PARSING
# ═══════════════════════════════════════════════════════════════════════════
def _parse_table_row(line: str) -> list:
    """`| a | b |` → ['a', 'b'] (pipe iniziale/finale opzionali)."""
    cells = line.strip().strip('|').split('|')
    return [c.strip() for c in cells]


def _inline_note_tokens(cell: str) -> str:
    """Le note {{nota: ...}} non sono renderizzabili nelle celle → testo inline."""
    return re.sub(r'\{\{\s*nota\s*:\s*(.*?)\}\}', r'(\1)', cell)


def _strip_outer_quotes(s: str) -> str:
    s = s.strip()
    for a, b in (('"', '"'), ('“', '”'), ("'", "'")):
        if len(s) >= 2 and s.startswith(a) and s.endswith(b):
            return s[1:-1].strip()
    # virgoletta solo in apertura o solo in chiusura (modello sciatto)
    return s.strip('"').strip()


def _build_table(caption: str, body_lines: list) -> Optional[TableAsset]:
    """Costruisce una TableAsset dalle righe interne al blocco. None se vuota."""
    header = None
    rows = []
    source = None
    for line in body_lines:
        if not line.strip():
            continue
        if MD_SEPARATOR_RE.match(line) and '|' in line and '-' in line:
            continue  # separatore markdown
        m = SOURCE_LINE_RE.match(line)
        if m and not line.lstrip().startswith('|'):
            source = m.group(1)
            continue
        if line.lstrip().startswith('|'):
            cells = [_inline_note_tokens(c) for c in _parse_table_row(line)]
            if header is None:
                header = cells
            else:
                rows.append(cells)
        # righe spurie non-| dentro il blocco: ignorate
    if not header:
        return None
    # normalizza le righe alla lunghezza dell'intestazione
    n = len(header)
    rows = [(r + [''] * n)[:n] for r in rows]
    return TableAsset(caption=caption.strip(), header=header, rows=rows, source=source)


def validate_chart_spec(spec) -> Optional[str]:
    """Ritorna un messaggio d'errore se la spec non è renderizzabile, None se ok."""
    if not isinstance(spec, dict):
        return "la specifica non è un oggetto JSON"
    ctype = str(spec.get('type', '')).lower().strip()
    if ctype not in CHART_TYPES:
        return f"tipo di grafico non supportato: '{ctype or '?'}'"
    labels = spec.get('labels')
    if not isinstance(labels, list) or not labels:
        return "'labels' mancante o vuoto"
    series = spec.get('series')
    if not isinstance(series, list) or not series:
        return "'series' mancante o vuoto"
    for s in series:
        if not isinstance(s, dict):
            return "ogni serie deve essere un oggetto {name, values}"
        values = s.get('values')
        if not isinstance(values, list) or not values:
            return "serie senza 'values'"
        try:
            [float(v) for v in values]
        except (TypeError, ValueError):
            return "'values' contiene valori non numerici"
    if ctype == 'pie':
        first = [float(v) for v in series[0]['values']]
        if any(v < 0 for v in first):
            return "un grafico a torta non ammette valori negativi"
        if all(v == 0 for v in first):
            return "un grafico a torta richiede almeno un valore positivo"
    return None


def _build_chart(caption: str, body_lines: list) -> ChartAsset:
    """Costruisce una ChartAsset dal corpo JSON (error valorizzato se invalida)."""
    body = '\n'.join(body_lines).strip()
    start, end = body.find('{'), body.rfind('}')
    if start == -1 or end <= start:
        return ChartAsset(caption=caption.strip(), error="nessun oggetto JSON nel blocco")
    try:
        spec = json.loads(body[start:end + 1])
    except (ValueError, TypeError) as e:
        return ChartAsset(caption=caption.strip(), error=f"JSON non valido: {e}")
    err = validate_chart_spec(spec)
    if err:
        return ChartAsset(caption=caption.strip(), spec=spec, error=err)
    return ChartAsset(caption=caption.strip(), spec=spec)


def parse_segments(content: str) -> list:
    """
    Divide il contenuto piatto in segmenti text/table/chart/hint.
    Lenient: blocco non chiuso entro MAX_BLOCK_LINES → marcatore rimosso e
    corpo trattato come testo; tabella senza righe → rimossa. Non solleva mai.
    """
    lines = (content or '').split('\n')
    segments = []
    text_buf = []

    def flush_text():
        if text_buf:
            joined = '\n'.join(text_buf)
            if joined.strip():  # segmenti di solo whitespace (tra due asset) scartati
                segments.append(Segment(kind='text', text=joined))
            text_buf.clear()

    def find_close(start: int, close_re) -> int:
        for j in range(start, min(start + MAX_BLOCK_LINES + 1, len(lines))):
            if close_re.match(lines[j]):
                return j
        return -1

    i = 0
    while i < len(lines):
        line = lines[i]

        m = TABLE_OPEN_RE.match(line)
        if m:
            close = find_close(i + 1, TABLE_CLOSE_RE)
            if close == -1:
                i += 1  # marcatore orfano: rimosso, il corpo resta testo
                continue
            table = _build_table(m.group(1), lines[i + 1:close])
            if table:
                flush_text()
                segments.append(Segment(kind='table', asset=table))
            i = close + 1
            continue

        m = CHART_OPEN_RE.match(line)
        if m:
            close = find_close(i + 1, CHART_CLOSE_RE)
            if close == -1:
                i += 1
                continue
            flush_text()
            segments.append(Segment(kind='chart', asset=_build_chart(m.group(1), lines[i + 1:close])))
            i = close + 1
            continue

        if TABLE_CLOSE_RE.match(line) or CHART_CLOSE_RE.match(line):
            i += 1  # chiusura orfana: rimossa
            continue

        m = HINT_RE.match(line)
        if m:
            flush_text()
            segments.append(Segment(kind='hint', asset=HintAsset(text=_strip_outer_quotes(m.group(1)))))
            i += 1
            continue

        text_buf.append(line)
        i += 1

    flush_text()
    return segments


# ═══════════════════════════════════════════════════════════════════════════
# NUMERAZIONE + INDICI
# ═══════════════════════════════════════════════════════════════════════════
def _normalize_title(t: str) -> str:
    return (t or '').strip().casefold()


def assign_asset_numbers(segments: list, chapters_structure: Optional[dict] = None) -> None:
    """
    Assegna label "Tabella N.M" / "Figura N.M" per capitolo (contatore Figura
    condiviso da grafici e future immagini). I capitoli speciali (is_special,
    Introduzione/Conclusione/Bibliografia) non contano; titolo `# ` sconosciuto
    → contatore sequenziale di fallback. Asset prima di qualunque capitolo o
    dentro uno speciale → numerazione globale senza prefisso ("Tabella 3").
    """
    title_to_num = {}
    special_titles = set(SPECIAL_TITLES)
    chapters = (chapters_structure or {}).get('chapters') or []
    seq = 0
    for ch in chapters:
        title = _normalize_title(ch.get('chapter_title') or ch.get('title') or '')
        if not title:
            continue
        if ch.get('is_special'):
            special_titles.add(title)
        else:
            seq += 1
            title_to_num[title] = ch.get('chapter_index') or seq

    current_ch = None
    fallback_counter = max(title_to_num.values(), default=0)
    table_n = figure_n = 0
    global_table_n = global_figure_n = 0

    for seg in segments:
        if seg.kind == 'text':
            # aggiorna il capitolo corrente in base ai titoli `# ` del segmento
            for line in seg.text.split('\n'):
                if line.startswith('# ') and not line.startswith('## '):
                    title = _normalize_title(line[2:])
                    table_n = figure_n = 0
                    if title in special_titles:
                        current_ch = None
                    elif title in title_to_num:
                        current_ch = title_to_num[title]
                        fallback_counter = max(fallback_counter, current_ch)
                    else:
                        fallback_counter += 1
                        current_ch = fallback_counter
        elif seg.kind == 'table':
            if current_ch is not None:
                table_n += 1
                seg.asset.label = f"Tabella {current_ch}.{table_n}"
            else:
                global_table_n += 1
                seg.asset.label = f"Tabella {global_table_n}"
        elif seg.kind == 'chart':
            if current_ch is not None:
                figure_n += 1
                seg.asset.label = f"Figura {current_ch}.{figure_n}"
            else:
                global_figure_n += 1
                seg.asset.label = f"Figura {global_figure_n}"


def format_caption(asset) -> str:
    """"Tabella 1.2 – didascalia" (o solo didascalia se label assente)."""
    label = getattr(asset, 'label', None)
    caption = (getattr(asset, 'caption', '') or '').strip()
    if label and caption:
        return f"{label} – {caption}"
    return label or caption


def build_tables_index(segments: list) -> list:
    """[(label, caption), ...] in ordine di documento."""
    return [(s.asset.label or 'Tabella', s.asset.caption)
            for s in segments if s.kind == 'table']


def build_figures_index(segments: list) -> list:
    return [(s.asset.label or 'Figura', s.asset.caption)
            for s in segments if s.kind == 'chart']


# ═══════════════════════════════════════════════════════════════════════════
# SANITIZE (a generazione)
# ═══════════════════════════════════════════════════════════════════════════
def _table_block_lines(table: TableAsset) -> list:
    out = [f"[TABELLA: {table.caption}]"]
    out.append('| ' + ' | '.join(table.header) + ' |')
    for row in table.rows:
        out.append('| ' + ' | '.join(row) + ' |')
    if table.source:
        out.append(f"Fonte: {table.source}")
    out.append('[/TABELLA]')
    return out


def _chart_block_lines(chart: ChartAsset) -> list:
    return [
        f"[GRAFICO: {chart.caption}]",
        json.dumps(chart.spec, ensure_ascii=False),
        '[/GRAFICO]',
    ]


def sanitize_generated_assets(text: str) -> str:
    """
    Normalizza gli asset appena generati dal modello:
      - marcatori orfani rimossi (il corpo resta testo);
      - righe tabella pareggiate all'intestazione, {{nota}} nelle celle inlined;
      - blocco grafico con JSON invalido → degradato a riga HINT;
      - blocchi ricomposti in forma canonica (JSON compatto su una riga).
    """
    if not text or ('[TABELLA' not in text and '[GRAFICO' not in text):
        return text

    out_lines = []
    for seg in parse_segments(text):
        if seg.kind == 'text':
            out_lines.append(seg.text)
        elif seg.kind == 'table':
            out_lines.extend([''] + _table_block_lines(seg.asset) + [''])
        elif seg.kind == 'chart':
            if seg.asset.error:
                logger.warning(f"Grafico degradato a HINT ({seg.asset.error}): {seg.asset.caption[:80]}")
                out_lines.extend([
                    '',
                    f'HINT: "Inserire qui un grafico su: {seg.asset.caption or "dati della sezione"}'
                    f' (dati non disponibili o in formato non valido)."',
                    '',
                ])
            else:
                out_lines.extend([''] + _chart_block_lines(seg.asset) + [''])
        elif seg.kind == 'hint':
            out_lines.extend(['', f'HINT: "{seg.asset.text}"', ''])

    result = '\n'.join(out_lines)
    return re.sub(r'\n{3,}', '\n\n', result)


# ═══════════════════════════════════════════════════════════════════════════
# SENTINELLE (protezione attraverso riscritture LLM / anti-AI)
# ═══════════════════════════════════════════════════════════════════════════
def protect_asset_blocks(text: str):
    """
    Sostituisce ogni blocco [TABELLA]/[GRAFICO] e ogni riga HINT con una
    sentinella ZZASSET{k}ZZ su paragrafo isolato. Ritorna (testo, mappa
    {sentinella: blocco_originale}).
    """
    if not text:
        return text, {}

    mapping = {}
    out_lines = []
    lines = text.split('\n')

    def find_close(start: int, close_re) -> int:
        for j in range(start, min(start + MAX_BLOCK_LINES + 1, len(lines))):
            if close_re.match(lines[j]):
                return j
        return -1

    i = 0
    while i < len(lines):
        line = lines[i]
        close = -1
        if TABLE_OPEN_RE.match(line):
            close = find_close(i + 1, TABLE_CLOSE_RE)
        elif CHART_OPEN_RE.match(line):
            close = find_close(i + 1, CHART_CLOSE_RE)

        if close != -1:
            sentinel = f"ZZASSET{len(mapping)}ZZ"
            mapping[sentinel] = '\n'.join(lines[i:close + 1])
            out_lines.extend(['', sentinel, ''])
            i = close + 1
            continue

        if HINT_RE.match(line):
            sentinel = f"ZZASSET{len(mapping)}ZZ"
            mapping[sentinel] = line.strip()
            out_lines.extend(['', sentinel, ''])
            i += 1
            continue

        out_lines.append(line)
        i += 1

    protected = re.sub(r'\n{3,}', '\n\n', '\n'.join(out_lines))
    return protected, mapping


def restore_asset_blocks(text: str, mapping: dict) -> str:
    """
    Ripristina i blocchi al posto delle sentinelle. Sentinella scomparsa dopo
    le riscritture → blocco ri-appeso in fondo (mai perso); occorrenze duplicate
    → solo la prima sostituita, le altre rimosse.
    """
    if not mapping:
        return text
    result = text or ''
    lost = []
    for sentinel, block in mapping.items():
        if sentinel in result:
            first, _, rest = result.partition(sentinel)
            rest = rest.replace(sentinel, '')
            result = first + '\n\n' + block + '\n\n' + rest
        else:
            logger.warning(f"Sentinella {sentinel} persa nelle riscritture: blocco ri-appeso in fondo")
            lost.append(block)
    # eventuali sentinelle spurie inventate dal modello
    result = SENTINEL_RE.sub('', result)
    if lost:
        result = result.rstrip() + '\n\n' + '\n\n'.join(lost) + '\n'
    return re.sub(r'\n{3,}', '\n\n', result)


def count_words_excluding_assets(text: str) -> int:
    """Parole del solo testo discorsivo (tabelle/grafici/HINT esclusi)."""
    if not text:
        return 0
    protected, _ = protect_asset_blocks(text)
    return len(SENTINEL_RE.sub(' ', protected).split())


# ═══════════════════════════════════════════════════════════════════════════
# RENDERING GRAFICI (matplotlib, backend Agg)
# ═══════════════════════════════════════════════════════════════════════════
_CHART_GRAYS = ['#404040', '#707070', '#9e9e9e', '#c4c4c4', '#dddddd']
_CHART_HATCHES = ['', '//', 'xx', '..', '\\\\']
_CHART_MARKERS = ['o', 's', '^', 'D', 'v']
_CHART_LINESTYLES = ['-', '--', '-.', ':', '-']


def render_chart_png(chart: ChartAsset, width_in: float = 6.0, dpi: int = 200) -> bytes:
    """
    Renderizza la spec del grafico in PNG con stile accademico (serif, scala
    di grigi, spines essenziali). Solleva ChartRenderError se il grafico non è
    renderizzabile o se THESIS_CHARTS_ENABLED è False.
    """
    if not getattr(config, 'THESIS_CHARTS_ENABLED', True):
        raise ChartRenderError("generazione grafici disabilitata (THESIS_CHARTS_ENABLED=false)")
    if chart.error or not chart.spec:
        raise ChartRenderError(chart.error or "specifica del grafico mancante")

    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
    except Exception as e:  # noqa: BLE001
        raise ChartRenderError(f"matplotlib non disponibile: {e}")

    spec = chart.spec
    ctype = str(spec['type']).lower().strip()
    labels = [str(l) for l in spec['labels']]
    series = []
    for s in spec['series']:
        values = [float(v) for v in s['values']]
        n = min(len(values), len(labels))
        series.append({'name': str(s.get('name', '') or ''), 'values': values[:n]})
    n_points = min(len(labels), min(len(s['values']) for s in series))
    if n_points < 1:
        raise ChartRenderError("nessun punto dati")
    labels = labels[:n_points]
    series = [{'name': s['name'], 'values': s['values'][:n_points]} for s in series]

    rc = {
        'font.family': 'serif',
        'font.size': 10,
        'axes.titlesize': 10,
        'axes.labelsize': 10,
        'axes.edgecolor': '#404040',
        'axes.linewidth': 0.8,
        'xtick.color': '#404040',
        'ytick.color': '#404040',
        'text.color': '#202020',
        'axes.labelcolor': '#202020',
        'legend.frameon': False,
        'figure.dpi': dpi,
    }

    try:
        with plt.rc_context(rc):
            fig, ax = plt.subplots(figsize=(width_in, width_in * 0.62))
            show_legend = len(series) > 1 or (series[0]['name'] and ctype != 'pie')

            if ctype == 'bar':
                import numpy as np  # dipendenza di matplotlib
                x = np.arange(len(labels))
                total_w = 0.72
                bw = total_w / len(series)
                for k, s in enumerate(series):
                    offset = (k - (len(series) - 1) / 2) * bw
                    ax.bar(x + offset, s['values'], bw * 0.92,
                           label=s['name'] or f"Serie {k + 1}",
                           color=_CHART_GRAYS[k % len(_CHART_GRAYS)],
                           hatch=_CHART_HATCHES[k % len(_CHART_HATCHES)],
                           edgecolor='#303030', linewidth=0.6)
                ax.set_xticks(x)
                ax.set_xticklabels(labels, rotation=30 if max(len(l) for l in labels) > 8 else 0,
                                   ha='right' if max(len(l) for l in labels) > 8 else 'center')
            elif ctype == 'line':
                for k, s in enumerate(series):
                    ax.plot(labels, s['values'],
                            label=s['name'] or f"Serie {k + 1}",
                            color=_CHART_GRAYS[k % len(_CHART_GRAYS)],
                            marker=_CHART_MARKERS[k % len(_CHART_MARKERS)],
                            linestyle=_CHART_LINESTYLES[k % len(_CHART_LINESTYLES)],
                            markersize=4, linewidth=1.4)
                if max(len(l) for l in labels) > 8:
                    plt.setp(ax.get_xticklabels(), rotation=30, ha='right')
            elif ctype == 'scatter':
                for k, s in enumerate(series):
                    ax.scatter(labels, s['values'],
                               label=s['name'] or f"Serie {k + 1}",
                               color=_CHART_GRAYS[k % len(_CHART_GRAYS)],
                               marker=_CHART_MARKERS[k % len(_CHART_MARKERS)], s=28)
                if max(len(l) for l in labels) > 8:
                    plt.setp(ax.get_xticklabels(), rotation=30, ha='right')
            elif ctype == 'pie':
                values = series[0]['values']
                wedge_colors = [_CHART_GRAYS[k % len(_CHART_GRAYS)] for k in range(len(values))]
                ax.pie(values, labels=labels, colors=wedge_colors, startangle=90,
                       autopct='%1.1f%%', pctdistance=0.75,
                       wedgeprops={'edgecolor': 'white', 'linewidth': 1.2},
                       textprops={'fontsize': 9})
                ax.axis('equal')
                show_legend = False
            else:
                raise ChartRenderError(f"tipo non supportato: {ctype}")

            if ctype != 'pie':
                if spec.get('x_label'):
                    ax.set_xlabel(str(spec['x_label']))
                if spec.get('y_label'):
                    ax.set_ylabel(str(spec['y_label']))
                ax.spines['top'].set_visible(False)
                ax.spines['right'].set_visible(False)
                ax.grid(axis='y', color='#d0d0d0', linewidth=0.5, alpha=0.8)
                ax.set_axisbelow(True)
            if show_legend:
                ax.legend(fontsize=9)

            fig.tight_layout()
            buf = BytesIO()
            fig.savefig(buf, format='png', dpi=dpi, facecolor='white')
            plt.close(fig)
            return buf.getvalue()
    except ChartRenderError:
        raise
    except Exception as e:  # noqa: BLE001
        try:
            plt.close('all')
        except Exception:  # noqa: BLE001
            pass
        raise ChartRenderError(f"errore di rendering: {e}")


# ═══════════════════════════════════════════════════════════════════════════
# CONVERSIONI md / txt
# ═══════════════════════════════════════════════════════════════════════════
def table_to_markdown(table: TableAsset) -> str:
    def esc(c):
        return c.replace('|', '/')
    lines = ['| ' + ' | '.join(esc(c) for c in table.header) + ' |']
    lines.append('|' + '|'.join([' --- '] * len(table.header)) + '|')
    for row in table.rows:
        lines.append('| ' + ' | '.join(esc(c) for c in row) + ' |')
    return '\n'.join(lines)


def table_to_plain_lines(table: TableAsset, max_col_width: int = 28) -> list:
    """Tabella come testo allineato a colonne (per l'export txt)."""
    all_rows = [table.header] + table.rows
    n = len(table.header)
    widths = [
        min(max_col_width, max(len(str(r[c])) for r in all_rows))
        for c in range(n)
    ]

    def fmt(row):
        cells = []
        for c in range(n):
            cell = str(row[c])
            if len(cell) > widths[c]:
                cell = cell[:widths[c] - 1] + '…'
            cells.append(cell.ljust(widths[c]))
        return '  '.join(cells).rstrip()

    sep = '  '.join('-' * w for w in widths)
    return [fmt(table.header), sep] + [fmt(r) for r in table.rows]


# ═══════════════════════════════════════════════════════════════════════════
# BUILDER DOCX (python-docx)
# ═══════════════════════════════════════════════════════════════════════════
HINT_FILL = 'FFF3CD'    # ambra chiaro
HINT_BORDER = 'B8860B'  # ambra scuro
TABLE_HEADER_FILL = 'F2F2F2'
HINT_PREFIX = '⚠ SUGGERIMENTO — DA SOSTITUIRE: '


def _docx_caption_paragraph(doc, text: str, ds: dict, italic: bool = False):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    font_name = ds.get('font_name', 'Times New Roman')
    size = Pt(max(8, int(ds.get('caption_font_size', ds.get('font_size', 12) - 1))))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # label in grassetto ("Tabella 1.2"), didascalia normale
    m = re.match(r'^((?:Tabella|Figura)\s+[\d.]+)(\s*–\s*)(.*)$', text, re.S)
    parts = [(m.group(1), True), (m.group(2) + m.group(3), False)] if m else [(text, False)]
    for chunk, bold in parts:
        run = p.add_run(chunk)
        run.font.name = font_name
        run.font.size = size
        run.font.bold = bold
        run.font.italic = italic
    return p


def _set_cell_shading(cell, fill: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_table_borders(table, color: str = '666666', size: int = 4):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(size))
        el.set(qn('w:color'), color)
        borders.append(el)
    tbl_pr.append(borders)


def add_docx_table(doc, table: TableAsset, ds: dict):
    """Tabella accademica: didascalia sopra, bordi sottili, header in evidenza."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    font_name = ds.get('font_name', 'Times New Roman')
    cell_size = Pt(max(8, int(ds.get('font_size', 12) - 1)))

    _docx_caption_paragraph(doc, format_caption(table), ds)

    n = len(table.header)
    t = doc.add_table(rows=1 + len(table.rows), cols=n)
    try:
        t.style = 'Table Grid'
    except KeyError:
        pass  # stile assente nel template: i bordi espliciti bastano
    _set_table_borders(t)
    t.autofit = True

    header_fill = ds.get('table_header_fill', TABLE_HEADER_FILL)
    for c, text in enumerate(table.header):
        cell = t.rows[0].cells[c]
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        run.font.name = font_name
        run.font.size = cell_size
        run.font.bold = True
        _set_cell_shading(cell, header_fill)
    for r, row in enumerate(table.rows):
        for c, text in enumerate(row):
            cell = t.rows[r + 1].cells[c]
            cell.text = ''
            run = cell.paragraphs[0].add_run(text)
            run.font.name = font_name
            run.font.size = cell_size

    if table.source:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Fonte: {table.source}")
        run.font.name = font_name
        run.font.size = Pt(max(7, int(ds.get('font_size', 12) - 2)))
        run.font.italic = True
    else:
        doc.add_paragraph()  # respiro dopo la tabella


def add_docx_hint(doc, hint: HintAsset, ds: dict):
    """Riquadro ambra ben visibile: segnaposto da sostituire manualmente."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    font_name = ds.get('font_name', 'Times New Roman')
    size = Pt(int(ds.get('font_size', 12)))

    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), ds.get('hint_fill', HINT_FILL))
    ppr.append(shd)
    borders = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '8')
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), HINT_BORDER)
        borders.append(el)
    ppr.append(borders)

    run = p.add_run(HINT_PREFIX)
    run.font.name = font_name
    run.font.size = size
    run.font.bold = True
    run = p.add_run(hint.text)
    run.font.name = font_name
    run.font.size = size


def add_docx_chart(doc, chart: ChartAsset, ds: dict):
    """Grafico PNG centrato con didascalia sotto; su errore → riquadro HINT."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    try:
        png = render_chart_png(chart)
    except ChartRenderError as e:
        logger.warning(f"Grafico '{chart.caption[:60]}' non renderizzato nel docx: {e}")
        add_docx_hint(doc, HintAsset(
            text=f"Grafico non generabile ({e}). Inserire manualmente: {chart.caption}"
        ), ds)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(BytesIO(png), width=Inches(5.8))
    _docx_caption_paragraph(doc, format_caption(chart), ds, italic=True)


# ═══════════════════════════════════════════════════════════════════════════
# WORD-WRAP GENERICO (per le tabelle dell'export PDF)
# ═══════════════════════════════════════════════════════════════════════════
def wrap_text_to_width(text: str, width: float, measure_fn: Callable[[str], float]) -> list:
    """
    Spezza `text` in righe di larghezza <= width secondo `measure_fn` (es.
    fitz.get_text_length). Le parole più larghe della colonna vengono spezzate
    a caratteri. Ritorna sempre almeno una riga.
    """
    words = (text or '').split()
    if not words:
        return ['']
    lines = []
    current = ''

    def push_word(word):
        nonlocal current
        if measure_fn(word) <= width:
            current = word
            return
        # parola singola più larga della colonna: spezza a caratteri
        chunk = ''
        for ch in word:
            if measure_fn(chunk + ch) > width and chunk:
                lines.append(chunk)
                chunk = ch
            else:
                chunk += ch
        current = chunk

    for word in words:
        if not current:
            push_word(word)
            continue
        candidate = f"{current} {word}"
        if measure_fn(candidate) <= width:
            current = candidate
        else:
            lines.append(current)
            push_word(word)
    if current:
        lines.append(current)
    return lines or ['']
