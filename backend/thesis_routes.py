"""
Router FastAPI per la gestione delle tesi/relazioni.

Questo modulo contiene tutti gli endpoint relativi alla funzionalità
di generazione tesi, inclusi lookup, CRUD, allegati e fasi di generazione.
"""

import uuid
import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional, List

from token_utils import cap_output_tokens

logger = logging.getLogger(__name__)

from fastapi import APIRouter, HTTPException, Request, UploadFile, File, BackgroundTasks, Depends, Form
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session as DBSession, defer
from sqlalchemy import text

from models import (
    ThesisCreateRequest, ThesisResponse, ThesisListResponse,
    ThesisAttachmentResponse, ThesisAttachmentsListResponse,
    GenerateChaptersResponse, ConfirmChaptersRequest,
    GenerateSectionsResponse, ConfirmSectionsRequest,
    StartContentGenerationResponse, GenerationStatusResponse,
    ChapterGenerationStatus, SectionGenerationStatus, LookupDataResponse,
    WritingStyleResponse, ContentDepthResponse,
    AudienceKnowledgeLevelResponse, AudienceSizeResponse,
    IndustryResponse, TargetAudienceResponse,
    ThesisStatus, ThesisWikiStatus, ChapterInfo, ThesisUrlAttachmentRequest,
    ThesisResearchSearchRequest, ThesisResearchSummarizeRequest,
    ThesisAddPapersRequest, ThesisAddPapersResponse,
    WikiIngestRequest, WikiStatusResponse, WikiLintReportResponse,
    WikiContentResponse,
    PaperKeywordSuggestResponse,
)
from db_models import (
    User, Thesis, ThesisAttachment, ThesisGenerationJob,
    WritingStyle, ContentDepthLevel, AudienceKnowledgeLevel,
    AudienceSize, Industry, TargetAudience, Session
)
from database import SessionLocal, get_db
from auth import get_current_active_user, require_permission
from credits import estimate_credits, deduct_credits, add_credits
from attachment_processor import (
    process_attachment, save_uploaded_file, delete_attachment_file,
    build_attachments_context, cleanup_thesis_attachments,
    sanitize_text_for_db,
)
from ai_client import get_ai_client, humanize_text_with_claude
from ai_exceptions import InsufficientCreditsError
from session_manager import session_manager
from template_service import get_template_by_id, get_page_dimensions, get_export_templates
from thesis_assets import (
    ChartRenderError,
    HintAsset,
    assign_asset_numbers,
    add_docx_chart,
    add_docx_hint,
    add_docx_table,
    build_figures_index,
    build_tables_index,
    count_words_excluding_assets,
    format_caption,
    parse_segments,
    protect_asset_blocks,
    render_chart_png,
    restore_asset_blocks,
    sanitize_generated_assets,
    add_docx_math,
    sanitize_math_outside_assets,
    table_to_markdown,
    table_to_plain_lines,
    wrap_text_to_width,
)
from thesis_math import (
    MathRenderError,
    has_inline_math,
    inline_math_to_unicode,
    iter_inline_math,
    latex_to_omml,
    latex_to_unicode,
    protect_math_spans,
    render_math_png,
    restore_math_spans,
    unprotect_math_spans,
)
from research_providers import UnifiedPaper
from research_service import DEFAULT_SOURCES, PROVIDER_REGISTRY, run_search_pipeline
from research_summarizer import (
    SummaryResult,
    paper_to_attachment_text,
    render_paper_with_summary,
    summarize_paper,
)
from keyword_extractor import extract_paper_keywords_from_attachments
import config
import ssrf_guard
from rate_limit import FETCH_LIMIT, limiter

# Mime type convenzionale per allegati di tipo "paper accademico"
PAPER_MIME_TYPE = "application/x-research-paper"

# Router
router = APIRouter(prefix="/api/thesis", tags=["Thesis Generation"])


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def get_thesis_by_id(db: DBSession, thesis_id: str, user_id: str) -> Thesis:
    """Recupera una tesi verificando l'ownership."""
    thesis = db.query(Thesis).filter(
        Thesis.id == thesis_id,
        Thesis.user_id == user_id
    ).first()

    if not thesis:
        raise HTTPException(status_code=404, detail="Tesi non trovata")

    return thesis


_RESTRICT_INSTRUCTION = """

═══════════════════════════════════════════════════════════════════════════════
VINCOLO DI FONTE (RESTRICT)
═══════════════════════════════════════════════════════════════════════════════
Le informazioni autorevoli per questa tesi sono SOLO quelle nella BASE DI
CONOSCENZA (LLM WIKI) sopra. Non aggiungere fatti, dati o citazioni provenienti
dalla tua conoscenza generale.
- Se un dato necessario non e' presente nel wiki, scrivi "[fonte non disponibile]"
  invece di inventarlo.
- Le citazioni inline ([1], {{nota:...}}) devono riferirsi UNICAMENTE a fonti
  elencate nella sezione "Fonti citate" del wiki.
- Non inventare autori, titoli, anni, DOI, dati statistici. Se il wiki non li
  contiene, ometti la citazione e segnala "[fonte non disponibile]".
"""


def _build_context_for_thesis(thesis: Thesis, db: DBSession, query_extra: str = "") -> str:
    """
    Costruisce il context da iniettare nei prompt (chapters/sections/content).

    Se la tesi ha un wiki gia' ingerito (wiki_status in {ingested, linted}),
    usa il retriever che pesca da wiki/temi/sintesi/concetti/fonti con TF-IDF
    sul query (titolo + key_topics + descrizione + query_extra). Quando
    restrict_to_sources=True appende il blocco "VINCOLO DI FONTE (RESTRICT)"
    direttamente al context: cosi' i prompt builder esistenti non devono
    cambiare firma.

    Altrimenti, fallback al vecchio build_attachments_context (back-compat per
    tesi pre-feature o per chi disabilita lo step Knowledge Base). In quel
    caso il VINCOLO viene comunque appeso se restrict_to_sources=True.
    """
    restrict = bool(thesis.restrict_to_sources)
    wiki_status = (thesis.wiki_status or "none").lower()

    base_context: str = ""
    if wiki_status in ("ingested", "linted"):
        try:
            from llm_wiki import wiki_retriever as _wr
            query_parts = [
                thesis.title or "",
                " ".join(thesis.key_topics or []),
                thesis.description or "",
                query_extra,
            ]
            query = " ".join(p for p in query_parts if p)
            result = _wr.build_context(
                str(thesis.id),
                query=query,
                max_chars=config.THESIS_MAX_CONTEXT_CHARS,
                restrict_to_sources=restrict,
            )
            base_context = result.context
        except Exception:  # noqa: BLE001
            logger.exception(
                "Wiki retriever fallito tesi %s, fallback ad attachments_context",
                thesis.id,
            )

    if not base_context:
        # Fallback: back-compat (tesi pre-LLM-Wiki o wiki non ancora ingerita)
        attachments = db.query(ThesisAttachment).filter(
            ThesisAttachment.thesis_id == thesis.id
        ).all()
        base_context = build_attachments_context(
            [a.to_dict() | {"extracted_text": a.extracted_text} for a in attachments]
        )

    if restrict:
        return base_context + _RESTRICT_INSTRUCTION
    return base_context


def build_thesis_data_dict(thesis: Thesis, db: DBSession) -> dict:
    """Costruisce il dizionario con tutti i dati della tesi per i prompt."""
    data = {
        "title": thesis.title,
        "description": thesis.description,
        "key_topics": thesis.key_topics or [],
        "num_chapters": thesis.num_chapters,
        "sections_per_chapter": thesis.sections_per_chapter,
        "words_per_section": thesis.words_per_section,
        "citation_style": getattr(thesis, 'citation_style', 'footnotes') or 'footnotes',
    }

    # Carica i dati di lookup
    if thesis.writing_style_id:
        style = db.query(WritingStyle).get(thesis.writing_style_id)
        if style:
            data["writing_style_name"] = style.name
            data["writing_style_hint"] = style.prompt_hint or ""

    if thesis.content_depth_id:
        depth = db.query(ContentDepthLevel).get(thesis.content_depth_id)
        if depth:
            data["content_depth_name"] = depth.name

    if thesis.knowledge_level_id:
        level = db.query(AudienceKnowledgeLevel).get(thesis.knowledge_level_id)
        if level:
            data["knowledge_level_name"] = level.name
            data["knowledge_level_hint"] = level.prompt_hint or ""

    if thesis.audience_size_id:
        size = db.query(AudienceSize).get(thesis.audience_size_id)
        if size:
            data["audience_size_name"] = size.name

    if thesis.industry_id:
        industry = db.query(Industry).get(thesis.industry_id)
        if industry:
            data["industry_name"] = industry.name

    if thesis.target_audience_id:
        target = db.query(TargetAudience).get(thesis.target_audience_id)
        if target:
            data["target_audience_name"] = target.name
            data["target_audience_hint"] = target.prompt_hint or ""

    return data


def _build_chapters_from_custom_outline(custom_outline: dict, include_sections: bool = False) -> dict:
    """
    Costruisce `chapters_structure` direttamente dai dati `custom_outline` forniti
    dall'utente, senza chiamare l'AI. Usato come short-circuit quando
    thesis.use_custom_outline == True.

    Args:
        custom_outline: dict con chiave "chapters" -> list di {title, brief_description, sections}.
        include_sections: se True include le sezioni in ogni capitolo (fase 2).
                           Se False (fase 1), restituisce solo titolo+descrizione capitolo.

    Returns:
        dict {"chapters": [...]} compatibile con il formato gia' usato dal frontend.
    """
    chapters_out = []
    for idx, ch in enumerate((custom_outline or {}).get("chapters", []) or [], start=1):
        title = (ch.get("title") or "").strip()
        brief = (ch.get("brief_description") or "").strip()
        entry = {
            "index": idx,
            "chapter_index": idx,
            "title": title,
            "chapter_title": title,
            "brief_description": brief,
            "description": brief,
        }
        if include_sections:
            sections_out = []
            for sidx, sec in enumerate(ch.get("sections", []) or [], start=1):
                sec_title = (sec.get("title") or "").strip()
                kps = [str(kp).strip() for kp in (sec.get("key_points") or []) if str(kp).strip()]
                sections_out.append({
                    "index": sidx,
                    "title": sec_title,
                    "key_points": kps,
                })
            entry["sections"] = sections_out
        chapters_out.append(entry)
    return {"chapters": chapters_out}


# ============================================================================
# LOOKUP ENDPOINTS
# ============================================================================

@router.get("/lookup", response_model=LookupDataResponse)
async def get_all_lookup_data(
    current_user: User = Depends(get_current_active_user),
    db: DBSession = Depends(get_db)
):
    """Restituisce tutti i dati di lookup in una singola chiamata."""
    writing_styles = db.query(WritingStyle).filter(
        WritingStyle.is_active == True
    ).order_by(WritingStyle.sort_order).all()

    content_depths = db.query(ContentDepthLevel).filter(
        ContentDepthLevel.is_active == True
    ).order_by(ContentDepthLevel.sort_order).all()

    knowledge_levels = db.query(AudienceKnowledgeLevel).filter(
        AudienceKnowledgeLevel.is_active == True
    ).order_by(AudienceKnowledgeLevel.sort_order).all()

    audience_sizes = db.query(AudienceSize).filter(
        AudienceSize.is_active == True
    ).order_by(AudienceSize.sort_order).all()

    industries = db.query(Industry).filter(
        Industry.is_active == True
    ).order_by(Industry.sort_order).all()

    target_audiences = db.query(TargetAudience).filter(
        TargetAudience.is_active == True
    ).order_by(TargetAudience.sort_order).all()

    return LookupDataResponse(
        writing_styles=[WritingStyleResponse(**s.to_dict()) for s in writing_styles],
        content_depths=[ContentDepthResponse(**d.to_dict()) for d in content_depths],
        knowledge_levels=[AudienceKnowledgeLevelResponse(**l.to_dict()) for l in knowledge_levels],
        audience_sizes=[AudienceSizeResponse(**s.to_dict()) for s in audience_sizes],
        industries=[IndustryResponse(**i.to_dict()) for i in industries],
        target_audiences=[TargetAudienceResponse(**t.to_dict()) for t in target_audiences]
    )


@router.get("/lookup/writing-styles")
async def get_writing_styles(
    current_user: User = Depends(get_current_active_user),
    db: DBSession = Depends(get_db)
):
    """Restituisce gli stili di scrittura disponibili."""
    styles = db.query(WritingStyle).filter(
        WritingStyle.is_active == True
    ).order_by(WritingStyle.sort_order).all()
    return {"styles": [s.to_dict() for s in styles]}


@router.get("/lookup/content-depths")
async def get_content_depths(
    current_user: User = Depends(get_current_active_user),
    db: DBSession = Depends(get_db)
):
    """Restituisce i livelli di profondità contenuto."""
    levels = db.query(ContentDepthLevel).filter(
        ContentDepthLevel.is_active == True
    ).order_by(ContentDepthLevel.sort_order).all()
    return {"levels": [l.to_dict() for l in levels]}


@router.get("/lookup/knowledge-levels")
async def get_knowledge_levels(
    current_user: User = Depends(get_current_active_user),
    db: DBSession = Depends(get_db)
):
    """Restituisce i livelli di conoscenza del pubblico."""
    levels = db.query(AudienceKnowledgeLevel).filter(
        AudienceKnowledgeLevel.is_active == True
    ).order_by(AudienceKnowledgeLevel.sort_order).all()
    return {"levels": [l.to_dict() for l in levels]}


@router.get("/lookup/audience-sizes")
async def get_audience_sizes(
    current_user: User = Depends(get_current_active_user),
    db: DBSession = Depends(get_db)
):
    """Restituisce le dimensioni del pubblico."""
    sizes = db.query(AudienceSize).filter(
        AudienceSize.is_active == True
    ).order_by(AudienceSize.sort_order).all()
    return {"sizes": [s.to_dict() for s in sizes]}


@router.get("/lookup/industries")
async def get_industries(
    current_user: User = Depends(get_current_active_user),
    db: DBSession = Depends(get_db)
):
    """Restituisce i settori/industrie."""
    industries = db.query(Industry).filter(
        Industry.is_active == True
    ).order_by(Industry.sort_order).all()
    return {"industries": [i.to_dict() for i in industries]}


@router.get("/lookup/target-audiences")
async def get_target_audiences(
    current_user: User = Depends(get_current_active_user),
    db: DBSession = Depends(get_db)
):
    """Restituisce i destinatari target."""
    audiences = db.query(TargetAudience).filter(
        TargetAudience.is_active == True
    ).order_by(TargetAudience.sort_order).all()
    return {"audiences": [a.to_dict() for a in audiences]}


# ============================================================================
# THESIS CRUD ENDPOINTS
# ============================================================================

@router.post("", response_model=ThesisResponse)
async def create_thesis(
    request: ThesisCreateRequest,
    current_user: User = Depends(require_permission('thesis')),
    db: DBSession = Depends(get_db)
):
    """
    Crea una nuova tesi/relazione.

    Richiede tutti i parametri di configurazione.
    Addebita la tariffa flat tesi unica (default 1000 crediti, uguale per tutti
    gli utenti, configurabile da admin). Gli admin StyleForge non vengono addebitati.
    """
    # Verifica sessione se specificata
    session_uuid = None
    if request.session_id:
        session = db.query(Session).filter(
            Session.session_id == request.session_id,
            Session.user_id == current_user.id
        ).first()
        if session:
            session_uuid = session.id

    # Nessun addebito alla creazione: la tesi si paga PER STEP (Capitoli/Sezioni/
    # Contenuto). Gli admin non pagano nulla; le tesi nuove non-admin pagano per step.

    # Indice custom (se l'utente ha scelto la modalita' "definisci tu l'indice")
    use_custom_outline = bool(getattr(request, 'use_custom_outline', False))
    custom_outline_dict = None
    custom_num_chapters = request.num_chapters
    custom_sections_per_chapter = request.sections_per_chapter
    if use_custom_outline and request.custom_outline is not None:
        custom_outline_dict = request.custom_outline.model_dump()
        chapters_in = custom_outline_dict.get("chapters") or []
        if chapters_in:
            custom_num_chapters = len(chapters_in)
            total_sec = sum(len(c.get("sections") or []) for c in chapters_in)
            custom_sections_per_chapter = max(1, round(total_sec / len(chapters_in)))

    # Crea la tesi
    thesis = Thesis(
        user_id=current_user.id,
        session_id=session_uuid,
        title=request.title,
        description=request.description,
        key_topics=request.key_topics,
        writing_style_id=request.writing_style_id,
        content_depth_id=request.content_depth_id,
        num_chapters=custom_num_chapters,
        sections_per_chapter=custom_sections_per_chapter,
        words_per_section=request.words_per_section,
        knowledge_level_id=request.knowledge_level_id,
        audience_size_id=request.audience_size_id,
        industry_id=request.industry_id,
        target_audience_id=request.target_audience_id,
        ai_provider=request.ai_provider.value if request.ai_provider else config.THESIS_AI_PROVIDER,
        citation_style=request.citation_style or "footnotes",
        status='draft',
        # True solo per admin (gratis) -> nessun addebito per step. Le tesi nuove
        # non-admin restano False e pagano per step (Capitoli/Sezioni/Contenuto).
        credits_charged=bool(current_user.is_admin),
        # LLM Wiki: vincolo "solo fonti selezionate" (default ON)
        restrict_to_sources=bool(getattr(request, 'restrict_to_sources', True)),
        wiki_status='none',
        # Indice custom (alternativa ai parametri numerici)
        use_custom_outline=use_custom_outline,
        custom_outline=custom_outline_dict,
    )

    db.add(thesis)
    db.commit()
    db.refresh(thesis)

    return ThesisResponse(**thesis.to_dict())


@router.get("", response_model=ThesisListResponse)
async def list_theses(
    status: Optional[str] = None,
    current_user: User = Depends(get_current_active_user),
    db: DBSession = Depends(get_db)
):
    """Elenca tutte le tesi dell'utente (payload leggero: niente corpo generato)."""
    # Deferisce i campi pesanti: non vengono nemmeno letti dal DB per la lista.
    query = db.query(Thesis).filter(Thesis.user_id == current_user.id).options(
        defer(Thesis.generated_content),
        defer(Thesis.chapters_structure),
        defer(Thesis.custom_outline),
        defer(Thesis.wiki_lint_report),
    )

    if status:
        query = query.filter(Thesis.status == status)

    theses = query.order_by(Thesis.created_at.desc()).all()

    return ThesisListResponse(
        theses=[ThesisResponse(**t.to_summary_dict()) for t in theses],
        total=len(theses)
    )


@router.get("/{thesis_id}", response_model=ThesisResponse)
async def get_thesis(
    thesis_id: str,
    current_user: User = Depends(get_current_active_user),
    db: DBSession = Depends(get_db)
):
    """Ottiene i dettagli di una tesi."""
    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))
    return ThesisResponse(**thesis.to_dict())


@router.delete("/{thesis_id}")
async def delete_thesis(
    thesis_id: str,
    current_user: User = Depends(get_current_active_user),
    db: DBSession = Depends(get_db)
):
    """Elimina una tesi e tutti i suoi dati."""
    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    # Elimina allegati dal filesystem
    cleanup_thesis_attachments(thesis_id)

    # Elimina cartella LLM Wiki (se presente)
    try:
        from llm_wiki import wiki_workspace as _ww
        _ww.cleanup(thesis_id)
    except Exception:  # noqa: BLE001
        logger.exception("Errore cleanup wiki tesi %s", thesis_id)

    # Elimina dal database (cascade eliminerà allegati e job)
    db.delete(thesis)
    db.commit()

    return {"message": "Tesi eliminata con successo"}


# ============================================================================
# ATTACHMENTS ENDPOINTS
# ============================================================================

@router.post("/{thesis_id}/attachments", response_model=ThesisAttachmentsListResponse)
async def upload_attachments(
    thesis_id: str,
    files: List[UploadFile] = File(...),
    current_user: User = Depends(get_current_active_user),
    db: DBSession = Depends(get_db)
):
    """
    Carica allegati per una tesi.

    Supporta PDF, DOCX, TXT. Estrae automaticamente il testo.
    """
    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    # Verifica limite allegati
    existing_count = db.query(ThesisAttachment).filter(
        ThesisAttachment.thesis_id == thesis.id
    ).count()

    if existing_count + len(files) > config.THESIS_MAX_ATTACHMENTS:
        raise HTTPException(
            status_code=400,
            detail=f"Superato il limite di {config.THESIS_MAX_ATTACHMENTS} allegati"
        )

    uploaded = []

    for file in files:
        # Leggi contenuto
        content = await file.read()

        # Salva file
        file_path = save_uploaded_file(content, file.filename, thesis_id)

        try:
            # Processa e estrai testo
            attachment_data = process_attachment(file_path, file.filename)

            # Salva nel database
            attachment = ThesisAttachment(
                thesis_id=thesis.id,
                filename=attachment_data["filename"],
                original_filename=attachment_data["original_filename"],
                file_path=attachment_data["file_path"],
                file_size=attachment_data["file_size"],
                mime_type=attachment_data["mime_type"],
                extracted_text=attachment_data["extracted_text"]
            )

            db.add(attachment)
            db.commit()
            db.refresh(attachment)

            uploaded.append(ThesisAttachmentResponse(**attachment.to_dict()))

        except Exception as e:
            # Se fallisce, elimina il file
            delete_attachment_file(str(file_path))
            raise HTTPException(status_code=400, detail=str(e))

    return ThesisAttachmentsListResponse(
        attachments=uploaded,
        total=len(uploaded)
    )


@router.get("/{thesis_id}/attachments", response_model=ThesisAttachmentsListResponse)
async def list_attachments(
    thesis_id: str,
    current_user: User = Depends(get_current_active_user),
    db: DBSession = Depends(get_db)
):
    """Elenca gli allegati di una tesi."""
    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    attachments = db.query(ThesisAttachment).filter(
        ThesisAttachment.thesis_id == thesis.id
    ).order_by(ThesisAttachment.created_at).all()

    return ThesisAttachmentsListResponse(
        attachments=[ThesisAttachmentResponse(**a.to_dict()) for a in attachments],
        total=len(attachments)
    )


@router.delete("/{thesis_id}/attachments/{attachment_id}")
async def delete_attachment(
    thesis_id: str,
    attachment_id: str,
    current_user: User = Depends(get_current_active_user),
    db: DBSession = Depends(get_db)
):
    """Elimina un allegato."""
    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    attachment = db.query(ThesisAttachment).filter(
        ThesisAttachment.id == attachment_id,
        ThesisAttachment.thesis_id == thesis.id
    ).first()

    if not attachment:
        raise HTTPException(status_code=404, detail="Allegato non trovato")

    # Elimina file
    delete_attachment_file(attachment.file_path)

    # Elimina dal database
    db.delete(attachment)
    db.commit()

    return {"message": "Allegato eliminato con successo"}


@router.post("/{thesis_id}/attachments/urls", response_model=ThesisAttachmentsListResponse)
@limiter.limit(FETCH_LIMIT)
async def add_url_attachments(
    thesis_id: str,
    request: Request,
    payload: ThesisUrlAttachmentRequest,
    current_user: User = Depends(get_current_active_user),
    db: DBSession = Depends(get_db)
):
    """
    Aggiunge URL come allegati alla tesi.

    Scarica il contenuto delle pagine web e lo estrae come testo di riferimento.
    """
    import httpx
    from bs4 import BeautifulSoup

    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    # Verifica limite allegati
    existing_count = db.query(ThesisAttachment).filter(
        ThesisAttachment.thesis_id == thesis.id
    ).count()

    if existing_count + len(payload.urls) > config.THESIS_MAX_ATTACHMENTS:
        raise HTTPException(
            status_code=400,
            detail=f"Superato il limite di {config.THESIS_MAX_ATTACHMENTS} allegati"
        )

    # TUTTI gli URL si validano PRIMA di fetchare qualsiasi cosa, e uno solo
    # rifiutato ferma l'intera richiesta.
    #
    # Non e' pignoleria: validare-e-fetchare in un ciclo unico lascerebbe un
    # oracolo di scansione. Con gli errori silenziati per-URL, i tempi di
    # risposta raccontano la rete interna (connessione rifiutata = subito,
    # filtrata = 20s, aperta = contenuto). Rifiutando prima di ogni I/O, un URL
    # bloccato costa sempre uguale e non dice niente.
    validati = []
    for url in payload.urls:
        try:
            validati.append(ssrf_guard.check_url_shape(url).url)
        except ssrf_guard.SsrfBlocked as e:
            logger.warning("URL allegato rifiutato (%s): %s", e.reason, url)
            raise HTTPException(status_code=400, detail=f"URL non consentito: {url}")

    uploaded = []

    for url in validati:
        try:
            try:
                response = await ssrf_guard.safe_get(
                    url,
                    max_bytes=config.THESIS_URL_MAX_BYTES,
                    deadline_s=config.THESIS_URL_TIMEOUT,
                    headers={"User-Agent": "Mozilla/5.0 (compatible; StyleForge/1.0)"},
                )
            except ssrf_guard.SsrfBlocked as e:
                # La forma era a posto ma la destinazione no (IP interno dietro
                # un nome pubblico, o un redirect verso l'interno).
                logger.warning("Destinazione bloccata dalla guard: %s", e)
                raise HTTPException(status_code=400, detail=f"URL non consentito: {url}")

            if response.status_code >= 400:
                logger.warning("Errore HTTP per URL %s: %s", url, response.status_code)
                continue

            soup = BeautifulSoup(response.text, 'html.parser')

            # Estrai titolo
            og_title = soup.find('meta', property='og:title')
            title = og_title['content'] if og_title and og_title.get('content') else ''
            if not title:
                title_tag = soup.find('title')
                title = title_tag.get_text(strip=True) if title_tag else url

            # original_filename e' String(500): un <title> piu' lungo faceva
            # esplodere l'insert dentro l'except generico qui sotto, e l'URL
            # spariva senza spiegazioni.
            title = title[:500]

            # Estrai contenuto
            content_text = ''
            for selector in ['.entry-content', '.post-content', 'article .content', 'article', 'main']:
                el = soup.select_one(selector)
                if el:
                    for tag in el.find_all(['script', 'style', 'nav', 'aside', 'footer']):
                        tag.decompose()
                    content_text = el.get_text(separator='\n', strip=True)
                    break

            if not content_text:
                body = soup.find('body')
                if body:
                    for tag in body.find_all(['script', 'style', 'nav', 'header', 'footer', 'aside']):
                        tag.decompose()
                    content_text = body.get_text(separator='\n', strip=True)

            if len(content_text) > 8000:
                content_text = content_text[:8000] + "\n[...contenuto troncato...]"

            if not content_text:
                logger.warning(f"Nessun contenuto estratto da URL: {url}")
                continue

            attachment = ThesisAttachment(
                thesis_id=thesis.id,
                filename=f"url_{uuid.uuid4().hex[:8]}.html",
                original_filename=title or url[:500],
                file_path=url,
                file_size=len(content_text),
                mime_type="text/html",
                extracted_text=content_text
            )

            db.add(attachment)
            db.commit()
            db.refresh(attachment)

            uploaded.append(ThesisAttachmentResponse(**attachment.to_dict()))

        except HTTPException:
            # Il 400 di una destinazione bloccata NON va inghiottito dall'except
            # generico qui sotto: rifiutare in silenzio e' esattamente il modo in
            # cui questo endpoint nascondeva la SSRF.
            raise
        except ssrf_guard.GuardError as e:
            logger.warning("Fetch abortito dalla guard per %s: %s", url, e)
        except httpx.HTTPError as e:
            logger.warning(f"Errore HTTP per URL {url}: {e}")
        except Exception as e:
            logger.warning(f"Errore recupero URL {url}: {e}")

    if not uploaded:
        raise HTTPException(
            status_code=400,
            detail="Impossibile recuperare contenuto da nessuno degli URL forniti"
        )

    return ThesisAttachmentsListResponse(
        attachments=uploaded,
        total=len(uploaded)
    )


# ============================================================================
# RESEARCH ENDPOINTS (paper accademici dentro il flusso tesi)
# ============================================================================

@router.post("/{thesis_id}/research/search")
async def thesis_research_search(
    thesis_id: str,
    request: ThesisResearchSearchRequest,
    current_user: User = Depends(require_permission('thesis')),
    db: DBSession = Depends(get_db),
):
    """
    Cerca paper accademici per una tesi specifica.
    Stesso motore di /api/research/search ma gated dal permesso 'thesis'.
    """
    # Verifica ownership della tesi
    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    sources = request.sources
    if sources:
        invalid = [s for s in sources if s not in PROVIDER_REGISTRY]
        if invalid:
            raise HTTPException(
                status_code=400,
                detail=f"Fonti non valide: {', '.join(invalid)}",
            )
    else:
        sources = DEFAULT_SOURCES

    # Ricerca paper a pagamento (gratis per admin e vecchie tesi flat).
    if not thesis.credits_charged:
        credit_estimate = estimate_credits("research_search", {"num_sources": len(sources)}, db=db)
        deduct_credits(
            user=current_user,
            amount=credit_estimate["credits_needed"],
            operation_type="research_search",
            description=f"Ricerca paper per tesi: {request.topic[:80]}",
            db=db,
        )

    import httpx as _httpx
    try:
        result = await run_search_pipeline(
            topic=request.topic.strip(),
            sources=sources,
            filters=request.filters.model_dump() if request.filters else None,
            sort_by=request.sort_by,
            per_provider_limit=request.per_provider_limit,
            final_limit=request.final_limit,
            contact_email=(config.CONTACT_EMAIL or None),
            semantic_scholar_api_key=(config.SEMANTIC_SCHOLAR_API_KEY or None),
        )
    except _httpx.RequestError as e:
        logger.error("Errore di rete nella ricerca paper per tesi: %s", e)
        raise HTTPException(status_code=502, detail="Errore di rete nel contattare i provider accademici")
    except Exception as e:
        logger.exception("Errore imprevisto nella ricerca paper per tesi")
        raise HTTPException(status_code=500, detail=f"Errore nella ricerca: {str(e)[:200]}")

    return result


@router.post("/{thesis_id}/research/summarize", response_model=SummaryResult)
async def thesis_research_summarize(
    thesis_id: str,
    request: ThesisResearchSummarizeRequest,
    current_user: User = Depends(require_permission('thesis')),
    db: DBSession = Depends(get_db),
):
    """
    Genera un riassunto AI per un paper, dentro il wizard tesi.
    """
    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    try:
        paper = UnifiedPaper(**request.paper)
    except Exception:
        raise HTTPException(status_code=400, detail="Paper non valido")

    # Riassunto paper a pagamento (gratis per admin e vecchie tesi flat).
    if not thesis.credits_charged:
        credit_estimate = estimate_credits("research_summary", {}, db=db)
        deduct_credits(
            user=current_user,
            amount=credit_estimate["credits_needed"],
            operation_type="research_summary",
            description=f"Riassunto paper per tesi: {paper.title[:80]}",
            db=db,
        )

    try:
        summary = await summarize_paper(paper)
    except Exception as e:
        logger.exception("Errore riassunto paper (tesi)")
        raise HTTPException(status_code=500, detail=f"Errore nel riassunto: {str(e)[:200]}")

    return summary


@router.post("/{thesis_id}/suggest-paper-keywords", response_model=PaperKeywordSuggestResponse)
async def suggest_paper_keywords(
    thesis_id: str,
    current_user: User = Depends(require_permission('thesis')),
    db: DBSession = Depends(get_db),
):
    """
    Suggerisce 5-8 termini di ricerca utili per cercare paper accademici,
    estratti dai documenti gia' caricati dall'utente (allegati non-paper con
    extracted_text). Pensato per la search bar dello step "Paper" del wizard
    tesi, dopo che l'utente ha caricato i propri documenti nello step "Allegati".
    """
    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    eligible = [
        a for a in db.query(ThesisAttachment).filter(
            ThesisAttachment.thesis_id == thesis.id,
            ThesisAttachment.mime_type != PAPER_MIME_TYPE,
        ).all()
        if a.extracted_text and len(a.extracted_text.strip()) >= 100
    ]

    if not eligible:
        raise HTTPException(
            status_code=400,
            detail="Nessun documento testuale disponibile. Carica almeno un PDF/DOCX/TXT prima di richiedere i suggerimenti.",
        )

    # Cap: massimo 5 documenti per contenere costo e token.
    eligible_used = eligible[:5]

    # Suggerimento keyword a pagamento (gratis per admin e vecchie tesi flat).
    credits_consumed = 0
    if not thesis.credits_charged:
        credit_estimate = estimate_credits(
            "paper_keyword_suggest",
            {"num_attachments": len(eligible_used)},
            db=db,
        )
        credits_consumed = int(credit_estimate.get("credits_needed", 0))
        deduct_credits(
            user=current_user,
            amount=credits_consumed,
            operation_type="paper_keyword_suggest",
            description=f"Suggerimento keyword paper da {len(eligible_used)} documenti",
            db=db,
        )

    try:
        keywords = await extract_paper_keywords_from_attachments(
            [a.extracted_text or "" for a in eligible_used],
            title=thesis.title,
            topics=thesis.key_topics,
        )
    except InsufficientCreditsError:
        raise
    except Exception as e:
        logger.exception("Errore estrazione keyword da allegati")
        raise HTTPException(status_code=500, detail=f"Errore nell'estrazione delle keyword: {str(e)[:200]}")

    return PaperKeywordSuggestResponse(
        thesis_id=str(thesis.id),
        keywords=keywords,
        eligible_attachments_count=len(eligible_used),
        credits_consumed=credits_consumed,
    )


@router.post("/{thesis_id}/attachments/papers", response_model=ThesisAddPapersResponse)
@limiter.limit(FETCH_LIMIT)
async def add_paper_attachments(
    thesis_id: str,
    request: Request,
    payload: ThesisAddPapersRequest,
    current_user: User = Depends(require_permission('thesis')),
    db: DBSession = Depends(get_db),
):
    """
    Salva i paper selezionati come allegati della tesi.

    Dal 2026-05 (introduzione LLM Wiki): NON viene piu' generato un riassunto
    AI per-paper qui. Il riassunto/sintesi e' compito del workflow INGEST del
    wiki (eseguito poi via /wiki/ingest), che opera in modo cross-fonti e
    riconcilia entita'/concetti/contraddizioni. Quindi:

      - extracted_text = solo metadati + abstract (compatto, no chiamata LLM)
      - summary fornito dal client e' ancora accettato per back-compat ma NON
        viene piu' generato server-side; tesi gia' create con summary li usano
        come prima (no rotture)
      - summarized_count e credits_consumed restano per back-compat = 0
    """
    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    existing_count = db.query(ThesisAttachment).filter(
        ThesisAttachment.thesis_id == thesis.id
    ).count()

    if existing_count + len(payload.items) > config.THESIS_MAX_ATTACHMENTS:
        raise HTTPException(
            status_code=400,
            detail=f"Superato il limite di {config.THESIS_MAX_ATTACHMENTS} allegati",
        )

    created: List[ThesisAttachmentResponse] = []

    for item in payload.items:
        try:
            paper = UnifiedPaper(**item.paper)
        except Exception:
            raise HTTPException(status_code=400, detail="Paper non valido nel payload")

        # Se il client ha gia' un summary (vecchio flusso), lo include per
        # massima retro-compatibilita'. Non genera piu' niente server-side.
        summary: Optional[SummaryResult] = None
        if item.summary is not None:
            try:
                summary = SummaryResult(**item.summary)
            except Exception:
                summary = None

        if summary is not None:
            extracted_text = sanitize_text_for_db(render_paper_with_summary(paper, summary))
        else:
            # Solo metadati + abstract: il wiki-ingest ricavera' i takeaway.
            extracted_text = sanitize_text_for_db(paper_to_attachment_text(paper))

        # Placeholder per file_path (la colonna è NOT NULL)
        #
        # full_text_url arriva dal payload del client e la pipeline wiki lo
        # scarica dopo (paper_downloader._download_pdf): senza questo controllo
        # e' una SSRF differita e persistente. Qui si valida la FORMA, senza
        # rete: il DNS puo' cambiare fra ora e il download, quindi la verifica
        # sugli IP la rifa' la guard al momento del fetch. Validare solo al
        # fetch lascerebbe URL avvelenati in tabella; solo qui non proteggerebbe
        # le righe gia' salvate. Servono entrambi.
        if paper.full_text_url:
            try:
                file_path = ssrf_guard.check_url_shape(paper.full_text_url).url
            except ssrf_guard.SsrfBlocked as e:
                logger.warning("URL paper rifiutato: %s", e)
                raise HTTPException(
                    status_code=400,
                    detail=f"URL del paper non consentito: {paper.full_text_url}",
                )
        elif paper.doi:
            file_path = f"doi:{paper.doi}"
        else:
            file_path = f"paper:{paper.id}"

        title = (paper.title or "Paper accademico")[:500]

        attachment = ThesisAttachment(
            thesis_id=thesis.id,
            filename=f"paper_{uuid.uuid4().hex[:8]}.txt",
            original_filename=title,
            file_path=file_path,
            file_size=len(extracted_text),
            mime_type=PAPER_MIME_TYPE,
            extracted_text=extracted_text,
        )
        db.add(attachment)
        db.commit()
        db.refresh(attachment)

        created.append(ThesisAttachmentResponse(**attachment.to_dict()))

    return ThesisAddPapersResponse(
        attachments=created,
        total=len(created),
        summarized_count=0,
        credits_consumed=0,
    )


# ============================================================================
# LLM WIKI (second-brain) ENDPOINTS
# ============================================================================

def _refund_wiki_if_charged(db, thesis_id: str, user_id: str):
    """
    Se l'analisi documenti/paper era stata addebitata, rimborsa la quota e azzera
    il flag (così l'ingest è ri-tentabile). Usato quando l'ingest fallisce.
    """
    try:
        thesis = db.query(Thesis).get(thesis_id)
        if not thesis or not thesis.wiki_charged:
            return
        user = db.query(User).get(user_id)
        n_sources = db.query(ThesisAttachment).filter(ThesisAttachment.thesis_id == thesis_id).count()
        cost = int(estimate_credits('wiki_ingest', {'num_sources': n_sources}, db).get('credits_needed', 0) or 0)
        if user and cost > 0:
            add_credits(
                user, cost,
                f"Rimborso analisi documenti/paper (fallita) - {(thesis.title or '')[:50]}",
                db, transaction_type='refund',
            )
        thesis.wiki_charged = False
        db.commit()
    except Exception:  # noqa: BLE001
        logger.exception("Rimborso analisi KB fallito tesi %s", thesis_id)


def _wiki_ingest_task(thesis_id: str, user_id: str):
    """Background task: scarica paper, materializza upload, lancia ingest+lint."""
    from filelock import Timeout as _LockTimeout

    from llm_wiki import (
        paper_downloader as _pd,
        wiki_runner as _wr,
        wiki_workspace as _ww,
    )

    db = SessionLocal()
    lock = None
    try:
        thesis = db.query(Thesis).get(thesis_id)
        if not thesis:
            logger.error("Wiki ingest: tesi %s non trovata", thesis_id)
            return

        # Acquisisci lock filesystem (non bloccante: se gia' in corso, esci)
        try:
            lock = _ww.acquire_lock(thesis_id, timeout=0)
        except _LockTimeout:
            logger.warning("Wiki ingest: tesi %s gia' in corso (lock attivo)", thesis_id)
            return

        # Bootstrap (idempotente) + snapshot di backup
        _ww.bootstrap(thesis_id)
        backup = _ww.snapshot_wiki(thesis_id)

        # Progresso granulare (per la UI): scrive theses.wiki_progress su DB.
        _started_iso = datetime.utcnow().isoformat()

        def _set_progress(phase, percent, message, files=None):
            try:
                db.query(Thesis).filter(Thesis.id == thesis_id).update(
                    {Thesis.wiki_progress: {
                        "phase": phase,
                        "percent": (int(percent) if percent is not None else None),
                        "message": message,
                        "files": files or [],
                        "started_at": _started_iso,
                        "updated_at": datetime.utcnow().isoformat(),
                    }},
                    synchronize_session=False,
                )
                db.commit()
            except Exception:  # noqa: BLE001
                logger.exception("Wiki progress update fallito tesi %s", thesis_id)
                db.rollback()

        def _clear_progress():
            try:
                db.query(Thesis).filter(Thesis.id == thesis_id).update(
                    {Thesis.wiki_progress: None}, synchronize_session=False)
                db.commit()
            except Exception:  # noqa: BLE001
                db.rollback()

        # 1. Scarica paper / fallback abstract
        _set_progress("download", 2, "Scarico i paper selezionati…")
        attachments = db.query(ThesisAttachment).filter(
            ThesisAttachment.thesis_id == thesis.id
        ).all()
        try:
            _pd.materialize_papers(thesis_id, attachments)
        except Exception:  # noqa: BLE001
            logger.exception("Wiki ingest: errore materialize_papers tesi %s", thesis_id)

        # 2. Materializza upload utente
        _set_progress("prepare", 8, "Preparo i documenti caricati…")
        try:
            _ww.materialize_user_uploads(thesis_id, attachments)
        except Exception:  # noqa: BLE001
            logger.exception("Wiki ingest: errore materialize_user_uploads tesi %s", thesis_id)

        # 3. Aggiorna stato + path
        thesis.wiki_status = "ingesting"
        thesis.wiki_path = str(_ww.get_wiki_root(thesis_id))
        db.commit()

        # 4. INGEST con SDK Anthropic. run_ingest emette progresso 0..100 sui batch;
        # lo rimappiamo sulla fascia 10..90 dell'avanzamento complessivo.
        _set_progress("ingest", 10, "Indicizzo i documenti…")
        try:
            ingest_summary = _wr.run_ingest(
                thesis_id,
                on_progress=lambda pct, msg, files=None: _set_progress(
                    "ingest", 10 + int(pct * 0.8), msg, files,
                ),
            )
        except Exception as e:  # noqa: BLE001
            logger.exception("Wiki ingest fallito tesi %s", thesis_id)
            _clear_progress()
            thesis = db.query(Thesis).get(thesis_id)
            if thesis:
                thesis.wiki_status = "failed"
                thesis.wiki_lint_report = {"error": str(e)[:500]}
                db.commit()
            _refund_wiki_if_charged(db, thesis_id, user_id)
            # Rollback al backup pre-ingest
            if backup is not None:
                try:
                    _ww.restore_snapshot(thesis_id, backup)
                except Exception:  # noqa: BLE001
                    logger.exception("Restore snapshot fallito tesi %s", thesis_id)
            return

        # Se nessuna pagina e' stata creata e ci sono fonti raw, l'ingest e' fallito
        # silenziosamente (es. modello non ha usato i tool). Marca come failed.
        if ingest_summary.sources_count > 0 and ingest_summary.pages_created == 0:
            logger.error(
                "Wiki ingest tesi %s: %d fonti ma 0 pagine create, tool_calls=%d, errors=%s",
                thesis_id, ingest_summary.sources_count,
                ingest_summary.total_tool_calls, ingest_summary.errors,
            )
            _clear_progress()
            thesis = db.query(Thesis).get(thesis_id)
            if thesis:
                thesis.wiki_status = "failed"
                thesis.wiki_lint_report = {
                    "error": "Ingest completato senza scrivere pagine. Il modello non ha "
                             "popolato il wiki. Verifica i log del backend.",
                    "_ingest_summary": {
                        "sources_count": ingest_summary.sources_count,
                        "pages_created": ingest_summary.pages_created,
                        "total_tool_calls": ingest_summary.total_tool_calls,
                        "duration_sec": ingest_summary.duration_sec,
                        "errors": ingest_summary.errors[:5],
                    },
                }
                db.commit()
            _refund_wiki_if_charged(db, thesis_id, user_id)
            if backup is not None:
                try:
                    _ww.restore_snapshot(thesis_id, backup)
                except Exception:  # noqa: BLE001
                    logger.exception("Restore snapshot fallito tesi %s", thesis_id)
            return

        thesis = db.query(Thesis).get(thesis_id)
        if thesis:
            thesis.wiki_status = "ingested"
            thesis.wiki_ingested_at = datetime.utcnow()
            db.commit()

        _ww.cleanup_old_snapshots(thesis_id, keep=2)

        # 5. LINT (chained). Errori lint non rendono failed l'ingest.
        thesis = db.query(Thesis).get(thesis_id)
        if thesis:
            thesis.wiki_status = "linting"
            db.commit()
        _set_progress("lint", 90, "Controllo qualità del wiki…")
        try:
            # Lint + auto-fix: se il lint trova mancanze correggibili, un round di
            # AI le sistema e poi ri-controlla, così il report finale è già pulito.
            report, _fixed = _wr.run_lint_and_autofix(
                thesis_id,
                on_progress=lambda pct, msg, files=None: _set_progress("lint", pct, msg),
            )
            thesis = db.query(Thesis).get(thesis_id)
            if thesis:
                thesis.wiki_status = "linted"
                thesis.wiki_lint_report = report
                thesis.wiki_linted_at = datetime.utcnow()
                db.commit()
            _clear_progress()
        except Exception as e:  # noqa: BLE001
            logger.exception("Wiki lint fallito tesi %s", thesis_id)
            _clear_progress()
            thesis = db.query(Thesis).get(thesis_id)
            if thesis:
                # Resta 'ingested' (utilizzabile, il lint e' opzionale)
                thesis.wiki_status = "ingested"
                thesis.wiki_lint_report = {"error": str(e)[:500]}
                db.commit()
    finally:
        if lock is not None:
            try:
                lock.release()
            except Exception:  # noqa: BLE001
                pass
        db.close()


@router.post("/{thesis_id}/wiki/ingest", response_model=WikiStatusResponse)
async def start_wiki_ingest(
    thesis_id: str,
    request: WikiIngestRequest,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(require_permission('thesis')),
    db: DBSession = Depends(get_db),
):
    """
    Avvia ingest + lint del wiki della tesi (in background).

    - 409 se gia' in corso (wiki_status in ingesting/linting).
    - 200 con stato corrente se gia' completato e force=False.
    - 200 + ripartenza se force=True o se wiki_status in {none, failed, ingested}.
    """
    from llm_wiki import wiki_workspace as _ww

    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    # Stati transitori -> 409
    if thesis.wiki_status in ("ingesting", "linting"):
        raise HTTPException(
            status_code=409,
            detail=f"Wiki gia' in elaborazione (stato={thesis.wiki_status})",
        )

    # Se gia' linted e force=False, non rifare niente
    if thesis.wiki_status == "linted" and not request.force:
        return WikiStatusResponse(
            thesis_id=str(thesis.id),
            wiki_status=ThesisWikiStatus.LINTED,
            wiki_path=thesis.wiki_path,
            sources_count=len(_ww.list_raw_files(thesis_id)),
            pages_count=_ww.count_wiki_pages(thesis_id),
            wiki_ingested_at=thesis.wiki_ingested_at,
            wiki_linted_at=thesis.wiki_linted_at,
        )

    # Addebito dell'analisi documenti/paper (Knowledge Base), una sola volta per
    # tesi. Salta admin/vecchie tesi flat (credits_charged) e se già addebitata.
    # Un eventuale 402 (saldo insufficiente) blocca qui senza cambiare lo stato.
    if not thesis.credits_charged and not thesis.wiki_charged:
        n_sources = db.query(ThesisAttachment).filter(ThesisAttachment.thesis_id == thesis.id).count()
        wiki_estimate = estimate_credits('wiki_ingest', {'num_sources': n_sources}, db=db)
        deduct_credits(
            user=current_user,
            amount=int(wiki_estimate['credits_needed']),
            operation_type='wiki_ingest',
            description=f"Analisi documenti/paper (Knowledge Base) - {thesis.title[:50]}",
            db=db,
        )
        thesis.wiki_charged = True
        db.commit()

    # Reset stato e schedula task
    _now_iso = datetime.utcnow().isoformat()
    thesis.wiki_status = "ingesting"
    thesis.wiki_lint_report = None
    thesis.wiki_progress = {
        "phase": "starting",
        "percent": 0,
        "message": "Avvio dell'indicizzazione…",
        "files": [],
        "started_at": _now_iso,
        "updated_at": _now_iso,
    }
    db.commit()

    background_tasks.add_task(_wiki_ingest_task, str(thesis.id), str(current_user.id))

    return WikiStatusResponse(
        thesis_id=str(thesis.id),
        wiki_status=ThesisWikiStatus.INGESTING,
        wiki_path=thesis.wiki_path,
        sources_count=0,
        pages_count=0,
        progress=thesis.wiki_progress,
    )


@router.post("/{thesis_id}/wiki/lint", response_model=WikiStatusResponse)
async def start_wiki_lint(
    thesis_id: str,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(require_permission('thesis')),
    db: DBSession = Depends(get_db),
):
    """Rilancia il SOLO lint (l'ingest deve essere gia' avvenuto)."""
    from llm_wiki import wiki_workspace as _ww

    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    if thesis.wiki_status not in ("ingested", "linted"):
        raise HTTPException(
            status_code=400,
            detail=f"Wiki non ancora ingestita (stato={thesis.wiki_status})",
        )

    def _lint_only(tid: str):
        from llm_wiki import wiki_runner as _wr
        d = SessionLocal()
        try:
            t = d.query(Thesis).get(tid)
            if not t:
                return
            t.wiki_status = "linting"
            d.commit()
            try:
                report, _fixed = _wr.run_lint_and_autofix(tid)
                t = d.query(Thesis).get(tid)
                if t:
                    t.wiki_status = "linted"
                    t.wiki_lint_report = report
                    t.wiki_linted_at = datetime.utcnow()
                    t.wiki_progress = None
                    d.commit()
            except Exception as e:  # noqa: BLE001
                logger.exception("Lint fallito")
                t = d.query(Thesis).get(tid)
                if t:
                    t.wiki_status = "ingested"
                    t.wiki_lint_report = {"error": str(e)[:500]}
                    t.wiki_progress = None
                    d.commit()
        finally:
            d.close()

    _now_iso = datetime.utcnow().isoformat()
    thesis.wiki_progress = {
        "phase": "lint",
        "percent": 92,
        "message": "Controllo qualità del wiki…",
        "files": [],
        "started_at": _now_iso,
        "updated_at": _now_iso,
    }
    db.commit()

    background_tasks.add_task(_lint_only, str(thesis.id))

    return WikiStatusResponse(
        thesis_id=str(thesis.id),
        wiki_status=ThesisWikiStatus.LINTING,
        wiki_path=thesis.wiki_path,
        sources_count=len(_ww.list_raw_files(thesis_id)),
        pages_count=_ww.count_wiki_pages(thesis_id),
        progress=thesis.wiki_progress,
    )


@router.post("/{thesis_id}/wiki/cancel", response_model=WikiStatusResponse)
async def cancel_wiki_ingest(
    thesis_id: str,
    current_user: User = Depends(require_permission('thesis')),
    db: DBSession = Depends(get_db),
):
    """
    Sblocca manualmente un wiki bloccato in stato ingesting/linting.
    Forza wiki_status='failed' e rilascia il filelock se presente.
    Utile quando il processo backend e' stato killato lasciando lo stato
    transitorio nel DB.
    """
    from llm_wiki import wiki_workspace as _ww

    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    # Forza failed (anche se non e' transitorio: l'admin/utente vuole sbloccare)
    thesis.wiki_status = "failed"
    if not thesis.wiki_lint_report:
        thesis.wiki_lint_report = {"error": "Cancellato manualmente dall'utente"}
    db.commit()

    # Tenta di rilasciare il lock filesystem (best-effort)
    try:
        lock_file = _ww.get_wiki_root(thesis_id) / ".ingest.lock"
        if lock_file.exists():
            lock_file.unlink()
    except Exception:  # noqa: BLE001
        logger.exception("Impossibile rimuovere lock per tesi %s", thesis_id)

    return WikiStatusResponse(
        thesis_id=str(thesis.id),
        wiki_status=ThesisWikiStatus.FAILED,
        wiki_path=thesis.wiki_path,
        sources_count=len(_ww.list_raw_files(thesis_id)) if thesis.wiki_path else 0,
        pages_count=_ww.count_wiki_pages(thesis_id) if thesis.wiki_path else 0,
    )


@router.get("/{thesis_id}/wiki/status", response_model=WikiStatusResponse)
async def get_wiki_status(
    thesis_id: str,
    current_user: User = Depends(require_permission('thesis')),
    db: DBSession = Depends(get_db),
):
    """Stato del wiki: per polling lato client."""
    from llm_wiki import wiki_workspace as _ww

    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    return WikiStatusResponse(
        thesis_id=str(thesis.id),
        wiki_status=ThesisWikiStatus(thesis.wiki_status or "none"),
        wiki_path=thesis.wiki_path,
        sources_count=len(_ww.list_raw_files(thesis_id)) if thesis.wiki_path else 0,
        pages_count=_ww.count_wiki_pages(thesis_id) if thesis.wiki_path else 0,
        progress=thesis.wiki_progress,
        wiki_ingested_at=thesis.wiki_ingested_at,
        wiki_linted_at=thesis.wiki_linted_at,
    )


@router.get("/{thesis_id}/wiki/report", response_model=WikiLintReportResponse)
async def get_wiki_report(
    thesis_id: str,
    current_user: User = Depends(require_permission('thesis')),
    db: DBSession = Depends(get_db),
):
    """Restituisce il lint report del wiki (se disponibile)."""
    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    return WikiLintReportResponse(
        thesis_id=str(thesis.id),
        wiki_status=ThesisWikiStatus(thesis.wiki_status or "none"),
        report=thesis.wiki_lint_report,
        generated_at=thesis.wiki_linted_at,
    )


@router.get("/{thesis_id}/wiki/content", response_model=WikiContentResponse)
async def get_wiki_content(
    thesis_id: str,
    current_user: User = Depends(require_permission('thesis')),
    db: DBSession = Depends(get_db),
):
    """
    Informazioni estratte dai documenti (vista utente): pagine del wiki
    raggruppate per categoria (fonti/entità/concetti/temi/sintesi/domande)
    con titolo, riassunto e tag. Sostituisce il report tecnico per l'utente.
    """
    from llm_wiki import wiki_workspace as _ww

    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))
    data = _ww.read_extracted_content(thesis_id) if thesis.wiki_path else {
        "totals": {"pages": 0, "sources": 0}, "categories": [],
    }
    return WikiContentResponse(
        thesis_id=str(thesis.id),
        wiki_status=ThesisWikiStatus(thesis.wiki_status or "none"),
        totals=data.get("totals", {}),
        categories=data.get("categories", []),
    )


# ============================================================================
# GENERATION ENDPOINTS
# ============================================================================

def generate_chapters_task(thesis_id: str, user_id: str):
    """Task background per generare i capitoli."""
    db = SessionLocal()
    try:
        thesis = db.query(Thesis).get(thesis_id)
        if not thesis:
            return

        # Short-circuit: l'utente ha fornito un indice custom.
        # Skip della chiamata AI, popolazione diretta da custom_outline.
        if getattr(thesis, 'use_custom_outline', False) and thesis.custom_outline:
            logger.info(f"[custom_outline] generate_chapters_task: skip AI, popolo da custom_outline (tesi {thesis_id})")
            result = _build_chapters_from_custom_outline(thesis.custom_outline, include_sections=False)
            thesis.chapters_structure = result
            thesis.status = 'chapters_pending'
            thesis.current_phase = 1
            thesis.num_chapters = len(result.get("chapters", []))
            job = db.query(ThesisGenerationJob).filter(
                ThesisGenerationJob.thesis_id == thesis.id,
                ThesisGenerationJob.phase == 'chapters'
            ).order_by(ThesisGenerationJob.created_at.desc()).first()
            if job:
                job.status = 'completed'
                job.result = json.dumps(result)
                job.completed_at = datetime.utcnow()
            db.commit()
            return

        # Costruisci i dati per il prompt
        thesis_data = build_thesis_data_dict(thesis, db)

        # Costruisci contesto: Wiki retriever se disponibile, altrimenti fallback
        attachments_context = _build_context_for_thesis(thesis, db)

        # Genera capitoli con il provider AI selezionato
        provider = thesis.ai_provider or config.THESIS_AI_PROVIDER
        client = get_ai_client(provider)
        logger.info(f"Generazione capitoli con provider: {provider}")
        result = client.generate_chapters(thesis_data, attachments_context)

        # Salva risultato
        thesis.chapters_structure = result
        thesis.status = 'chapters_pending'
        thesis.current_phase = 1

        # Aggiorna job
        job = db.query(ThesisGenerationJob).filter(
            ThesisGenerationJob.thesis_id == thesis.id,
            ThesisGenerationJob.phase == 'chapters'
        ).order_by(ThesisGenerationJob.created_at.desc()).first()

        if job:
            job.status = 'completed'
            job.result = json.dumps(result)
            job.completed_at = datetime.utcnow()

        db.commit()

    except InsufficientCreditsError as e:
        logger.error(f"Crediti insufficienti durante generazione capitoli: {e.user_message}")
        job = db.query(ThesisGenerationJob).filter(
            ThesisGenerationJob.thesis_id == thesis_id,
            ThesisGenerationJob.phase == 'chapters'
        ).order_by(ThesisGenerationJob.created_at.desc()).first()
        if job:
            job.status = 'failed'
            job.error = f"CREDITI_INSUFFICIENTI: {e.user_message}"
            db.commit()
        thesis = db.query(Thesis).get(thesis_id)
        if thesis:
            thesis.status = 'failed'
            db.commit()

    except Exception as e:
        # Aggiorna job con errore
        job = db.query(ThesisGenerationJob).filter(
            ThesisGenerationJob.thesis_id == thesis_id,
            ThesisGenerationJob.phase == 'chapters'
        ).order_by(ThesisGenerationJob.created_at.desc()).first()

        if job:
            job.status = 'failed'
            job.error = str(e)
            db.commit()

        # Aggiorna stato tesi
        thesis = db.query(Thesis).get(thesis_id)
        if thesis:
            thesis.status = 'failed'
            db.commit()

    finally:
        db.close()


@router.post("/{thesis_id}/generate-chapters")
async def generate_chapters(
    thesis_id: str,
    current_user: User = Depends(require_permission('thesis')),
    db: DBSession = Depends(get_db)
):
    """
    FASE 1: Genera i titoli dei capitoli.

    Utilizza OpenAI per generare l'indice basato sui parametri.
    L'utente potrà modificare i titoli prima di confermare.
    Generazione sincrona per risposta immediata.
    """
    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    if thesis.status not in ['draft', 'failed']:
        raise HTTPException(
            status_code=400,
            detail=f"Impossibile generare capitoli: stato attuale '{thesis.status}'"
        )

    # Short-circuit: indice custom fornito dall'utente.
    # Nessuna chiamata AI, nessun debit di crediti aggiuntivi.
    if getattr(thesis, 'use_custom_outline', False) and thesis.custom_outline:
        try:
            result = _build_chapters_from_custom_outline(thesis.custom_outline, include_sections=False)
            thesis.chapters_structure = result
            thesis.status = 'chapters_pending'
            thesis.current_phase = 1
            thesis.num_chapters = len(result.get("chapters", []))
            db.commit()
            return {
                "thesis_id": str(thesis.id),
                "status": "chapters_pending",
                "chapters": result.get("chapters", []),
                "message": "Capitoli caricati dall'indice personalizzato. Puoi modificarli prima di confermare.",
                "from_custom_outline": True,
            }
        except Exception as e:
            thesis.status = 'failed'
            db.commit()
            raise HTTPException(status_code=500, detail=f"Errore nel caricamento dell'indice personalizzato: {str(e)}")

    # Caratteri allegati per lo scaling del costo capitoli.
    ch_attachments = db.query(ThesisAttachment).filter(ThesisAttachment.thesis_id == thesis.id).all()
    ch_attachment_chars = sum(len(a.extracted_text or '') for a in ch_attachments)

    # Addebito PER STEP (quota fissa + scaling). Salta se "tutto pagato" (admin /
    # vecchie flat) o se la fase è già stata addebitata (idempotenza su retry).
    charged_chapters_now = False
    chapters_cost = 0
    if not thesis.credits_charged and not thesis.chapters_charged:
        credit_estimate = estimate_credits('thesis_chapters', {'attachment_chars': ch_attachment_chars}, db=db)
        chapters_cost = int(credit_estimate['credits_needed'])
        deduct_credits(
            user=current_user,
            amount=chapters_cost,
            operation_type='thesis_chapters',
            description=f"Generazione struttura capitoli - {thesis.title[:50]}",
            db=db,
        )
        thesis.chapters_charged = True
        charged_chapters_now = True
        db.commit()

    try:
        # Costruisci i dati per il prompt
        thesis_data = build_thesis_data_dict(thesis, db)

        # Costruisci contesto: Wiki retriever se disponibile, altrimenti fallback
        attachments_context = _build_context_for_thesis(thesis, db)

        # Genera capitoli con il provider AI selezionato (sincrono)
        provider = thesis.ai_provider or config.THESIS_AI_PROVIDER
        client = get_ai_client(provider)
        logger.info(f"Generazione capitoli (sincrono) con provider: {provider}")
        result = client.generate_chapters(thesis_data, attachments_context)

        # Salva risultato
        thesis.chapters_structure = result
        thesis.status = 'chapters_pending'
        thesis.current_phase = 1
        db.commit()

        # Restituisci i capitoli generati
        chapters = result.get("chapters", [])
        return {
            "thesis_id": str(thesis.id),
            "status": "chapters_pending",
            "chapters": chapters,
            "message": "Capitoli generati con successo. Puoi modificarli prima di confermare."
        }

    except InsufficientCreditsError as e:
        if charged_chapters_now:
            add_credits(current_user, chapters_cost, f"Rimborso generazione capitoli (fallita) - {thesis.title[:50]}", db, transaction_type='refund')
            thesis.chapters_charged = False
        thesis.status = 'failed'
        db.commit()
        raise HTTPException(
            status_code=402,
            detail=e.user_message
        )
    except Exception as e:
        if charged_chapters_now:
            add_credits(current_user, chapters_cost, f"Rimborso generazione capitoli (fallita) - {thesis.title[:50]}", db, transaction_type='refund')
            thesis.chapters_charged = False
        thesis.status = 'failed'
        db.commit()
        raise HTTPException(
            status_code=500,
            detail=f"Errore nella generazione dei capitoli: {str(e)}"
        )


@router.put("/{thesis_id}/chapters")
async def confirm_chapters(
    thesis_id: str,
    request: ConfirmChaptersRequest,
    current_user: User = Depends(get_current_active_user),
    db: DBSession = Depends(get_db)
):
    """
    Conferma i titoli dei capitoli (eventualmente modificati dall'utente).
    """
    logger.info(f"=== CONFERMA CAPITOLI - INIZIO ===")
    logger.info(f"Thesis ID: {thesis_id}")
    logger.info(f"User ID: {current_user.id}")
    logger.info(f"Numero capitoli ricevuti: {len(request.chapters)}")

    # Log dettaglio di ogni capitolo ricevuto
    for i, c in enumerate(request.chapters):
        logger.info(f"  Capitolo {i+1}: title='{c.title}', index={c.index}, "
                     f"brief_desc='{c.brief_description}', desc='{c.description}', "
                     f"sections={c.sections is not None}")

    # Log del payload raw
    try:
        raw_payload = [c.model_dump() for c in request.chapters]
        logger.info(f"Payload raw: {json.dumps(raw_payload, ensure_ascii=False, default=str)[:2000]}")
    except Exception as log_err:
        logger.warning(f"Impossibile loggare payload raw: {log_err}")

    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    logger.info(f"Stato attuale tesi: {thesis.status}")
    logger.info(f"chapters_structure attuale: {json.dumps(thesis.chapters_structure, ensure_ascii=False, default=str)[:1000] if thesis.chapters_structure else 'None'}")

    # Permetti conferma se lo stato è chapters_pending o se già chapters_confirmed (retry)
    if thesis.status not in ['chapters_pending', 'chapters_confirmed']:
        logger.error(f"STATO NON AMMESSO: '{thesis.status}' non è in ['chapters_pending', 'chapters_confirmed']")
        raise HTTPException(
            status_code=400,
            detail=f"Impossibile confermare capitoli: stato attuale '{thesis.status}'"
        )

    try:
        # Aggiorna struttura con i capitoli confermati
        chapters_data = []
        for c in request.chapters:
            chapter_dict = c.model_dump()
            # Assicurati che brief_description sia presente
            if not chapter_dict.get('brief_description') and chapter_dict.get('description'):
                chapter_dict['brief_description'] = chapter_dict['description']
            chapters_data.append(chapter_dict)

        logger.info(f"chapters_data da salvare: {json.dumps(chapters_data, ensure_ascii=False, default=str)[:2000]}")

        thesis.chapters_structure = {"chapters": chapters_data}
        thesis.status = 'chapters_confirmed'
        thesis.num_chapters = len(request.chapters)

        db.commit()
        logger.info(f"=== CONFERMA CAPITOLI - SUCCESSO === tesi {thesis_id}")

        return {"message": "Capitoli confermati con successo", "status": "chapters_confirmed"}

    except Exception as e:
        logger.error(f"=== CONFERMA CAPITOLI - ERRORE === {type(e).__name__}: {str(e)}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        db.rollback()
        raise HTTPException(
            status_code=500,
            detail=f"Errore nel salvataggio dei capitoli: {str(e)}"
        )


def generate_sections_task(thesis_id: str, user_id: str):
    """Task background per generare le sezioni."""
    db = SessionLocal()
    try:
        thesis = db.query(Thesis).get(thesis_id)
        if not thesis:
            return

        # Short-circuit: indice custom -> sezioni gia' definite dall'utente.
        if getattr(thesis, 'use_custom_outline', False) and thesis.custom_outline:
            logger.info(f"[custom_outline] generate_sections_task: skip AI, popolo sezioni da custom_outline (tesi {thesis_id})")
            result = _build_chapters_from_custom_outline(thesis.custom_outline, include_sections=True)
            thesis.chapters_structure = result
            thesis.status = 'sections_pending'
            thesis.current_phase = 2
            chapters_out = result.get("chapters", [])
            if chapters_out:
                total_sec = sum(len(c.get("sections") or []) for c in chapters_out)
                avg_sec = max(1, round(total_sec / len(chapters_out)))
                thesis.sections_per_chapter = avg_sec
            job = db.query(ThesisGenerationJob).filter(
                ThesisGenerationJob.thesis_id == thesis.id,
                ThesisGenerationJob.phase == 'sections'
            ).order_by(ThesisGenerationJob.created_at.desc()).first()
            if job:
                job.status = 'completed'
                job.result = json.dumps(result)
                job.completed_at = datetime.utcnow()
            db.commit()
            return

        # Costruisci dati
        thesis_data = build_thesis_data_dict(thesis, db)
        chapters = thesis.chapters_structure.get("chapters", [])

        # Costruisci contesto: Wiki retriever (query include titoli capitoli) se
        # disponibile, altrimenti fallback ai vecchi extracted_text
        chapter_titles_q = " ".join(
            c.get("chapter_title") or c.get("title") or "" for c in chapters
        )
        attachments_context = _build_context_for_thesis(thesis, db, query_extra=chapter_titles_q)

        # Genera sezioni con il provider AI selezionato
        provider = thesis.ai_provider or config.THESIS_AI_PROVIDER
        client = get_ai_client(provider)
        logger.info(f"Generazione sezioni con provider: {provider}")
        result = client.generate_sections(thesis_data, chapters, attachments_context)

        # Salva risultato
        thesis.chapters_structure = result
        thesis.status = 'sections_pending'
        thesis.current_phase = 2

        # Aggiorna job
        job = db.query(ThesisGenerationJob).filter(
            ThesisGenerationJob.thesis_id == thesis.id,
            ThesisGenerationJob.phase == 'sections'
        ).order_by(ThesisGenerationJob.created_at.desc()).first()

        if job:
            job.status = 'completed'
            job.result = json.dumps(result)
            job.completed_at = datetime.utcnow()

        db.commit()

    except InsufficientCreditsError as e:
        logger.error(f"Crediti insufficienti durante generazione sezioni: {e.user_message}")
        job = db.query(ThesisGenerationJob).filter(
            ThesisGenerationJob.thesis_id == thesis_id,
            ThesisGenerationJob.phase == 'sections'
        ).order_by(ThesisGenerationJob.created_at.desc()).first()
        if job:
            job.status = 'failed'
            job.error = f"CREDITI_INSUFFICIENTI: {e.user_message}"
            db.commit()
        thesis = db.query(Thesis).get(thesis_id)
        if thesis:
            thesis.status = 'failed'
            db.commit()

    except Exception as e:
        job = db.query(ThesisGenerationJob).filter(
            ThesisGenerationJob.thesis_id == thesis_id,
            ThesisGenerationJob.phase == 'sections'
        ).order_by(ThesisGenerationJob.created_at.desc()).first()

        if job:
            job.status = 'failed'
            job.error = str(e)
            db.commit()

        thesis = db.query(Thesis).get(thesis_id)
        if thesis:
            thesis.status = 'failed'
            db.commit()

    finally:
        db.close()


@router.post("/{thesis_id}/generate-sections")
async def generate_sections(
    thesis_id: str,
    current_user: User = Depends(require_permission('thesis')),
    db: DBSession = Depends(get_db)
):
    """
    FASE 2: Genera i titoli delle sezioni per ogni capitolo.

    Richiede che i capitoli siano stati confermati.
    Generazione sincrona per risposta immediata.
    """
    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    if thesis.status != 'chapters_confirmed':
        raise HTTPException(
            status_code=400,
            detail=f"Devi prima confermare i capitoli. Stato attuale: '{thesis.status}'"
        )

    # Short-circuit: indice custom -> bypass AI e debit crediti
    if getattr(thesis, 'use_custom_outline', False) and thesis.custom_outline:
        try:
            result = _build_chapters_from_custom_outline(thesis.custom_outline, include_sections=True)
            thesis.chapters_structure = result
            thesis.status = 'sections_pending'
            thesis.current_phase = 2
            chapters_out = result.get("chapters", [])
            if chapters_out:
                total_sec = sum(len(c.get("sections") or []) for c in chapters_out)
                avg_sec = max(1, round(total_sec / len(chapters_out)))
                thesis.sections_per_chapter = avg_sec
            db.commit()
            return {
                "thesis_id": str(thesis.id),
                "status": "sections_pending",
                "chapters": chapters_out,
                "message": "Sezioni caricate dall'indice personalizzato. Puoi modificarle prima di confermare.",
                "from_custom_outline": True,
            }
        except Exception as e:
            thesis.status = 'failed'
            db.commit()
            raise HTTPException(status_code=500, detail=f"Errore nel caricamento delle sezioni custom: {str(e)}")

    # Addebito PER STEP (quota fissa), idempotente per fase.
    charged_sections_now = False
    sections_cost = 0
    if not thesis.credits_charged and not thesis.sections_charged:
        credit_estimate = estimate_credits('thesis_sections', {'num_chapters': thesis.num_chapters}, db=db)
        sections_cost = int(credit_estimate['credits_needed'])
        deduct_credits(
            user=current_user,
            amount=sections_cost,
            operation_type='thesis_sections',
            description=f"Generazione struttura sezioni - {thesis.title[:50]}",
            db=db,
        )
        thesis.sections_charged = True
        charged_sections_now = True
        db.commit()

    try:
        # Costruisci dati
        thesis_data = build_thesis_data_dict(thesis, db)
        chapters = thesis.chapters_structure.get("chapters", [])

        # Costruisci contesto: Wiki retriever con query mirata ai capitoli
        chapter_titles_q = " ".join(
            c.get("chapter_title") or c.get("title") or "" for c in chapters
        )
        attachments_context = _build_context_for_thesis(thesis, db, query_extra=chapter_titles_q)

        # Genera sezioni con il provider AI selezionato (sincrono)
        provider = thesis.ai_provider or config.THESIS_AI_PROVIDER
        client = get_ai_client(provider)
        logger.info(f"Generazione sezioni (sincrono) con provider: {provider}")
        result = client.generate_sections(thesis_data, chapters, attachments_context)

        # Salva risultato
        thesis.chapters_structure = result
        thesis.status = 'sections_pending'
        thesis.current_phase = 2
        db.commit()

        # Restituisci le sezioni generate
        chapters_with_sections = result.get("chapters", [])
        return {
            "thesis_id": str(thesis.id),
            "status": "sections_pending",
            "chapters": chapters_with_sections,
            "message": "Sezioni generate con successo. Puoi modificarle prima di confermare."
        }

    except InsufficientCreditsError as e:
        if charged_sections_now:
            add_credits(current_user, sections_cost, f"Rimborso generazione sezioni (fallita) - {thesis.title[:50]}", db, transaction_type='refund')
            thesis.sections_charged = False
        thesis.status = 'failed'
        db.commit()
        raise HTTPException(
            status_code=402,
            detail=e.user_message
        )
    except Exception as e:
        if charged_sections_now:
            add_credits(current_user, sections_cost, f"Rimborso generazione sezioni (fallita) - {thesis.title[:50]}", db, transaction_type='refund')
            thesis.sections_charged = False
        thesis.status = 'failed'
        db.commit()
        raise HTTPException(
            status_code=500,
            detail=f"Errore nella generazione delle sezioni: {str(e)}"
        )


@router.put("/{thesis_id}/sections")
async def confirm_sections(
    thesis_id: str,
    request: ConfirmSectionsRequest,
    current_user: User = Depends(get_current_active_user),
    db: DBSession = Depends(get_db)
):
    """
    Conferma i titoli delle sezioni (eventualmente modificati).
    """
    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    if thesis.status != 'sections_pending':
        raise HTTPException(
            status_code=400,
            detail=f"Impossibile confermare sezioni: stato attuale '{thesis.status}'"
        )

    # Aggiorna struttura
    thesis.chapters_structure = {
        "chapters": [c.model_dump() for c in request.chapters]
    }
    thesis.status = 'sections_confirmed'

    db.commit()

    return {"message": "Sezioni confermate con successo", "status": "sections_confirmed"}


def _humanize_content(content: str, trained_session_client, section_name: str = "Sezione") -> str:
    """
    Per la tesi, NON applica l'anti-AI processor.
    Il contenuto accademico deve restare formale e pulito, senza incisi
    informali, trattini o parentetiche tipiche dell'anti-AI.
    Se c'e' una sessione addestrata, applica solo lo stile dell'autore
    tramite Claude (senza post-processing anti-AI).
    """
    if not trained_session_client:
        # Nessuna sessione addestrata: restituisce il contenuto originale
        return content

    try:
        logger.info(f"Applicazione stile autore per: {section_name}")
        # Usa il client Claude della sessione per riscrivere nello stile appreso,
        # ma senza il post-processing anti-AI
        client = trained_session_client
        word_count = len(content.split())
        estimated_tokens = int(word_count * 2.5) + 2000
        dynamic_max_tokens = cap_output_tokens(max(estimated_tokens, 8192))

        style_prompt = f"""Riscrivi il seguente testo accademico applicando lo stile di scrittura
che hai appreso durante l'addestramento. Mantieni il registro formale e accademico.
NON aggiungere incisi informali, trattini, parentetiche colloquiali o espressioni come
"almeno", "diciamo", "va detto", "cioè". Il testo deve restare professionale.
Mantieni INTATTE tutte le citazioni e i riferimenti bibliografici.
Se compaiono token ZZASSETnZZ o ZZMATHnZZ (segnaposto di tabelle/figure/formule),
mantienili INTATTI: quelli su riga isolata restano su riga isolata, quelli in
mezzo a una frase restano ESATTAMENTE nella stessa posizione sintattica della
frase; non rimuoverli, non duplicarli, non spostarli in fondo.
Output SOLO il testo riscritto.

---
{content}
---"""

        client.conversation_history.append({"role": "user", "content": style_prompt})

        try:
            from anthropic import Anthropic
            response = client.client.messages.create(
                model=client.MODEL_ID,
                max_tokens=dynamic_max_tokens,
                system=client.system_prompt,
                messages=client.conversation_history,
                timeout=300.0
            )
            result = response.content[0].text
            client.conversation_history.append({"role": "assistant", "content": result})
            return result
        except Exception:
            client.conversation_history.pop()
            return content
    except InsufficientCreditsError:
        raise
    except Exception as e:
        logger.warning(f"Errore applicazione stile: {e}, uso contenuto originale")
        return content


def _apply_anti_ai(content: str, label: str = "Sezione", target_words: int = 0,
                   seed: Optional[int] = None) -> str:
    """
    Applica gli stage anti-rilevamento AI al contenuto di una tesi:
      0) PARAFRASI CONTROLLATA ricorsiva (DIPPER-style): max diversità lessicale +
         riordino, applicata 2x (leva principale). Flag THESIS_PARAPHRASE_ENABLED.
      1) (legacy/opzionale) riscrittura de-AI accademica via LLM. THESIS_REWRITE_ENABLED.
      2) pass algoritmico register-safe (anti_ai_processor, profilo 'academic').

    La parafrasi e la riscrittura preservano citazioni [x], note {{nota}}, registro
    e lunghezza (floor interno al 90%). NON va usato sulla bibliografia (lista formale).
    """
    import config
    from anti_ai_pipeline import apply_anti_ai_pipeline

    if not getattr(config, 'THESIS_ANTI_AI_ENABLED', True):
        return content
    if not content or not content.strip():
        return content

    # Delega alla pipeline anti-AI condivisa con i flag specifici della Tesi
    # (comportamento invariato: profilo 'academic', stessi default THESIS_*).
    return apply_anti_ai_pipeline(
        content,
        profile=getattr(config, 'THESIS_ANTI_AI_PROFILE', 'academic'),
        target_words=target_words,
        seed=seed,
        paraphrase_enabled=getattr(config, 'THESIS_PARAPHRASE_ENABLED', True),
        paraphrase_rounds=getattr(config, 'THESIS_PARAPHRASE_ROUNDS', 2),
        paraphrase_model=getattr(config, 'THESIS_PARAPHRASE_MODEL', None),
        rewrite_enabled=getattr(config, 'THESIS_REWRITE_ENABLED', False),
        rewrite_model=getattr(config, 'THESIS_REWRITE_MODEL', None),
        algo_enabled=getattr(config, 'THESIS_ALGO_ENABLED', True),
        label=label,
    )


def _ensure_word_count(client, content: str, target_words: int, context_info: str, max_tokens: int) -> str:
    """
    Verifica che il contenuto raggiunga il target di parole.
    Se è sotto il 70%, chiede al modello di continuare ed espandere.
    Effettua al massimo 2 tentativi di continuazione.
    """
    for attempt in range(2):
        # le tabelle/grafici/HINT non contano come prosa: il target è sul testo
        current_words = count_words_excluding_assets(content)
        if current_words >= target_words * 0.70:
            return content

        missing_words = target_words - current_words
        logger.info(
            f"Contenuto troppo corto ({current_words}/{target_words} parole) "
            f"per {context_info}. Tentativo di espansione {attempt + 1}/2..."
        )

        continuation_prompt = f"""Il testo seguente dovrebbe avere ALMENO {target_words} parole, ma ne ha solo circa {current_words}.

⚠️ DEVI aggiungere almeno {missing_words} parole NUOVE per raggiungere il target.

REGOLE:
- Continua il discorso da dove si è interrotto
- NON ripetere concetti già scritti — approfondisci con nuovi dettagli, esempi, analisi
- NON scrivere "in conclusione" o "per riassumere" — stai CONTINUANDO, non chiudendo
- Mantieni lo stesso stile e tono del testo esistente
- Se il testo contiene citazioni [x], mantienile e puoi aggiungerne di nuove (solo fonti REALI)
- Se il testo contiene blocchi [TABELLA]/[GRAFICO] o righe HINT, NON riprodurli
  né aggiungerne di nuovi nella continuazione: scrivi solo prosa
- Se il testo contiene formule tra $...$ o $$...$$, NON riprodurle nella
  continuazione; nuove formule solo se davvero necessarie al discorso

TESTO ESISTENTE DA CONTINUARE:

{content[-3000:]}

═══════════════════════════════════════════════════════════════
SCRIVI la continuazione (almeno {missing_words} parole):"""

        try:
            continuation = client.generate_text(continuation_prompt, max_tokens=max_tokens)
            content = content.rstrip() + "\n\n" + continuation.strip()
        except Exception as e:
            logger.warning(f"Errore nella continuazione: {e}")
            break

    return content


def _refund_content_if_charged(db, thesis_id: str, user_id: str):
    """
    Se il contenuto era stato addebitato, rimborsa la quota e azzera il flag
    (così la generazione contenuto è ri-tentabile). Usato quando il task fallisce.
    """
    try:
        thesis = db.query(Thesis).get(thesis_id)
        if not thesis or not thesis.content_charged:
            return
        user = db.query(User).get(user_id)
        cost = int(estimate_credits('thesis_content', {
            'num_chapters': thesis.num_chapters,
            'sections_per_chapter': thesis.sections_per_chapter,
            'words_per_section': thesis.words_per_section,
        }, db).get('credits_needed', 0) or 0)
        if user and cost > 0:
            add_credits(
                user, cost,
                f"Rimborso generazione contenuto (fallita) - {(thesis.title or '')[:50]}",
                db, transaction_type='refund',
            )
        thesis.content_charged = False
        db.commit()
    except Exception:  # noqa: BLE001
        logger.exception("Rimborso contenuto fallito tesi %s", thesis_id)


def _extract_human_style_examples(thesis, db, max_examples: int = 2, words_each: int = 160) -> str:
    """
    Estrae 1-2 brani brevi di prosa UMANA dalle fonti della tesi (allegati con
    extracted_text o file raw/ del wiki) come ancora di stile (few-shot) nel prompt.
    Ritorna stringa vuota se non ci sono fonti testuali sufficienti.
    """
    texts = []
    try:
        atts = db.query(ThesisAttachment).filter(
            ThesisAttachment.thesis_id == thesis.id,
            ThesisAttachment.extracted_text.isnot(None),
        ).all()
        texts = [a.extracted_text.strip() for a in atts
                 if a.extracted_text and len(a.extracted_text.split()) >= 80]
    except Exception:  # noqa: BLE001
        logger.exception("Esempi stile umano: errore lettura allegati tesi %s", thesis.id)

    if not texts and getattr(thesis, 'wiki_path', None):
        try:
            from llm_wiki import wiki_workspace as _ww
            for p in _ww.list_raw_files(str(thesis.id))[:3]:
                try:
                    t = p.read_text(encoding="utf-8", errors="replace").strip()
                    if len(t.split()) >= 80:
                        texts.append(t)
                except OSError:
                    continue
        except Exception:  # noqa: BLE001
            logger.exception("Esempi stile umano: errore lettura raw wiki tesi %s", thesis.id)

    if not texts:
        return ""

    examples = []
    for t in texts[:max_examples]:
        words = t.split()
        # Brano "centrale" (evita header/abstract iniziali)
        start = min(len(words) // 4, max(0, len(words) - words_each))
        snippet = " ".join(words[start:start + words_each]).strip()
        if snippet:
            examples.append(f"--- Esempio ---\n{snippet}")
    return "\n\n".join(examples)


def generate_content_task(thesis_id: str, user_id: str):
    """Task background per generare il contenuto completo."""
    db = SessionLocal()
    try:
        thesis = db.query(Thesis).get(thesis_id)
        if not thesis:
            return

        thesis_data = build_thesis_data_dict(thesis, db)
        chapters = thesis.chapters_structure.get("chapters", [])

        # Costruisci contesto: Wiki retriever se disponibile, altrimenti fallback
        attachments_context = _build_context_for_thesis(thesis, db)

        # Few-shot di prosa umana (dalle fonti) per un draft meno rilevabile come AI.
        human_examples = _extract_human_style_examples(thesis, db)

        # Verifica se c'è una sessione addestrata per umanizzazione avanzata
        trained_session_client = None
        author_style_context = ""
        if thesis.session_id:
            session = db.query(Session).get(thesis.session_id)
            if session and session.is_trained:
                author_style_context = "Applica lo stile dell'autore appreso durante l'addestramento."
                # Carica il client della sessione addestrata per umanizzazione completa
                try:
                    trained_session_client = session_manager.get_session(
                        session.session_id,
                        user_id
                    )
                    if not trained_session_client.is_trained:
                        trained_session_client = None
                        logger.info(f"Sessione {session.session_id} non addestrata, uso umanizzazione algoritmica")
                except Exception as e:
                    logger.warning(f"Impossibile caricare sessione addestrata: {e}")
                    trained_session_client = None

        # Usa il provider AI selezionato per la generazione contenuto
        provider = thesis.ai_provider or config.THESIS_AI_PROVIDER
        client = get_ai_client(provider)
        logger.info(f"Generazione contenuto con provider: {provider}")

        # Import prompt builders per capitoli speciali
        from thesis_prompts import build_introduction_prompt, build_conclusion_prompt, build_bibliography_prompt

        generated_chapters_content = []
        raw_chapters_content = []  # Contenuto PRE-umanizzazione per la bibliografia
        previous_summary = ""

        # Total: sezioni normali + 3 (introduzione, conclusione, bibliografia)
        total_sections = sum(len(c.get("sections", [])) for c in chapters) + 3
        completed_sections = 0

        # Calcola max_tokens dinamico per le generazioni
        words_per_section = thesis_data.get('words_per_section', 5000)
        dynamic_max_tokens = cap_output_tokens(max(int(words_per_section * 2.5) + 2000, 16000))

        # Raccoglie titoli dei capitoli per i prompt di intro/conclusione
        chapters_titles = [
            c.get('chapter_title') or c.get('title', f"Capitolo {i+1}")
            for i, c in enumerate(chapters)
        ]

        # ===================================================================
        # FASE 1: Genera contenuto dei capitoli normali (con citazioni [x])
        # ===================================================================
        for chapter in chapters:
            chapter_content = f"\n\n# {chapter.get('chapter_title', 'Capitolo')}\n\n"
            raw_chapter_content = ""

            for section in chapter.get("sections", []):
                # Genera contenuto sezione
                raw_content = client.generate_section_content(
                    thesis_data=thesis_data,
                    chapter=chapter,
                    section=section,
                    previous_sections_summary=previous_summary,
                    attachments_context=attachments_context,
                    author_style_context=author_style_context,
                    human_style_examples=human_examples
                )

                # Normalizza gli asset generati (tabelle pareggiate, grafici con
                # JSON rotto degradati a HINT, marcatori orfani rimossi) e la
                # matematica (\(..\)/\[..\] → $/$$, display su riga isolata)
                raw_content = sanitize_generated_assets(raw_content)
                raw_content = sanitize_math_outside_assets(raw_content)

                # Verifica word count e richiedi continuazione se troppo corto
                section_label = f"Cap. {chapter.get('chapter_index', '?')} - {section.get('title', 'Sezione')}"
                raw_content = _ensure_word_count(
                    client, raw_content, words_per_section,
                    section_label, dynamic_max_tokens
                )
                raw_content = sanitize_generated_assets(raw_content)
                raw_content = sanitize_math_outside_assets(raw_content)

                # Salva contenuto raw per la bibliografia (con citazioni [x] intatte)
                raw_chapter_content += f"\n{raw_content}\n"

                # Proteggi tabelle/grafici/HINT e formule con sentinelle: le
                # riscritture LLM (stile autore + anti-AI) non devono toccarli.
                # Prima gli asset (la matematica nelle celle è già sequestrata),
                # poi le formule del testo.
                protected_content, section_asset_map = protect_asset_blocks(raw_content)
                protected_content, section_math_map = protect_math_spans(protected_content)

                # Applica umanizzazione (stile autore, solo se sessione addestrata)
                content = _humanize_content(protected_content, trained_session_client, section.get('title', 'Sezione'))
                # Anti-AI come ultimo layer: riscrittura de-AI accademica + pass algoritmico
                content = _apply_anti_ai(content, section_label, target_words=words_per_section)
                # Ripristina formule e asset al posto delle sentinelle (mai persi)
                content = restore_math_spans(content, section_math_map)
                content = restore_asset_blocks(content, section_asset_map)

                section_text = f"\n## {section.get('title', 'Sezione')}\n\n{content}\n"
                chapter_content += section_text

                # Aggiorna riassunto per coerenza
                if len(content) > 500:
                    previous_summary += f"\n- {section.get('title', 'Sezione')}: {content[:300]}..."

                completed_sections += 1

                # Aggiorna progress
                progress = int((completed_sections / total_sections) * 100)
                thesis.generation_progress = progress
                thesis.total_words_generated += count_words_excluding_assets(content)
                db.commit()

            generated_chapters_content.append(chapter_content)
            raw_chapters_content.append(raw_chapter_content)

        # ===================================================================
        # FASE 2: Genera INTRODUZIONE
        # ===================================================================
        logger.info("Generazione Introduzione...")
        intro_prompt = build_introduction_prompt(
            thesis_data=thesis_data,
            chapters_titles=chapters_titles,
            attachments_context=attachments_context,
            author_style_context=author_style_context
        )
        intro_content = client.generate_text(intro_prompt, max_tokens=dynamic_max_tokens)
        intro_content = sanitize_generated_assets(intro_content)
        intro_content = sanitize_math_outside_assets(intro_content)
        intro_content = _ensure_word_count(
            client, intro_content, words_per_section, "Introduzione", dynamic_max_tokens
        )
        # Protezione difensiva: l'introduzione non dovrebbe contenere asset,
        # ma se il modello ne emette non devono essere mangiati dalle riscritture
        protected_intro, intro_asset_map = protect_asset_blocks(intro_content)
        protected_intro, intro_math_map = protect_math_spans(protected_intro)
        intro_content = _humanize_content(protected_intro, trained_session_client, "Introduzione")
        intro_content = _apply_anti_ai(intro_content, "Introduzione", target_words=words_per_section)
        intro_content = restore_math_spans(intro_content, intro_math_map)
        intro_content = restore_asset_blocks(intro_content, intro_asset_map)

        completed_sections += 1
        progress = int((completed_sections / total_sections) * 100)
        thesis.generation_progress = progress
        thesis.total_words_generated += count_words_excluding_assets(intro_content)
        db.commit()

        # ===================================================================
        # FASE 3: Genera CONCLUSIONE
        # ===================================================================
        logger.info("Generazione Conclusione...")
        # Costruisci riassunto completo per la conclusione
        conclusion_summary = previous_summary
        conclusion_prompt = build_conclusion_prompt(
            thesis_data=thesis_data,
            content_summary=conclusion_summary,
            chapters_titles=chapters_titles,
            author_style_context=author_style_context
        )
        conclusion_content = client.generate_text(conclusion_prompt, max_tokens=dynamic_max_tokens)
        conclusion_content = sanitize_generated_assets(conclusion_content)
        conclusion_content = sanitize_math_outside_assets(conclusion_content)
        conclusion_content = _ensure_word_count(
            client, conclusion_content, words_per_section, "Conclusione", dynamic_max_tokens
        )
        protected_conclusion, conclusion_asset_map = protect_asset_blocks(conclusion_content)
        protected_conclusion, conclusion_math_map = protect_math_spans(protected_conclusion)
        conclusion_content = _humanize_content(protected_conclusion, trained_session_client, "Conclusione")
        conclusion_content = _apply_anti_ai(conclusion_content, "Conclusione", target_words=words_per_section)
        conclusion_content = restore_math_spans(conclusion_content, conclusion_math_map)
        conclusion_content = restore_asset_blocks(conclusion_content, conclusion_asset_map)

        completed_sections += 1
        progress = int((completed_sections / total_sections) * 100)
        thesis.generation_progress = progress
        thesis.total_words_generated += count_words_excluding_assets(conclusion_content)
        db.commit()

        # ===================================================================
        # FASE 4: Genera BIBLIOGRAFIA
        # ===================================================================
        logger.info("Generazione Bibliografia...")
        # Usa il contenuto RAW (pre-umanizzazione) per trovare le citazioni [x]
        # perché l'umanizzazione potrebbe averle alterate.
        # I blocchi asset e le formule vengono esclusi: i JSON dei grafici e i
        # pedici LaTeX contengono [numeri] che falserebbero la regex citazioni.
        all_raw_text, _ = protect_asset_blocks("\n".join(raw_chapters_content))
        all_raw_text, _ = protect_math_spans(all_raw_text)
        # Fallback: se il raw non ha citazioni, prova anche con il contenuto umanizzato
        import re as _re
        raw_citations = _re.findall(r'\[\d+\]', all_raw_text)
        if not raw_citations:
            # Prova con il contenuto umanizzato (l'anti-AI ora preserva le citazioni)
            all_raw_text, _ = protect_asset_blocks("\n".join(generated_chapters_content))
            all_raw_text, _ = protect_math_spans(all_raw_text)
        bibliography_prompt = build_bibliography_prompt(
            thesis_data=thesis_data,
            all_content=all_raw_text
        )
        # Usa sempre Claude per la bibliografia: i modelli OpenAI a volte si rifiutano
        # di generare riferimenti bibliografici ("I'm sorry, I can't provide...")
        try:
            from ai_client import get_ai_client as _get_ai_client
            bib_client = _get_ai_client("claude")
            logger.info("Bibliografia: uso Claude per evitare rifiuti di generazione")
        except Exception as bib_err:
            logger.warning(f"Claude non disponibile per bibliografia, uso provider default: {bib_err}")
            bib_client = client
        bibliography_content = bib_client.generate_text(bibliography_prompt)
        # NON umanizzare la bibliografia (è una lista formale)

        # Verifica che la risposta non sia un rifiuto dell'AI
        refusal_patterns = ["i'm sorry", "i can't", "i cannot", "i apologize", "unable to provide",
                           "not able to", "cannot provide", "can't provide"]
        if any(p in bibliography_content.lower() for p in refusal_patterns):
            logger.warning("Bibliografia: rilevato rifiuto AI, ritento con prompt diretto")
            # Ritenta con un prompt più diretto
            fallback_prompt = (
                f"Genera {len(raw_citations)} voci bibliografiche in formato APA italiano per una tesi su: "
                f"{thesis_data.get('title', '')}. "
                f"Settore: {thesis_data.get('industry_name', 'Generale')}. "
                f"Usa autori e opere REALI e note nel campo, niente fonti inventate.\n"
                f"FORMATO APA OBBLIGATORIO (titolo in corsivo Markdown *titolo*, città: editore):\n"
                f"  LIBRO:    [1] Cognome, N. (Anno). *Titolo dell'opera*. Città: Casa editrice.\n"
                f"  ARTICOLO: [2] Cognome, N. (Anno). Titolo articolo. *Nome Rivista*, vol(num), pp-pp.\n"
                f"  REPORT:   [3] Organizzazione. (Anno). *Titolo report*. Città: Editore.\n"
                f"Esempio: [1] Bonura, A. (2021). *Legislazione e innovazioni normative*. Palermo: USR Sicilia.\n"
                f"Output SOLO la lista, da [1] a [{len(raw_citations)}]."
            )
            bibliography_content = bib_client.generate_text(fallback_prompt)

        completed_sections += 1
        thesis.generation_progress = 100
        db.commit()

        # ===================================================================
        # FASE 5: Assembla contenuto finale
        # ===================================================================
        # Ordine: Introduzione → Capitoli → Conclusione → Bibliografia
        final_content_parts = []

        # Introduzione
        final_content_parts.append(f"\n\n# Introduzione\n\n{intro_content}\n")

        # Capitoli normali
        final_content_parts.extend(generated_chapters_content)

        # Conclusione
        final_content_parts.append(f"\n\n# Conclusione\n\n{conclusion_content}\n")

        # Bibliografia
        final_content_parts.append(f"\n\n# Bibliografia\n\n{bibliography_content}\n")

        # Salva contenuto finale
        thesis.generated_content = "\n".join(final_content_parts)

        # ===================================================================
        # FASE 6: Aggiorna chapters_structure con capitoli speciali per TOC
        # ===================================================================
        updated_chapters = []

        # Introduzione (primo)
        updated_chapters.append({
            "chapter_index": 0,
            "chapter_title": "Introduzione",
            "is_special": True
        })

        # Capitoli normali (rinumerati da 1)
        for i, ch in enumerate(chapters):
            ch_copy = dict(ch)
            ch_copy["chapter_index"] = i + 1
            updated_chapters.append(ch_copy)

        # Conclusione (penultimo)
        updated_chapters.append({
            "chapter_index": len(chapters) + 1,
            "chapter_title": "Conclusione",
            "is_special": True
        })

        # Bibliografia (ultimo)
        updated_chapters.append({
            "chapter_index": len(chapters) + 2,
            "chapter_title": "Bibliografia",
            "is_special": True
        })

        thesis.chapters_structure = {"chapters": updated_chapters}
        thesis.status = 'completed'
        thesis.current_phase = 3
        thesis.generation_progress = 100
        thesis.completed_at = datetime.utcnow()

        # Aggiorna job
        job = db.query(ThesisGenerationJob).filter(
            ThesisGenerationJob.thesis_id == thesis.id,
            ThesisGenerationJob.phase == 'content'
        ).order_by(ThesisGenerationJob.created_at.desc()).first()

        if job:
            job.status = 'completed'
            job.completed_at = datetime.utcnow()

        db.commit()

    except InsufficientCreditsError as e:
        logger.error(f"Crediti insufficienti durante generazione contenuto: {e.user_message}")
        job = db.query(ThesisGenerationJob).filter(
            ThesisGenerationJob.thesis_id == thesis_id,
            ThesisGenerationJob.phase == 'content'
        ).order_by(ThesisGenerationJob.created_at.desc()).first()

        if job:
            job.status = 'failed'
            job.error = f"CREDITI_INSUFFICIENTI: {e.user_message}"
            db.commit()

        thesis = db.query(Thesis).get(thesis_id)
        if thesis:
            thesis.status = 'failed'
            db.commit()
        _refund_content_if_charged(db, thesis_id, user_id)

    except Exception as e:
        job = db.query(ThesisGenerationJob).filter(
            ThesisGenerationJob.thesis_id == thesis_id,
            ThesisGenerationJob.phase == 'content'
        ).order_by(ThesisGenerationJob.created_at.desc()).first()

        if job:
            job.status = 'failed'
            job.error = str(e)
            db.commit()

        thesis = db.query(Thesis).get(thesis_id)
        if thesis:
            thesis.status = 'failed'
            db.commit()
        _refund_content_if_charged(db, thesis_id, user_id)

    finally:
        db.close()


@router.post("/{thesis_id}/generate-content", response_model=StartContentGenerationResponse)
async def start_content_generation(
    thesis_id: str,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(require_permission('thesis')),
    db: DBSession = Depends(get_db)
):
    """
    FASE 3: Avvia la generazione del contenuto.

    Genera ogni sezione una alla volta, applicando umanizzazione.
    """
    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    if thesis.status != 'sections_confirmed':
        raise HTTPException(
            status_code=400,
            detail=f"Devi prima confermare le sezioni. Stato attuale: '{thesis.status}'"
        )

    chapters = thesis.chapters_structure.get("chapters", [])
    total_sections = sum(len(c.get("sections", [])) for c in chapters)

    # Addebito PER STEP (quota fissa), idempotente per fase. Il rimborso in caso di
    # fallimento avviene nel task generate_content_task (vedi gestione 'failed').
    if not thesis.credits_charged and not thesis.content_charged:
        credit_estimate = estimate_credits('thesis_content', {
            'num_chapters': thesis.num_chapters,
            'sections_per_chapter': thesis.sections_per_chapter,
            'words_per_section': thesis.words_per_section,
        }, db=db)
        deduct_credits(
            user=current_user,
            amount=int(credit_estimate['credits_needed']),
            operation_type='thesis_content',
            description=f"Generazione contenuto tesi - {thesis.title[:50]}",
            db=db,
        )
        thesis.content_charged = True

    # Aggiorna stato
    thesis.status = 'generating'
    thesis.current_phase = 3
    thesis.generation_progress = 0

    # Crea job
    job_id = f"thesis_content_{uuid.uuid4().hex[:8]}"
    job = ThesisGenerationJob(
        thesis_id=thesis.id,
        job_id=job_id,
        phase='content',
        status='pending'
    )
    db.add(job)
    db.commit()

    # Avvia task
    background_tasks.add_task(generate_content_task, str(thesis.id), str(current_user.id))

    return StartContentGenerationResponse(
        thesis_id=str(thesis.id),
        job_id=job_id,
        status='generating',
        message="Generazione contenuto avviata",
        total_sections=total_sections
    )


@router.get("/{thesis_id}/generation-status", response_model=GenerationStatusResponse)
async def get_generation_status(
    thesis_id: str,
    current_user: User = Depends(get_current_active_user),
    db: DBSession = Depends(get_db)
):
    """
    Ottiene lo stato dettagliato della generazione.
    """
    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    chapters = thesis.chapters_structure.get("chapters", []) if thesis.chapters_structure else []
    total_sections = sum(len(c.get("sections", [])) for c in chapters)
    completed_sections = int(total_sections * thesis.generation_progress / 100) if thesis.generation_progress else 0

    # Calcola capitolo/sezione corrente
    current_chapter = 0
    current_section = 0
    sections_counted = 0

    for i, ch in enumerate(chapters):
        ch_sections = len(ch.get("sections", []))
        if sections_counted + ch_sections > completed_sections:
            current_chapter = i
            current_section = completed_sections - sections_counted
            break
        sections_counted += ch_sections

    # Costruisci stato per capitolo
    chapters_status = []
    for i, ch in enumerate(chapters):
        ch_sections = len(ch.get("sections", []))
        ch_completed = 0

        if i < current_chapter:
            ch_completed = ch_sections
            ch_status = 'completed'
        elif i == current_chapter:
            ch_completed = current_section
            ch_status = 'in_progress' if thesis.status == 'generating' else 'pending'
        else:
            ch_status = 'pending'

        # Costruisci stato per ogni sezione del capitolo
        sections_status = []
        for j, sec in enumerate(ch.get("sections", [])):
            if i < current_chapter:
                sec_status = 'completed'
            elif i == current_chapter:
                if j < current_section:
                    sec_status = 'completed'
                elif j == current_section and thesis.status == 'generating':
                    sec_status = 'in_progress'
                else:
                    sec_status = 'pending'
            else:
                sec_status = 'pending'

            sections_status.append(SectionGenerationStatus(
                section_index=j,
                title=sec.get("title", f"Sezione {j+1}"),
                status=sec_status,
                words_count=0  # TODO: calcolare parole reali per sezione
            ))

        chapters_status.append(ChapterGenerationStatus(
            chapter_index=i,
            chapter_title=ch.get("chapter_title", f"Capitolo {i+1}"),
            total_sections=ch_sections,
            completed_sections=ch_completed,
            status=ch_status,
            sections=sections_status
        ))

    return GenerationStatusResponse(
        thesis_id=str(thesis.id),
        status=ThesisStatus(thesis.status),
        current_phase=thesis.current_phase,
        generation_progress=thesis.generation_progress,
        current_chapter=current_chapter if thesis.status == 'generating' else None,
        current_section=current_section if thesis.status == 'generating' else None,
        total_sections=total_sections,
        completed_sections=completed_sections,
        chapters=chapters_status,
        estimated_time_remaining=None  # TODO: calcolare in base a media
    )


# ============================================================================
# FOOTNOTE PROCESSING UTILITIES
# ============================================================================

import re as _re

_FOOTNOTE_PATTERN = _re.compile(r'\{\{nota:\s*(.*?)\}\}')
# Markdown italic: *testo* (non greedy), evita ** che è bold
_ITALIC_PATTERN = _re.compile(r'(?<!\*)\*(?!\*)([^\*\n]+?)\*(?!\*)')


def iter_italic_segments(text: str):
    """Splitta `text` in tuple (segment, italic_bool) basandosi sul markdown *italic*.

    Esempio: "Cognome, N. (2021). *Titolo*. Roma: X." →
        [("Cognome, N. (2021). ", False), ("Titolo", True), (". Roma: X.", False)]
    Se non ci sono asterischi, yield un singolo segmento non-italic.
    """
    last = 0
    for m in _ITALIC_PATTERN.finditer(text):
        if m.start() > last:
            yield text[last:m.start()], False
        yield m.group(1), True
        last = m.end()
    if last < len(text):
        yield text[last:], False


def strip_italic_markers(text: str) -> str:
    """Rimuove i marker Markdown *italic* lasciando solo il testo (per export non-Markdown)."""
    return _ITALIC_PATTERN.sub(lambda m: m.group(1), text)


def extract_footnotes_from_line(line: str) -> list:
    """Trova tutte le {{nota: ...}} in una riga. Ritorna lista di (start, end, testo_nota)."""
    return [(m.start(), m.end(), m.group(1).strip()) for m in _FOOTNOTE_PATTERN.finditer(line)]


def strip_footnotes_for_plain(content: str, start_num: int = 1) -> tuple:
    """
    Per export TXT/MD: sostituisce {{nota:...}} con numeri e raccoglie le note.
    Ritorna (testo_processato, lista_note, next_num).

    Le formule vengono protette prima del match: le graffe }} del LaTeX
    (es. \\sqrt{\\frac{a}{b}}) non devono chiudere prematuramente la {{nota:}}.
    """
    notes = []
    num = start_num
    protected, math_map = protect_math_spans(content)

    def replacer(m):
        nonlocal num
        note_text = unprotect_math_spans(m.group(1).strip(), math_map)
        notes.append((num, note_text))
        result = f"[{num}]"
        num += 1
        return result

    processed = unprotect_math_spans(_FOOTNOTE_PATTERN.sub(replacer, protected), math_map)
    return processed, notes, num


# ============================================================================
# EXPORT ENDPOINTS
# ============================================================================

def generate_table_of_contents(chapters_structure: dict, format_type: str = "txt") -> str:
    """
    Genera l'indice della tesi basato sulla struttura dei capitoli.

    Supporta capitoli speciali (Introduzione, Conclusione, Bibliografia)
    che non hanno sezioni e vengono mostrati solo come titolo.

    Args:
        chapters_structure: Dizionario con la struttura dei capitoli
        format_type: "txt", "md" o "pdf"

    Returns:
        Stringa formattata con l'indice
    """
    if not chapters_structure or "chapters" not in chapters_structure:
        return ""

    chapters = chapters_structure.get("chapters", [])
    if not chapters:
        return ""

    if format_type == "md":
        # Formato Markdown
        toc = "## Indice\n\n"
        for ch_idx, chapter in enumerate(chapters):
            ch_title = chapter.get("chapter_title") or chapter.get("title", f"Capitolo {ch_idx + 1}")
            is_special = chapter.get("is_special", False)

            if is_special:
                # Capitoli speciali (Introduzione, Conclusione, Bibliografia)
                toc += f"**{ch_title}**\n\n"
            else:
                ch_num = chapter.get("chapter_index", ch_idx + 1)
                toc += f"**Capitolo {ch_num}: {ch_title}**\n\n"

                sections = chapter.get("sections", [])
                for sec_idx, section in enumerate(sections):
                    sec_num = section.get("index", sec_idx + 1)
                    sec_title = section.get("title", f"Sezione {sec_num}")
                    toc += f"  - {ch_num}.{sec_num}: {sec_title}\n"
                toc += "\n"

        toc += "---\n\n"
        return toc

    else:
        # Formato TXT (anche per PDF e DOCX)
        separator = "═" * 65
        toc = f"{separator}\n"
        toc += "                           INDICE\n"
        toc += f"{separator}\n\n"

        for ch_idx, chapter in enumerate(chapters):
            ch_title = chapter.get("chapter_title") or chapter.get("title", f"Capitolo {ch_idx + 1}")
            is_special = chapter.get("is_special", False)

            if is_special:
                # Capitoli speciali senza sezioni
                toc += f"{ch_title}\n\n"
            else:
                ch_num = chapter.get("chapter_index", ch_idx + 1)
                toc += f"Capitolo {ch_num}: {ch_title}\n"

                sections = chapter.get("sections", [])
                for sec_idx, section in enumerate(sections):
                    sec_num = section.get("index", sec_idx + 1)
                    sec_title = section.get("title", f"Sezione {sec_num}")
                    toc += f"    {ch_num}.{sec_num}: {sec_title}\n"
                toc += "\n"

        toc += f"{separator}\n\n"
        return toc


@router.get("/{thesis_id}/export")
async def export_thesis(
    thesis_id: str,
    format: str = "pdf",
    template_id: str = None,
    current_user: User = Depends(get_current_active_user),
    db: DBSession = Depends(get_db)
):
    """
    Esporta la tesi completata nel formato richiesto.

    Formati supportati: pdf, txt, md, docx
    Include automaticamente l'indice all'inizio del documento.
    """
    thesis = get_thesis_by_id(db, thesis_id, str(current_user.id))

    if thesis.status != 'completed':
        raise HTTPException(
            status_code=400,
            detail=f"La tesi non è ancora completata. Stato: '{thesis.status}'"
        )

    if not thesis.generated_content:
        raise HTTPException(status_code=404, detail="Nessun contenuto generato")

    content = thesis.generated_content
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = "".join(c for c in thesis.title[:50] if c.isalnum() or c in ' _-').strip()
    cit_style = getattr(thesis, 'citation_style', 'footnotes') or 'footnotes'

    # Genera l'indice
    toc = generate_table_of_contents(thesis.chapters_structure, format)

    # Segmenta il contenuto (testo / tabelle / grafici / HINT) e numera gli
    # asset per capitolo; i contenuti senza marcatori restano un solo segmento
    # di testo e tutti i percorsi si comportano come prima.
    segments = parse_segments(content)
    assign_asset_numbers(segments, thesis.chapters_structure or {})
    tables_index = build_tables_index(segments)
    figures_index = build_figures_index(segments)

    if format == "txt":
        # Export TXT con indice e note a piè di pagina come endnotes
        has_footnotes = cit_style == 'footnotes'
        all_notes = []
        next_note_num = 1
        parts = []
        for seg in segments:
            if seg.kind == 'text':
                seg_text = seg.text
                if has_footnotes:
                    seg_text, seg_notes, next_note_num = strip_footnotes_for_plain(seg_text, next_note_num)
                    all_notes.extend(seg_notes)
                # Formule inline in Unicode leggibile (β, ω, √)
                seg_text = inline_math_to_unicode(seg_text)
                # In plain text gli asterischi del corsivo Markdown non hanno senso → rimuovi
                parts.append(strip_italic_markers(seg_text))
            elif seg.kind == 'table':
                t_lines = [format_caption(seg.asset), ''] + table_to_plain_lines(seg.asset)
                if seg.asset.source:
                    t_lines.append(f"Fonte: {seg.asset.source}")
                parts.append('\n'.join(t_lines))
            elif seg.kind == 'chart':
                parts.append(
                    f"{format_caption(seg.asset)}\n[Grafico disponibile negli export PDF/DOCX]"
                )
            elif seg.kind == 'math':
                eq_line = f"    {latex_to_unicode(seg.asset.latex)}"
                if seg.asset.label:
                    eq_line += f"    {seg.asset.label}"
                parts.append(eq_line)
            elif seg.kind == 'hint':
                parts.append(f">>> HINT (da sostituire): {seg.asset.text} <<<")
        processed_content = '\n\n'.join(parts)

        full_content = f"{thesis.title}\n{'=' * len(thesis.title)}\n\n"
        full_content += toc
        for idx_title, idx_entries in (("INDICE DELLE TABELLE", tables_index),
                                       ("INDICE DELLE FIGURE", figures_index)):
            if idx_entries:
                full_content += f"{idx_title}\n"
                for label, caption in idx_entries:
                    full_content += f"    {label} – {caption}\n"
                full_content += "\n"
        full_content += processed_content
        if all_notes:
            full_content += "\n\n" + "=" * 60 + "\nNOTE\n" + "=" * 60 + "\n\n"
            for num, note_text in all_notes:
                full_content += f"[{num}] {strip_italic_markers(inline_math_to_unicode(note_text))}\n"

        file_path = config.RESULTS_DIR / f"thesis_{safe_title}_{timestamp}.txt"
        file_path.write_text(full_content, encoding='utf-8')

        return FileResponse(
            path=file_path,
            filename=f"tesi_{safe_title}.txt",
            media_type="text/plain"
        )

    elif format == "md":
        # Export Markdown con indice e note come footnotes
        has_footnotes = cit_style == 'footnotes'
        all_notes = []
        next_note_num = 1
        parts = []
        for seg in segments:
            if seg.kind == 'text':
                seg_text = seg.text
                if has_footnotes:
                    seg_text, seg_notes, next_note_num = strip_footnotes_for_plain(seg_text, next_note_num)
                    all_notes.extend(seg_notes)
                parts.append(seg_text)
            elif seg.kind == 'table':
                block = f"**{format_caption(seg.asset)}**\n\n{table_to_markdown(seg.asset)}"
                if seg.asset.source:
                    block += f"\n\n*Fonte: {seg.asset.source}*"
                parts.append(block)
            elif seg.kind == 'chart':
                block = f"**{format_caption(seg.asset)}**"
                if seg.asset.spec and not seg.asset.error:
                    # fallback leggibile: i dati del grafico come tabella markdown
                    spec = seg.asset.spec
                    labels = [str(l) for l in spec.get('labels', [])]
                    header = '| ' + ' | '.join([''] + [s.get('name') or f"Serie {i+1}"
                                                       for i, s in enumerate(spec.get('series', []))]) + ' |'
                    sep = '|' + '|'.join([' --- '] * (len(spec.get('series', [])) + 1)) + '|'
                    data_rows = []
                    for li, lab in enumerate(labels):
                        row = [lab]
                        for s in spec.get('series', []):
                            vals = s.get('values', [])
                            row.append(str(vals[li]) if li < len(vals) else '')
                        data_rows.append('| ' + ' | '.join(row) + ' |')
                    block += '\n\n' + '\n'.join([header, sep] + data_rows)
                    if spec.get('source'):
                        block += f"\n\n*Fonte: {spec['source']}*"
                block += "\n\n*Grafico renderizzato negli export PDF e DOCX.*"
                parts.append(block)
            elif seg.kind == 'math':
                # $..$ / $$..$$ sono math markdown standard: sorgente verbatim
                tag = f" \\tag{{{seg.asset.label.strip('()')}}}" if seg.asset.label else ""
                parts.append(f"$$ {seg.asset.latex}{tag} $$")
            elif seg.kind == 'hint':
                parts.append(f"> ⚠️ **HINT (da sostituire):** {seg.asset.text}")
        processed_content = '\n\n'.join(parts)

        md_content = f"# {thesis.title}\n\n"
        md_content += toc
        for idx_title, idx_entries in (("Indice delle tabelle", tables_index),
                                       ("Indice delle figure", figures_index)):
            if idx_entries:
                md_content += f"## {idx_title}\n\n"
                for label, caption in idx_entries:
                    md_content += f"- **{label}** – {caption}\n"
                md_content += "\n"
        md_content += processed_content
        if all_notes:
            md_content += "\n\n---\n\n### Note\n\n"
            for num, note_text in all_notes:
                md_content += f"[^{num}]: {note_text}\n\n"

        file_path = config.RESULTS_DIR / f"thesis_{safe_title}_{timestamp}.md"
        file_path.write_text(md_content, encoding='utf-8')

        return FileResponse(
            path=file_path,
            filename=f"tesi_{safe_title}.md",
            media_type="text/markdown"
        )

    elif format == "docx":
        # Export DOCX con indice — usa template (23 parametri)
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches, Cm, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from lxml import etree

        # Footnote tracking for DOCX
        docx_footnote_id = [1]  # Mutable counter
        docx_all_footnotes = []  # Collect all footnotes for endnotes fallback

        def _ensure_footnotes_part(doc):
            """Crea o ottieni la FootnotesPart per il documento."""
            from docx.opc.part import Part as OpcPart
            from docx.opc.packuri import PackURI

            # Cerca se esiste già una relazione footnotes
            FOOTNOTES_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
            for rel in doc.part.rels.values():
                if rel.reltype == FOOTNOTES_REL_TYPE:
                    return rel.target_part

            # Crea la footnotes part
            footnotes_xml = (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
                '<w:footnote w:type="separator" w:id="-1">'
                '<w:p><w:r><w:separator/></w:r></w:p>'
                '</w:footnote>'
                '<w:footnote w:type="continuationSeparator" w:id="0">'
                '<w:p><w:r><w:continuationSeparator/></w:r></w:p>'
                '</w:footnote>'
                '</w:footnotes>'
            )
            footnotes_part = OpcPart(
                PackURI('/word/footnotes.xml'),
                'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml',
                footnotes_xml.encode('utf-8'),
                doc.part.package
            )
            doc.part.relate_to(footnotes_part, FOOTNOTES_REL_TYPE)
            return footnotes_part

        def add_footnote(doc, paragraph, footnote_text, footnote_id, fn_font_name="Times New Roman", fn_font_size=10):
            """Aggiunge una footnote reale al documento DOCX."""
            try:
                footnotes_part = _ensure_footnotes_part(doc)
                fns_element = etree.fromstring(footnotes_part.blob)

                # Crea l'elemento footnote
                footnote_el = OxmlElement('w:footnote')
                footnote_el.set(qn('w:id'), str(footnote_id))

                # Paragrafo nella footnote
                fn_para = OxmlElement('w:p')

                # Run con il numero della footnote (nella footnote stessa)
                fn_ref_run = OxmlElement('w:r')
                fn_ref_rPr = OxmlElement('w:rPr')
                fn_ref_style = OxmlElement('w:rStyle')
                fn_ref_style.set(qn('w:val'), 'FootnoteReference')
                fn_ref_rPr.append(fn_ref_style)
                fn_ref_run.append(fn_ref_rPr)
                fn_ref_elem = OxmlElement('w:footnoteRef')
                fn_ref_run.append(fn_ref_elem)
                fn_para.append(fn_ref_run)

                # Spazio dopo il numero
                space_run = OxmlElement('w:r')
                space_t = OxmlElement('w:t')
                space_t.set(qn('xml:space'), 'preserve')
                space_t.text = ' '
                space_run.append(space_t)
                fn_para.append(space_run)

                # Testo della nota (con supporto Markdown *italic* per i titoli APA)
                for seg_text, seg_italic in iter_italic_segments(footnote_text):
                    if not seg_text:
                        continue
                    fn_seg_run = OxmlElement('w:r')
                    fn_seg_rPr = OxmlElement('w:rPr')
                    fn_seg_sz = OxmlElement('w:sz')
                    fn_seg_sz.set(qn('w:val'), str(fn_font_size * 2))  # half-points
                    fn_seg_rPr.append(fn_seg_sz)
                    fn_seg_szCs = OxmlElement('w:szCs')
                    fn_seg_szCs.set(qn('w:val'), str(fn_font_size * 2))
                    fn_seg_rPr.append(fn_seg_szCs)
                    if fn_font_name:
                        fn_seg_rFonts = OxmlElement('w:rFonts')
                        fn_seg_rFonts.set(qn('w:ascii'), fn_font_name)
                        fn_seg_rFonts.set(qn('w:hAnsi'), fn_font_name)
                        fn_seg_rPr.append(fn_seg_rFonts)
                    if seg_italic:
                        fn_seg_i = OxmlElement('w:i')
                        fn_seg_rPr.append(fn_seg_i)
                    fn_seg_run.append(fn_seg_rPr)
                    fn_seg_t = OxmlElement('w:t')
                    fn_seg_t.set(qn('xml:space'), 'preserve')
                    fn_seg_t.text = seg_text
                    fn_seg_run.append(fn_seg_t)
                    fn_para.append(fn_seg_run)

                footnote_el.append(fn_para)
                fns_element.append(footnote_el)

                # Aggiorna il blob
                footnotes_part._blob = etree.tostring(fns_element, xml_declaration=True, encoding='UTF-8', standalone=True)

                # Aggiungi il riferimento nel paragrafo del documento
                fn_inline_run = OxmlElement('w:r')
                fn_inline_rPr = OxmlElement('w:rPr')
                fn_inline_style = OxmlElement('w:rStyle')
                fn_inline_style.set(qn('w:val'), 'FootnoteReference')
                fn_inline_rPr.append(fn_inline_style)
                fn_inline_run.append(fn_inline_rPr)
                fn_inline_ref = OxmlElement('w:footnoteReference')
                fn_inline_ref.set(qn('w:id'), str(footnote_id))
                fn_inline_run.append(fn_inline_ref)
                paragraph._element.append(fn_inline_run)

            except Exception:
                raise  # Let the caller handle the fallback

        template = get_template_by_id(template_id, db)
        ds = template.get("docx", {})

        # Parametri base
        font_name = ds.get("font_name", "Times New Roman")
        font_sz = ds.get("font_size", 12)
        font_title_sz = ds.get("font_title_size", 26)
        title_align_str = ds.get("title_alignment", "center")
        body_align_str = ds.get("body_alignment", "left")
        line_sp = ds.get("line_spacing", 1.5)
        para_sp_after = ds.get("paragraph_spacing_after", 6)
        chapter_sp_before = ds.get("chapter_spacing_before", 18)
        section_sp_before = ds.get("section_spacing_before", 12)
        include_toc_docx = ds.get("include_toc", True)
        include_page_nums = ds.get("include_page_numbers", True)
        page_num_pos = ds.get("page_number_position", "bottom_center")
        toc_indent_val = ds.get("toc_indent", 0.5)
        h1_size = ds.get("heading1_size", 16)
        h2_size = ds.get("heading2_size", 14)

        # Margini
        margin_top = ds.get("margin_top", 72)
        margin_bottom = ds.get("margin_bottom", 72)
        margin_left = ds.get("margin_left", 72)
        margin_right = ds.get("margin_right", 72)

        # Header/Footer
        include_header = ds.get("include_header", False)
        header_text = ds.get("header_text", "")
        include_footer = ds.get("include_footer", False)
        footer_text = ds.get("footer_text", "")

        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        title_alignment = align_map.get(title_align_str, WD_ALIGN_PARAGRAPH.CENTER)
        body_alignment = align_map.get(body_align_str, WD_ALIGN_PARAGRAPH.LEFT)

        file_path = config.RESULTS_DIR / f"thesis_{safe_title}_{timestamp}.docx"

        doc = DocxDocument()

        # ── Margini pagina ──
        section_doc = doc.sections[0]
        section_doc.top_margin = Pt(margin_top)
        section_doc.bottom_margin = Pt(margin_bottom)
        section_doc.left_margin = Pt(margin_left)
        section_doc.right_margin = Pt(margin_right)

        # Imposta stile Normal
        style = doc.styles['Normal']
        style_font = style.font
        style_font.name = font_name
        style_font.size = Pt(font_sz)
        style.paragraph_format.line_spacing = line_sp

        # Imposta font Heading 1
        try:
            h1_style = doc.styles['Heading 1']
            h1_style.font.name = font_name
            h1_style.font.size = Pt(h1_size)
            h1_style.paragraph_format.space_before = Pt(chapter_sp_before)
        except Exception:
            pass

        # Imposta font Heading 2
        try:
            h2_style = doc.styles['Heading 2']
            h2_style.font.name = font_name
            h2_style.font.size = Pt(h2_size)
            h2_style.paragraph_format.space_before = Pt(section_sp_before)
        except Exception:
            pass

        # ── Helper: inserisce campo PAGE in un paragrafo ──
        def _add_page_field(paragraph, pg_font_name, pg_font_size=9):
            run = paragraph.add_run()
            fld_begin = OxmlElement('w:fldChar')
            fld_begin.set(qn('w:fldCharType'), 'begin')
            run._r.append(fld_begin)
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = ' PAGE '
            run._r.append(instr)
            fld_end = OxmlElement('w:fldChar')
            fld_end.set(qn('w:fldCharType'), 'end')
            run._r.append(fld_end)
            for r in paragraph.runs:
                r.font.name = pg_font_name
                r.font.size = Pt(pg_font_size)

        # ── Numeri di pagina (con posizione configurabile) ──
        if include_page_nums:
            try:
                is_top = page_num_pos.startswith("top")
                is_right = page_num_pos.endswith("right")
                pg_align = WD_ALIGN_PARAGRAPH.RIGHT if is_right else WD_ALIGN_PARAGRAPH.CENTER

                if is_top:
                    target = section_doc.header
                    target.is_linked_to_previous = False
                    pg_para = target.paragraphs[0] if target.paragraphs else target.add_paragraph()
                    pg_para.alignment = pg_align
                    _add_page_field(pg_para, font_name)
                else:
                    target = section_doc.footer
                    target.is_linked_to_previous = False
                    pg_para = target.paragraphs[0] if target.paragraphs else target.add_paragraph()
                    pg_para.alignment = pg_align
                    _add_page_field(pg_para, font_name)
            except Exception:
                pass

        # ── Intestazione (header text) ──
        if include_header and header_text:
            try:
                section_doc.header.is_linked_to_previous = False
                # Se numeri pagina sono in alto, aggiungi testo su una riga separata
                if include_page_nums and page_num_pos.startswith("top"):
                    h_para = section_doc.header.add_paragraph()
                else:
                    h_para = section_doc.header.paragraphs[0] if section_doc.header.paragraphs else section_doc.header.add_paragraph()
                h_run = h_para.add_run(header_text)
                h_run.font.name = font_name
                h_run.font.size = Pt(9)
                h_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

        # ── Pie' di pagina (footer text) ──
        if include_footer and footer_text:
            try:
                section_doc.footer.is_linked_to_previous = False
                # Se numeri pagina sono in basso, aggiungi testo su una riga separata
                if include_page_nums and page_num_pos.startswith("bottom"):
                    f_para = section_doc.footer.add_paragraph()
                else:
                    f_para = section_doc.footer.paragraphs[0] if section_doc.footer.paragraphs else section_doc.footer.add_paragraph()
                f_run = f_para.add_run(footer_text)
                f_run.font.name = font_name
                f_run.font.size = Pt(9)
                f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

        # ── Titolo principale ──
        title_para = doc.add_heading(thesis.title, level=0)
        title_para.alignment = title_alignment
        for run in title_para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_title_sz)

        doc.add_paragraph()  # Spazio dopo titolo

        # Indice
        chapters_for_toc = thesis.chapters_structure.get("chapters", []) if thesis.chapters_structure else []
        if chapters_for_toc and include_toc_docx:
            toc_heading = doc.add_heading('Indice', level=1)
            for run in toc_heading.runs:
                run.font.name = font_name

            for ch_idx, chapter in enumerate(chapters_for_toc):
                ch_title = chapter.get("chapter_title") or chapter.get("title", f"Capitolo {ch_idx + 1}")
                is_special = chapter.get("is_special", False)

                if is_special:
                    toc_para = doc.add_paragraph()
                    toc_run = toc_para.add_run(ch_title)
                    toc_run.bold = True
                    toc_run.font.size = Pt(font_sz - 1)
                    toc_run.font.name = font_name
                else:
                    ch_num = chapter.get("chapter_index", ch_idx + 1)
                    toc_para = doc.add_paragraph()
                    toc_run = toc_para.add_run(f"Capitolo {ch_num}: {ch_title}")
                    toc_run.bold = True
                    toc_run.font.size = Pt(font_sz - 1)
                    toc_run.font.name = font_name

                    sections = chapter.get("sections", [])
                    for sec_idx, section_item in enumerate(sections):
                        sec_num = section_item.get("index", sec_idx + 1)
                        sec_title = section_item.get("title", f"Sezione {sec_num}")
                        sec_para = doc.add_paragraph(
                            f"    {ch_num}.{sec_num}: {sec_title}",
                            style='List Bullet'
                        )
                        sec_para.paragraph_format.left_indent = Inches(toc_indent_val)
                        for run in sec_para.runs:
                            run.font.size = Pt(font_sz - 2)
                            run.font.name = font_name

            # ── Indici delle tabelle/figure (solo se presenti) ──
            for idx_title, idx_entries in (("Indice delle tabelle", tables_index),
                                           ("Indice delle figure", figures_index)):
                if not idx_entries:
                    continue
                idx_heading = doc.add_heading(idx_title, level=1)
                for run in idx_heading.runs:
                    run.font.name = font_name
                for label, caption in idx_entries:
                    idx_para = doc.add_paragraph()
                    idx_run = idx_para.add_run(f"{label} – {caption}")
                    idx_run.font.size = Pt(font_sz - 1)
                    idx_run.font.name = font_name

            doc.add_page_break()

        # ── Contenuto con body_alignment e footnotes ──
        def _append_math_run(para, latex):
            """Formula inline: OMML nativo → PNG mathtext → testo Unicode corsivo."""
            try:
                para._p.append(latex_to_omml(latex))
                return
            except MathRenderError as e:
                logger.warning(f"Formula inline non convertita in OMML: {e}")
            try:
                from io import BytesIO
                mp = render_math_png(latex, fontsize=float(font_sz), dpi=300)
                para.add_run().add_picture(BytesIO(mp.png), width=Pt(min(mp.width_pt, 420.0)))
                return
            except MathRenderError as e:
                logger.warning(f"Formula inline non renderizzata in PNG: {e}")
            run = para.add_run(latex_to_unicode(latex))
            run.font.name = font_name
            run.font.size = Pt(font_sz)
            run.italic = True

        def render_docx_text_line(line):
            if line.startswith('# '):
                h = doc.add_heading(inline_math_to_unicode(line[2:]), level=1)
                h.paragraph_format.space_before = Pt(chapter_sp_before)
                for run in h.runs:
                    run.font.name = font_name
            elif line.startswith('## '):
                h = doc.add_heading(inline_math_to_unicode(line[3:]), level=2)
                h.paragraph_format.space_before = Pt(section_sp_before)
                for run in h.runs:
                    run.font.name = font_name
            elif line.strip():
                # Estrazione note sulla riga con formule PROTETTE: le graffe
                # }} del LaTeX (\frac{a}{b}}) non devono chiudere la {{nota:}}
                pline, line_math_map = protect_math_spans(line)
                footnotes_in_line = extract_footnotes_from_line(pline)

                def _add_runs_with_italic(para, text):
                    """Run del paragrafo: Markdown *italic* + formule $...$ (OMML).

                    Le formule vengono estratte PRIMA dello split del corsivo,
                    così un asterisco dentro la matematica non viene
                    interpretato come marcatore.
                    """
                    for kind, latex, raw in iter_inline_math(text):
                        if kind == 'math':
                            _append_math_run(para, latex)
                            continue
                        for seg_text, seg_italic in iter_italic_segments(raw):
                            if not seg_text:
                                continue
                            run = para.add_run(seg_text)
                            run.font.name = font_name
                            run.font.size = Pt(font_sz)
                            if seg_italic:
                                run.italic = True

                if footnotes_in_line:
                    para = doc.add_paragraph()
                    para.alignment = body_alignment
                    para.paragraph_format.space_after = Pt(para_sp_after)
                    para.paragraph_format.line_spacing = line_sp

                    last_end = 0
                    for fn_start, fn_end, fn_text in footnotes_in_line:
                        # Testo prima della nota (con eventuale italic Markdown)
                        before_text = unprotect_math_spans(pline[last_end:fn_start], line_math_map)
                        if before_text:
                            _add_runs_with_italic(para, before_text)
                        # Aggiungi la footnote (italic gestito dentro add_footnote)
                        fn_text = unprotect_math_spans(fn_text, line_math_map)
                        try:
                            # Eventuale matematica nel testo della nota → Unicode
                            add_footnote(doc, para, inline_math_to_unicode(fn_text),
                                         docx_footnote_id[0], font_name, font_sz - 2)
                            docx_footnote_id[0] += 1
                        except Exception:
                            # Fallback: aggiungi come testo in apice
                            sup_run = para.add_run(f"[{docx_footnote_id[0]}]")
                            sup_run.font.name = font_name
                            sup_run.font.size = Pt(font_sz - 2)
                            sup_run.font.superscript = True
                            docx_footnote_id[0] += 1
                        last_end = fn_end

                    # Testo dopo l'ultima nota
                    remaining = unprotect_math_spans(pline[last_end:], line_math_map)
                    if remaining:
                        _add_runs_with_italic(para, remaining)
                else:
                    para = doc.add_paragraph()
                    para.alignment = body_alignment
                    para.paragraph_format.space_after = Pt(para_sp_after)
                    para.paragraph_format.line_spacing = line_sp
                    _add_runs_with_italic(para, line)
            # Righe vuote: non aggiungere nulla (spazio naturale)

        for seg in segments:
            if seg.kind == 'text':
                for line in seg.text.split('\n'):
                    render_docx_text_line(line)
                continue
            # Un asset malformato non deve MAI far fallire l'export:
            # degrada a riquadro HINT.
            try:
                if seg.kind == 'table':
                    add_docx_table(doc, seg.asset, ds)
                elif seg.kind == 'chart':
                    add_docx_chart(doc, seg.asset, ds)
                elif seg.kind == 'math':
                    add_docx_math(doc, seg.asset, ds)
                elif seg.kind == 'hint':
                    add_docx_hint(doc, seg.asset, ds)
            except Exception as e:  # noqa: BLE001
                logger.warning(f"Asset non renderizzabile nel DOCX: {e}")
                try:
                    add_docx_hint(doc, HintAsset(
                        text=f"Elemento non renderizzabile: {format_caption(seg.asset) or 'elemento visivo'}"
                    ), ds)
                except Exception:  # noqa: BLE001
                    pass

        doc.save(str(file_path))

        return FileResponse(
            path=file_path,
            filename=f"tesi_{safe_title}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    else:
        # Export PDF (default) con indice — usa template
        import fitz

        template = get_template_by_id(template_id, db)
        ps = template.get("pdf", {})

        page_width, page_height = get_page_dimensions(ps.get("page_size", "A4"))
        margin_top = ps.get("margin_top", 50)
        margin_bottom = ps.get("margin_bottom", 50)
        margin_left = ps.get("margin_left", 50)
        margin_right = ps.get("margin_right", 50)
        font_body = ps.get("font_body", "helv")
        font_size = ps.get("font_body_size", 11)
        font_title_size = ps.get("font_title_size", 24)
        font_chapter_size = ps.get("font_chapter_size", 18)
        font_section_size = ps.get("font_section_size", 14)
        line_height_mult = ps.get("line_height_multiplier", 1.5)
        include_toc_pdf = ps.get("include_toc", True)
        include_page_numbers = ps.get("include_page_numbers", True)
        page_number_position = ps.get("page_number_position", "bottom_center")
        include_header = ps.get("include_header", False)
        header_text = ps.get("header_text", "")
        include_footer = ps.get("include_footer", False)
        footer_text = ps.get("footer_text", "")
        title_align = ps.get("title_alignment", "center")
        body_align = ps.get("body_alignment", "left")
        chapter_spacing = ps.get("chapter_spacing_before", 20)
        section_spacing = ps.get("section_spacing_before", 15)
        paragraph_spacing = ps.get("paragraph_spacing", 0)
        toc_separator_color = ps.get("toc_separator_color", [0.7, 0.7, 0.7])
        bg_image_file = ps.get("background_image", "")
        bg_image_mode = ps.get("background_image_mode", "all_pages")
        bg_opacity = ps.get("background_opacity", 0.15)
        bg_image_fit = ps.get("background_image_fit", "tile")

        # Resolve background image path
        bg_image_path = None
        if bg_image_file:
            candidate = config.UPLOAD_DIR / "template_backgrounds" / bg_image_file
            if candidate.exists():
                bg_image_path = str(candidate)

        line_height = font_size * line_height_mult
        content_width = page_width - margin_left - margin_right

        # Calcolo posizione x per allineamenti
        def calc_text_x(text, fontsize, fontname, alignment):
            """Calcola la posizione x basata sull'allineamento."""
            if alignment == "center":
                text_width = fitz.get_text_length(text, fontname=fontname, fontsize=fontsize)
                return margin_left + (content_width - text_width) / 2
            elif alignment == "right":
                text_width = fitz.get_text_length(text, fontname=fontname, fontsize=fontsize)
                return page_width - margin_right - text_width
            return margin_left  # left / justify (default)

        def insert_justified_line(page, x_start, y_pos, words_list, fontsize, fontname, available_width, is_last_line=False):
            """Inserisce una riga di testo giustificato distribuendo gli spazi tra le parole."""
            if is_last_line or len(words_list) <= 1:
                # Ultima riga o singola parola: allinea a sinistra
                page.insert_text((x_start, y_pos), ' '.join(words_list), fontsize=fontsize, fontname=fontname)
                return
            # Calcola lo spazio extra da distribuire
            text_no_spaces = ''.join(words_list)
            text_width = fitz.get_text_length(text_no_spaces, fontname=fontname, fontsize=fontsize)
            total_space = available_width - text_width
            space_between = total_space / (len(words_list) - 1)
            # Inserisci parola per parola
            cx = x_start
            for i, word in enumerate(words_list):
                page.insert_text((cx, y_pos), word, fontsize=fontsize, fontname=fontname)
                word_w = fitz.get_text_length(word, fontname=fontname, fontsize=fontsize)
                cx += word_w + space_between

        file_path = config.RESULTS_DIR / f"thesis_{safe_title}_{timestamp}.pdf"

        pdf_doc = fitz.open()
        page_count = [0]  # Mutable per contare le pagine

        def new_pdf_page():
            """Crea una nuova pagina e incrementa il contatore."""
            p = pdf_doc.new_page(width=page_width, height=page_height)
            page_count[0] += 1
            return p

        current_page = new_pdf_page()
        y = margin_top

        # Titolo principale (con word-wrap)
        title_words = thesis.title.split()
        title_current_line = []
        for t_word in title_words:
            t_test = ' '.join(title_current_line + [t_word])
            t_tw = fitz.get_text_length(t_test, fontname=font_body, fontsize=font_title_size)
            if t_tw < content_width:
                title_current_line.append(t_word)
            else:
                if title_current_line:
                    t_str = ' '.join(title_current_line)
                    t_x = calc_text_x(t_str, font_title_size, font_body, title_align)
                    current_page.insert_text((t_x, y + font_title_size), t_str, fontsize=font_title_size, fontname=font_body)
                    y += font_title_size + 4
                title_current_line = [t_word]
        if title_current_line:
            t_str = ' '.join(title_current_line)
            t_x = calc_text_x(t_str, font_title_size, font_body, title_align)
            current_page.insert_text((t_x, y + font_title_size), t_str, fontsize=font_title_size, fontname=font_body)
            y += font_title_size + 4
        y += 16

        # Separatore
        y += 20

        # Indice
        if toc and include_toc_pdf:
            toc_title_size = font_section_size
            current_page.insert_text(
                (margin_left, y),
                "INDICE",
                fontsize=toc_title_size,
                fontname=font_body
            )
            y += toc_title_size + 10

            # Linea separatrice
            sep_color = tuple(toc_separator_color) if isinstance(toc_separator_color, list) else (0.7, 0.7, 0.7)
            current_page.draw_line(
                fitz.Point(margin_left, y),
                fitz.Point(page_width - margin_right, y),
                color=sep_color,
                width=1
            )
            y += 15

            # Contenuto indice
            chapters = thesis.chapters_structure.get("chapters", []) if thesis.chapters_structure else []
            for ch_idx, chapter in enumerate(chapters):
                if y + line_height * 2 > page_height - margin_bottom:
                    current_page = new_pdf_page()
                    y = margin_top

                ch_title = chapter.get("chapter_title") or chapter.get("title", f"Capitolo {ch_idx + 1}")
                is_special = chapter.get("is_special", False)

                if is_special:
                    # Word-wrap special chapter titles
                    toc_text = ch_title
                    toc_words = toc_text.split()
                    toc_current_line = []
                    for toc_word in toc_words:
                        toc_test_line = ' '.join(toc_current_line + [toc_word])
                        toc_tw = fitz.get_text_length(toc_test_line, fontname=font_body, fontsize=font_size)
                        if toc_tw < content_width:
                            toc_current_line.append(toc_word)
                        else:
                            if toc_current_line:
                                if y + line_height > page_height - margin_bottom:
                                    current_page = new_pdf_page()
                                    y = margin_top
                                current_page.insert_text((margin_left, y), ' '.join(toc_current_line), fontsize=font_size, fontname=font_body)
                                y += line_height
                            toc_current_line = [toc_word]
                    if toc_current_line:
                        if y + line_height > page_height - margin_bottom:
                            current_page = new_pdf_page()
                            y = margin_top
                        current_page.insert_text((margin_left, y), ' '.join(toc_current_line), fontsize=font_size, fontname=font_body)
                        y += line_height
                else:
                    ch_num = chapter.get("chapter_index", ch_idx + 1)

                    # Word-wrap chapter titles
                    toc_text = f"Capitolo {ch_num}: {ch_title}"
                    toc_words = toc_text.split()
                    toc_current_line = []
                    for toc_word in toc_words:
                        toc_test_line = ' '.join(toc_current_line + [toc_word])
                        toc_tw = fitz.get_text_length(toc_test_line, fontname=font_body, fontsize=font_size)
                        if toc_tw < content_width:
                            toc_current_line.append(toc_word)
                        else:
                            if toc_current_line:
                                if y + line_height > page_height - margin_bottom:
                                    current_page = new_pdf_page()
                                    y = margin_top
                                current_page.insert_text((margin_left, y), ' '.join(toc_current_line), fontsize=font_size, fontname=font_body)
                                y += line_height
                            toc_current_line = [toc_word]
                    if toc_current_line:
                        if y + line_height > page_height - margin_bottom:
                            current_page = new_pdf_page()
                            y = margin_top
                        current_page.insert_text((margin_left, y), ' '.join(toc_current_line), fontsize=font_size, fontname=font_body)
                        y += line_height

                    sections = chapter.get("sections", [])
                    for sec_idx, section in enumerate(sections):
                        if y + line_height > page_height - margin_bottom:
                            current_page = new_pdf_page()
                            y = margin_top

                        sec_num = section.get("index", sec_idx + 1)
                        sec_title = section.get("title", f"Sezione {sec_num}")

                        # Word-wrap section titles (indented by 20)
                        sec_text = f"{ch_num}.{sec_num}: {sec_title}"
                        sec_available_width = content_width - 20
                        sec_words = sec_text.split()
                        sec_current_line = []
                        for sec_word in sec_words:
                            sec_test_line = ' '.join(sec_current_line + [sec_word])
                            sec_tw = fitz.get_text_length(sec_test_line, fontname=font_body, fontsize=font_size - 1)
                            if sec_tw < sec_available_width:
                                sec_current_line.append(sec_word)
                            else:
                                if sec_current_line:
                                    if y + line_height > page_height - margin_bottom:
                                        current_page = new_pdf_page()
                                        y = margin_top
                                    current_page.insert_text((margin_left + 20, y), ' '.join(sec_current_line), fontsize=font_size - 1, fontname=font_body)
                                    y += line_height * 0.9
                                sec_current_line = [sec_word]
                        if sec_current_line:
                            if y + line_height > page_height - margin_bottom:
                                current_page = new_pdf_page()
                                y = margin_top
                            current_page.insert_text((margin_left + 20, y), ' '.join(sec_current_line), fontsize=font_size - 1, fontname=font_body)
                            y += line_height * 0.9

                y += 5  # Spazio tra capitoli

            # ── Indici delle tabelle/figure (solo se presenti) ──
            for idx_title, idx_entries in (("INDICE DELLE TABELLE", tables_index),
                                           ("INDICE DELLE FIGURE", figures_index)):
                if not idx_entries:
                    continue
                if y + font_section_size + line_height * 2 > page_height - margin_bottom:
                    current_page = new_pdf_page()
                    y = margin_top
                y += 10
                current_page.insert_text((margin_left, y), idx_title,
                                         fontsize=font_section_size, fontname=font_body)
                y += font_section_size + 8
                for idx_label, idx_caption in idx_entries:
                    entry_text = strip_italic_markers(f"{idx_label} - {idx_caption}")
                    entry_lines = wrap_text_to_width(
                        entry_text, content_width - 20,
                        lambda s: fitz.get_text_length(s, fontname=font_body, fontsize=font_size - 1)
                    )
                    for entry_line in entry_lines:
                        if y + line_height > page_height - margin_bottom:
                            current_page = new_pdf_page()
                            y = margin_top
                        current_page.insert_text((margin_left + 20, y), entry_line,
                                                 fontsize=font_size - 1, fontname=font_body)
                        y += line_height * 0.9
                y += 5

            # Separatore dopo indice
            y += 15
            current_page.draw_line(
                fitz.Point(margin_left, y),
                fitz.Point(page_width - margin_right, y),
                color=sep_color,
                width=1
            )
            y += 30

        # Nuova pagina per il contenuto
        current_page = new_pdf_page()
        y = margin_top

        # Footnote tracking for PDF
        pdf_footnote_num = [1]  # Progressive footnote number
        page_footnotes = []  # Footnotes for current page
        fn_font_size = max(font_size - 3, 7)
        fn_line_height = fn_font_size * 1.3
        fn_separator_space = 15  # Space for separator line above footnotes

        def get_footnotes_height():
            """Calcola altezza necessaria per le note a piè di pagina correnti."""
            if not page_footnotes:
                return 0
            return fn_separator_space + len(page_footnotes) * fn_line_height + 5

        def render_page_footnotes():
            """Renderizza le note raccolte in fondo alla pagina corrente."""
            if not page_footnotes:
                return
            fn_y = page_height - margin_bottom - get_footnotes_height() + fn_separator_space
            # Linea separatrice
            current_page.draw_line(
                fitz.Point(margin_left, fn_y - 8),
                fitz.Point(margin_left + content_width * 0.3, fn_y - 8),
                color=(0.5, 0.5, 0.5),
                width=0.5
            )
            for fn_num, fn_text in page_footnotes:
                fn_label = f"{fn_num} "
                label_width = fitz.get_text_length(fn_label, fontname=font_body, fontsize=fn_font_size)
                current_page.insert_text(
                    (margin_left, fn_y),
                    fn_label,
                    fontsize=fn_font_size,
                    fontname=font_body,
                    color=(0.3, 0.3, 0.3)
                )
                # Wrap footnote text (rimuoviamo i marker Markdown *italic* nel PDF:
                # qui non c'è rendering italic nativo, evitiamo asterischi letterali)
                fn_text = strip_italic_markers(fn_text)
                fn_words = fn_text.split()
                fn_current_line = []
                fn_x_start = margin_left + label_width
                fn_content_width = content_width - label_width
                first_line = True
                for fw in fn_words:
                    test = ' '.join(fn_current_line + [fw])
                    tw = fitz.get_text_length(test, fontname=font_body, fontsize=fn_font_size)
                    if tw < fn_content_width:
                        fn_current_line.append(fw)
                    else:
                        if fn_current_line:
                            x_pos = fn_x_start if first_line else margin_left + label_width
                            if not first_line:
                                fn_y += fn_line_height
                            current_page.insert_text(
                                (x_pos, fn_y if first_line else fn_y),
                                ' '.join(fn_current_line),
                                fontsize=fn_font_size,
                                fontname=font_body,
                                color=(0.3, 0.3, 0.3)
                            )
                            first_line = False
                        fn_current_line = [fw]
                if fn_current_line:
                    x_pos = fn_x_start if first_line else margin_left + label_width
                    if not first_line:
                        fn_y += fn_line_height
                    current_page.insert_text(
                        (x_pos, fn_y),
                        ' '.join(fn_current_line),
                        fontsize=fn_font_size,
                        fontname=font_body,
                        color=(0.3, 0.3, 0.3)
                    )
                fn_y += fn_line_height

        def get_available_y():
            """Altezza massima disponibile per il contenuto (sottraendo footnotes)."""
            return page_height - margin_bottom - get_footnotes_height()

        def check_new_page_needed(needed_height):
            """Verifica se serve nuova pagina. Se sì, renderizza footnotes e crea nuova pagina."""
            nonlocal current_page, y, page_footnotes
            if y + needed_height > get_available_y():
                render_page_footnotes()
                page_footnotes = []
                current_page = new_pdf_page()
                y = margin_top

        # ── Renderer asset (tabelle / grafici / riquadri HINT) ──
        _bold_font_map = {'helv': 'hebo', 'tiro': 'tibo', 'cour': 'cobo'}
        font_body_bold = _bold_font_map.get(font_body, font_body)
        # capacità verticale di una pagina vuota (per troncare contenuti monstre)
        page_capacity = page_height - margin_top - margin_bottom - fn_separator_space

        def _pdf_safe(s):
            """I font base-14 non coprono tutti i glifi Unicode: normalizza i più comuni."""
            return (str(s).replace('–', '-').replace('—', '-')
                    .replace('‘', "'").replace('’', "'")
                    .replace('“', '"').replace('”', '"')
                    .replace('…', '...').replace('⚠', '').replace('️', '').strip())

        def _latin1_math(s):
            """Formule inline → Unicode SOLO se i glifi restano latin-1 (base-14):
            per il greco meglio il sorgente LaTeX leggibile che glifi persi."""
            if '$' not in s:
                return s
            converted = inline_math_to_unicode(s)
            try:
                converted.encode('latin-1')
                return converted
            except UnicodeEncodeError:
                return s

        def render_pdf_caption(text):
            """Didascalia centrata in grigio sotto/sopra un asset."""
            nonlocal y
            cap_size = max(8, font_size - 1)
            cap_lines = wrap_text_to_width(
                _pdf_safe(strip_italic_markers(text)), content_width,
                lambda s: fitz.get_text_length(s, fontname=font_body, fontsize=cap_size))
            for cap_line in cap_lines:
                check_new_page_needed(cap_size + 4)
                cl_w = fitz.get_text_length(cap_line, fontname=font_body, fontsize=cap_size)
                current_page.insert_text((margin_left + (content_width - cl_w) / 2, y), cap_line,
                                         fontsize=cap_size, fontname=font_body, color=(0.25, 0.25, 0.25))
                y += cap_size + 3
            y += 6

        def render_pdf_table(table):
            """Tabella a griglia: header ripetuto ai salti pagina, celle wrappate."""
            nonlocal y
            n = max(1, len(table.header))
            cell_size = max(7.0, font_size - 1.5)
            cell_lh = cell_size * 1.25
            pad = 4
            col_w = content_width / n
            inner_w = max(10.0, col_w - 2 * pad)
            grid_color = (0.35, 0.35, 0.35)
            max_cell_lines = max(1, int((page_capacity - 2 * pad) / cell_lh) - 2)

            cell_space_w = fitz.get_text_length(' ', fontname=font_body, fontsize=cell_size)

            def _cell_lines_one(text, fname):
                """Righe di una cella: liste di [(item, spazio_prima)] con altezza
                e ascendente propri. Item ('t', parola, w) | ('m', MathPng, w, h, d):
                le formule $...$ diventano PNG mathtext allineati al baseline."""
                def _t(word):
                    return ('t', word, fitz.get_text_length(
                        word, fontname=fname, fontsize=cell_size))

                text = str(text)
                if not has_inline_math(text):
                    wrapped = wrap_text_to_width(
                        _pdf_safe(strip_italic_markers(text)), inner_w,
                        lambda s: fitz.get_text_length(s, fontname=fname, fontsize=cell_size))
                    return [([(_t(ln), False)], cell_lh, cell_size) for ln in wrapped]

                # Cluster = item senza spazio interno nel sorgente: la
                # punteggiatura resta attaccata alla formula (come nel corpo)
                glued_clusters = []

                def add_item(item, glued):
                    if glued and glued_clusters:
                        glued_clusters[-1].append(item)
                    else:
                        glued_clusters.append([item])

                last_ends_nospace = False
                for kind, latex, raw in iter_inline_math(text):
                    if kind == 'math':
                        try:
                            mp = render_math_png(latex, fontsize=cell_size, dpi=300)
                            scale = min(1.0, inner_w / mp.width_pt) if mp.width_pt else 1.0
                            add_item(('m', mp, mp.width_pt * scale,
                                      mp.height_pt * scale, mp.depth_pt * scale),
                                     last_ends_nospace)
                            last_ends_nospace = True
                            continue
                        except MathRenderError as e:
                            logger.warning(f"Formula non renderizzata in cella PDF: {e}")
                            raw = _latin1_math(raw)  # sorgente/Unicode come testo
                    starts_ws = raw[:1].isspace()
                    words = _pdf_safe(strip_italic_markers(raw)).split()
                    if not words:
                        if raw:
                            last_ends_nospace = False  # soli spazi: separa
                        continue
                    for wi, w_ in enumerate(words):
                        add_item(_t(w_), glued=(wi == 0 and not starts_ws
                                                and last_ends_nospace))
                    last_ends_nospace = not raw[-1:].isspace()

                def cluster_w(cluster):
                    return sum(it[2] for it in cluster)

                # cluster più largo della cella: meglio spezzarlo che sbordare
                clusters = [c for cluster in glued_clusters
                            for c in ([[it] for it in cluster]
                                      if len(cluster) > 1 and cluster_w(cluster) > inner_w
                                      else [cluster])]
                wrapped, cur, cur_w = [], [], 0.0
                for cluster in clusters:
                    cw = cluster_w(cluster)
                    needed = cw if not cur else cur_w + cell_space_w + cw
                    if cur and needed > inner_w:
                        wrapped.append(cur)
                        cur, cur_w = [cluster], cw
                    else:
                        cur.append(cluster)
                        cur_w = needed
                if cur:
                    wrapped.append(cur)

                out = []
                for ln_clusters in wrapped:
                    flat = [(it, ci > 0 and ii == 0)
                            for ci, cluster in enumerate(ln_clusters)
                            for ii, it in enumerate(cluster)]
                    asc, desc = cell_size, cell_size * 0.25
                    for it, _sp in flat:
                        if it[0] == 'm':
                            asc = max(asc, it[3] - it[4])
                            desc = max(desc, it[4])
                    out.append((flat, max(cell_lh, asc + desc + 2), asc))
                return out

            def cell_lines(cells, fname):
                out = []
                for ctext in cells:
                    cl = _cell_lines_one(ctext, fname)
                    if len(cl) > max_cell_lines:
                        logger.warning("Cella di tabella più alta di una pagina: contenuto troncato nel PDF")
                        cl = cl[:max_cell_lines]
                        flat, lh, asc = cl[-1]
                        cl[-1] = (flat + [(('t', '...', fitz.get_text_length(
                            '...', fontname=fname, fontsize=cell_size)), True)], lh, asc)
                    out.append(cl)
                return out

            def row_height(lines_per_cell):
                return max(sum(lh for _, lh, _ in cl) for cl in lines_per_cell) + 2 * pad

            def draw_row(lines_per_cell, fname, fill=None):
                nonlocal y
                rh = row_height(lines_per_cell)
                if fill:
                    current_page.draw_rect(
                        fitz.Rect(margin_left, y, margin_left + content_width, y + rh),
                        fill=fill, color=None)
                for c in range(n + 1):
                    grid_x = margin_left + c * col_w
                    current_page.draw_line(fitz.Point(grid_x, y), fitz.Point(grid_x, y + rh),
                                           color=grid_color, width=0.5)
                current_page.draw_line(fitz.Point(margin_left, y + rh),
                                       fitz.Point(margin_left + content_width, y + rh),
                                       color=grid_color, width=0.5)
                for c, cl in enumerate(lines_per_cell):
                    slot_top = y + pad
                    for flat, lh, asc in cl:
                        baseline = slot_top + asc
                        tx = margin_left + c * col_w + pad
                        for it, spaced in flat:
                            if spaced:
                                tx += cell_space_w
                            if it[0] == 't':
                                current_page.insert_text((tx, baseline), it[1],
                                                         fontsize=cell_size, fontname=fname)
                            else:
                                mp_h, mp_d = it[3], it[4]
                                current_page.insert_image(
                                    fitz.Rect(tx, baseline - (mp_h - mp_d),
                                              tx + it[2], baseline + mp_d),
                                    stream=it[1].png)
                            tx += it[2]
                        slot_top += lh
                y += rh

            header_lines = cell_lines(table.header, font_body_bold)
            header_h = row_height(header_lines)

            def start_table_block():
                nonlocal y
                current_page.draw_line(fitz.Point(margin_left, y),
                                       fitz.Point(margin_left + content_width, y),
                                       color=grid_color, width=0.5)
                draw_row(header_lines, font_body_bold, fill=(0.93, 0.93, 0.93))

            # didascalia + header + prima riga restano insieme
            first_row_h = row_height(cell_lines(table.rows[0], font_body)) if table.rows else 0
            caption_est = (font_size + 2) * 2 + 6
            check_new_page_needed(caption_est + header_h + first_row_h)

            render_pdf_caption(format_caption(table))
            start_table_block()
            for row in table.rows:
                lines_per_cell = cell_lines(row, font_body)
                rh = row_height(lines_per_cell)
                if y + rh > get_available_y():
                    check_new_page_needed(rh)   # footnotes + nuova pagina
                    start_table_block()          # header ripetuto
                draw_row(lines_per_cell, font_body)

            if table.source:
                src_size = max(7, font_size - 2)
                y += src_size + 3  # baseline sotto il bordo inferiore della tabella
                src_lines = wrap_text_to_width(
                    _pdf_safe(f"Fonte: {table.source}"), content_width,
                    lambda s: fitz.get_text_length(s, fontname=font_body, fontsize=src_size))
                for src_line in src_lines:
                    check_new_page_needed(src_size + 3)
                    sl_w = fitz.get_text_length(src_line, fontname=font_body, fontsize=src_size)
                    current_page.insert_text((margin_left + (content_width - sl_w) / 2, y), src_line,
                                             fontsize=src_size, fontname=font_body, color=(0.35, 0.35, 0.35))
                    y += src_size + 3
            y += 10

        def render_pdf_hint(hint):
            """Riquadro ambra ben visibile: segnaposto da sostituire manualmente."""
            nonlocal y
            h_size = max(8, font_size - 1)
            h_lh = h_size * 1.35
            pad = 8
            body_lines = wrap_text_to_width(
                _pdf_safe(hint.text), content_width - 2 * pad,
                lambda s: fitz.get_text_length(s, fontname=font_body, fontsize=h_size))
            max_hint_lines = max(1, int((page_capacity - 2 * pad) / h_lh) - 2)
            if len(body_lines) > max_hint_lines:
                body_lines = body_lines[:max_hint_lines - 1] + [body_lines[max_hint_lines - 1] + '...']
            box_h = (len(body_lines) + 1) * h_lh + 2 * pad
            check_new_page_needed(box_h + 8)
            current_page.draw_rect(
                fitz.Rect(margin_left, y, margin_left + content_width, y + box_h),
                fill=(1.0, 0.953, 0.804), color=(0.72, 0.53, 0.04), width=1)
            ty = y + pad + h_size
            current_page.insert_text((margin_left + pad, ty), "SUGGERIMENTO - DA SOSTITUIRE:",
                                     fontsize=h_size, fontname=font_body_bold, color=(0.45, 0.32, 0.02))
            ty += h_lh
            for body_line in body_lines:
                current_page.insert_text((margin_left + pad, ty), body_line,
                                         fontsize=h_size, fontname=font_body, color=(0.25, 0.2, 0.05))
                ty += h_lh
            y += box_h + 10

        def render_pdf_chart(chart):
            """Grafico PNG centrato con didascalia sotto; su errore → riquadro HINT."""
            nonlocal y
            try:
                png = render_chart_png(chart)
            except ChartRenderError as e:
                logger.warning(f"Grafico '{(chart.caption or '')[:60]}' non renderizzato nel PDF: {e}")
                render_pdf_hint(HintAsset(
                    text=f"Grafico non generabile ({e}). Inserire manualmente: {chart.caption}"))
                return
            pix = fitz.Pixmap(png)
            aspect = (pix.height / pix.width) if pix.width else 0.62
            disp_w = min(content_width, 430)
            disp_h = disp_w * aspect
            if disp_h > page_capacity - 60:
                disp_h = page_capacity - 60
                disp_w = disp_h / aspect
            check_new_page_needed(disp_h + 10 + (font_size + 2) * 2)
            img_x = margin_left + (content_width - disp_w) / 2
            current_page.insert_image(fitz.Rect(img_x, y, img_x + disp_w, y + disp_h), stream=png)
            y += disp_h + 8
            render_pdf_caption(format_caption(chart))

        def render_pdf_math(math_asset):
            """Equazione display: PNG mathtext centrato, numero "(N.M)" a destra.

            Su MathRenderError degrada al sorgente LaTeX come testo (ASCII: i
            font base-14 non hanno i glifi greci, il PNG sì)."""
            nonlocal y
            try:
                # display=True applica già il fattore 1.15 dentro render_math_png
                mp = render_math_png(math_asset.latex, fontsize=font_size,
                                     dpi=300, display=True)
            except MathRenderError as e:
                logger.warning(f"Equazione display non renderizzata nel PDF: {e}")
                fallback = f"$$ {math_asset.latex} $$"
                if math_asset.label:
                    fallback += f" {math_asset.label}"
                render_pdf_text_line(_pdf_safe(fallback))
                return
            num_w = 0.0
            if math_asset.label:
                num_w = fitz.get_text_length(math_asset.label, fontname=font_body,
                                             fontsize=font_size)
            # la formula è centrata su content_width: perché non tocchi il
            # numero a destra deve restare dentro content_width - 2*num_w
            max_w = content_width - (2 * num_w + 12 if num_w else 0)
            scale = 1.0
            disp_w, disp_h = mp.width_pt, mp.height_pt
            if disp_w > max_w:
                scale = max_w / disp_w
                disp_w, disp_h = disp_w * scale, disp_h * scale
            if disp_h > page_capacity - 40:
                s2 = (page_capacity - 40) / disp_h
                scale *= s2
                disp_w, disp_h = disp_w * s2, disp_h * s2
            check_new_page_needed(disp_h + 12)
            y += 6
            img_x = margin_left + (content_width - disp_w) / 2
            current_page.insert_image(
                fitz.Rect(img_x, y, img_x + disp_w, y + disp_h), stream=mp.png)
            if math_asset.label:
                # numero al margine destro, allineato al baseline della formula
                baseline = y + disp_h - mp.depth_pt * scale
                current_page.insert_text(
                    (page_width - margin_right - num_w, baseline),
                    math_asset.label, fontsize=font_size, fontname=font_body)
            y += disp_h + 6

        def render_pdf_inline_math_line(line):
            """Riga di corpo con formule $...$: parole e PNG inline sul baseline.

            Wrap greedy a item (parola o formula) sulla larghezza colonna;
            giustificazione a gap uniforme come insert_justified_line; le
            formule sono allineate al baseline tramite la metrica depth."""
            nonlocal y
            # Cluster = sequenza indivisibile di item senza spazio interno, così
            # la punteggiatura resta attaccata alla formula ("ω/ωₙ," e non "ω/ωₙ ,")
            # e il wrap non spezza mai formula e virgola.
            clusters = []  # lista di liste di ('word', testo, w, _) | ('math', MathPng, w, scala)

            def add_item(item, glued):
                if glued and clusters:
                    clusters[-1].append(item)
                else:
                    clusters.append([item])

            last_ends_nospace = False
            for kind, latex, raw in iter_inline_math(line):
                if kind == 'math':
                    try:
                        mp = render_math_png(latex, fontsize=font_size, dpi=300)
                        scale = min(1.0, content_width / mp.width_pt) if mp.width_pt else 1.0
                        add_item(('math', mp, mp.width_pt * scale, scale), last_ends_nospace)
                        last_ends_nospace = True
                        continue
                    except MathRenderError as e:
                        logger.warning(f"Formula inline non renderizzata nel PDF: {e}")
                        raw = _pdf_safe(raw)  # sorgente come testo
                chunk = strip_italic_markers(raw)
                words = chunk.split()
                if not words:
                    if chunk:
                        last_ends_nospace = False  # soli spazi: separa
                    continue
                starts_ws = chunk[0].isspace()
                for wi, w in enumerate(words):
                    item = ('word', w, fitz.get_text_length(
                        w, fontname=font_body, fontsize=font_size), 1.0)
                    add_item(item, glued=(wi == 0 and not starts_ws and last_ends_nospace))
                last_ends_nospace = not chunk[-1].isspace()
            if not clusters:
                return
            space_w = fitz.get_text_length(' ', fontname=font_body, fontsize=font_size)

            def cluster_w(cluster):
                return sum(it[2] for it in cluster)

            # Un cluster più largo della colonna (formula quasi a tutta pagina
            # + punteggiatura incollata) sborderebbe: meglio spezzarlo che
            # uscire dal margine destro.
            clusters = [c for cluster in clusters
                        for c in ([[it] for it in cluster]
                                  if len(cluster) > 1 and cluster_w(cluster) > content_width
                                  else [cluster])]

            rows, row, row_w = [], [], 0.0
            for cluster in clusters:
                cw = cluster_w(cluster)
                needed = cw if not row else row_w + space_w + cw
                if row and needed > content_width:
                    rows.append(row)
                    row, row_w = [cluster], cw
                else:
                    row.append(cluster)
                    row_w = needed
            if row:
                rows.append(row)

            text_asc = font_size * 0.8    # ascendente tipico dei font base-14
            text_desc = font_size * 0.25
            for ri, current_row in enumerate(rows):
                asc, desc = text_asc, text_desc
                for cluster in current_row:
                    for item in cluster:
                        if item[0] == 'math':
                            mp, scale = item[1], item[3]
                            asc = max(asc, (mp.height_pt - mp.depth_pt) * scale)
                            desc = max(desc, mp.depth_pt * scale)
                row_h = max(line_height, asc + desc + 2)
                check_new_page_needed(row_h)
                # il baseline scende se la matematica è più alta del testo
                baseline = y + (asc - text_asc)
                is_last = ri == len(rows) - 1
                total_w = sum(cluster_w(c) for c in current_row)
                gaps = len(current_row) - 1
                if body_align == "justify" and not is_last and gaps > 0:
                    gap = (content_width - total_w) / gaps
                else:
                    gap = space_w
                # stesso allineamento del fast path (calc_text_x)
                row_total = total_w + gaps * gap
                if body_align == "center":
                    cx = margin_left + max(0.0, content_width - row_total) / 2
                elif body_align == "right":
                    cx = margin_left + max(0.0, content_width - row_total)
                else:
                    cx = margin_left
                for cluster in current_row:
                    for item in cluster:
                        if item[0] == 'word':
                            current_page.insert_text((cx, baseline), item[1],
                                                     fontsize=font_size, fontname=font_body)
                        else:
                            mp, scale = item[1], item[3]
                            h, d = mp.height_pt * scale, mp.depth_pt * scale
                            current_page.insert_image(
                                fitz.Rect(cx, baseline - (h - d), cx + item[2], baseline + d),
                                stream=mp.png)
                        cx += item[2]
                    cx += gap
                y += row_h
            if paragraph_spacing > 0:
                y += paragraph_spacing

        # Contenuto con footnotes
        def render_pdf_text_line(line):
            nonlocal y
            check_new_page_needed(line_height)

            # Gestisci titoli
            if line.startswith('# '):
                y += chapter_spacing
                check_new_page_needed(font_chapter_size + 10)
                # Word-wrap chapter title (eventuale math → Unicode se latin-1)
                ch_title_text = _latin1_math(line[2:])
                ch_words = ch_title_text.split()
                ch_current_line = []
                for ch_word in ch_words:
                    ch_test_line = ' '.join(ch_current_line + [ch_word])
                    ch_tw = fitz.get_text_length(ch_test_line, fontname=font_body, fontsize=font_chapter_size)
                    if ch_tw < content_width:
                        ch_current_line.append(ch_word)
                    else:
                        if ch_current_line:
                            check_new_page_needed(font_chapter_size + 4)
                            current_page.insert_text((margin_left, y), ' '.join(ch_current_line), fontsize=font_chapter_size, fontname=font_body)
                            y += font_chapter_size + 4
                        ch_current_line = [ch_word]
                if ch_current_line:
                    check_new_page_needed(font_chapter_size + 4)
                    current_page.insert_text((margin_left, y), ' '.join(ch_current_line), fontsize=font_chapter_size, fontname=font_body)
                    y += font_chapter_size + 4
                y += 4
            elif line.startswith('## '):
                y += section_spacing
                check_new_page_needed(font_section_size + 8)
                # Word-wrap section title (eventuale math → Unicode se latin-1)
                sec_title_text = _latin1_math(line[3:])
                sec_words = sec_title_text.split()
                sec_current_line = []
                for sec_word in sec_words:
                    sec_test_line = ' '.join(sec_current_line + [sec_word])
                    sec_tw = fitz.get_text_length(sec_test_line, fontname=font_body, fontsize=font_section_size)
                    if sec_tw < content_width:
                        sec_current_line.append(sec_word)
                    else:
                        if sec_current_line:
                            check_new_page_needed(font_section_size + 3)
                            current_page.insert_text((margin_left, y), ' '.join(sec_current_line), fontsize=font_section_size, fontname=font_body)
                            y += font_section_size + 3
                        sec_current_line = [sec_word]
                if sec_current_line:
                    check_new_page_needed(font_section_size + 3)
                    current_page.insert_text((margin_left, y), ' '.join(sec_current_line), fontsize=font_section_size, fontname=font_body)
                    y += font_section_size + 3
            elif line.strip():
                # Estrazione note sulla riga con formule PROTETTE: le graffe
                # }} del LaTeX (\frac{a}{b}}) non devono chiudere la {{nota:}}
                pline, line_math_map = protect_math_spans(line)
                footnotes_in_line = extract_footnotes_from_line(pline)

                if footnotes_in_line:
                    # Process line: strip {{nota:...}} and replace with superscript numbers
                    processed_line = ""
                    last_end = 0
                    line_fn_nums = []
                    for fn_start, fn_end, fn_text in footnotes_in_line:
                        processed_line += pline[last_end:fn_start]
                        fn_num = pdf_footnote_num[0]
                        processed_line += f"[{fn_num}]"
                        fn_text = _latin1_math(unprotect_math_spans(fn_text, line_math_map))
                        line_fn_nums.append((fn_num, fn_text))
                        page_footnotes.append((fn_num, fn_text))
                        pdf_footnote_num[0] += 1
                        last_end = fn_end
                    processed_line += pline[last_end:]
                    line = unprotect_math_spans(processed_line, line_math_map)

                # Formule inline: layout a run dedicato (PNG sul baseline).
                # Va PRIMA dello strip corsivi: un * dentro $...$ non è Markdown.
                if has_inline_math(line):
                    render_pdf_inline_math_line(line)
                    return

                # PDF non gestisce nativamente il corsivo Markdown: rimuoviamo gli asterischi
                # così i titoli APA non appaiono come letterali. (Il DOCX usa run italic.)
                line = strip_italic_markers(line)

                # Wrap text
                words = line.split()
                current_line = []
                wrapped_lines = []
                for word in words:
                    test_line = ' '.join(current_line + [word])
                    text_width = fitz.get_text_length(test_line, fontname=font_body, fontsize=font_size)
                    if text_width < content_width:
                        current_line.append(word)
                    else:
                        if current_line:
                            wrapped_lines.append(current_line)
                        current_line = [word]
                if current_line:
                    wrapped_lines.append(current_line)

                for li, wline in enumerate(wrapped_lines):
                    check_new_page_needed(line_height)
                    is_last = (li == len(wrapped_lines) - 1)
                    if body_align == "justify" and not is_last and len(wline) > 1:
                        insert_justified_line(current_page, margin_left, y, wline, font_size, font_body, content_width)
                    else:
                        text_str = ' '.join(wline)
                        text_x = calc_text_x(text_str, font_size, font_body, body_align)
                        current_page.insert_text((text_x, y), text_str, fontsize=font_size, fontname=font_body)
                    y += line_height

                # Spazio extra tra paragrafi
                if paragraph_spacing > 0:
                    y += paragraph_spacing
            else:
                y += line_height * 0.5

        for seg in segments:
            if seg.kind == 'text':
                for line in seg.text.split('\n'):
                    render_pdf_text_line(line)
                continue
            # Un asset malformato non deve MAI far fallire l'export:
            # degrada a riquadro HINT.
            try:
                if seg.kind == 'table':
                    render_pdf_table(seg.asset)
                elif seg.kind == 'chart':
                    render_pdf_chart(seg.asset)
                elif seg.kind == 'math':
                    render_pdf_math(seg.asset)
                elif seg.kind == 'hint':
                    render_pdf_hint(seg.asset)
            except Exception as e:  # noqa: BLE001
                logger.warning(f"Asset non renderizzabile nel PDF: {e}")
                try:
                    render_pdf_hint(HintAsset(
                        text=f"Elemento non renderizzabile: {format_caption(seg.asset) or 'elemento visivo'}"
                    ))
                except Exception:  # noqa: BLE001
                    pass

        # Renderizza le ultime footnotes
        render_page_footnotes()

        # Aggiungi sfondo/header/footer/numeri pagina a tutte le pagine
        total_pages = len(pdf_doc)
        for page_idx in range(total_pages):
            page = pdf_doc[page_idx]

            # Background image (behind content)
            if bg_image_path:
                apply_bg = (bg_image_mode == "all_pages") or (bg_image_mode == "first_page_only" and page_idx == 0)
                if apply_bg:
                    try:
                        from PIL import Image
                        import io as _io

                        img = Image.open(bg_image_path).convert("RGBA")
                        if bg_opacity < 1.0:
                            white_bg = Image.new("RGBA", img.size, (255, 255, 255, 255))
                            img = Image.blend(white_bg, img, bg_opacity)

                        pw, ph = int(page_width), int(page_height)

                        if bg_image_fit == "tile":
                            # Tile: repeat image across the page
                            canvas = Image.new("RGB", (pw, ph), (255, 255, 255))
                            iw, ih = img.size
                            for ty in range(0, ph, ih):
                                for tx in range(0, pw, iw):
                                    canvas.paste(img, (tx, ty), img if img.mode == "RGBA" else None)
                            final_img = canvas
                        elif bg_image_fit == "original":
                            # Original size from top-left corner
                            canvas = Image.new("RGB", (pw, ph), (255, 255, 255))
                            canvas.paste(img, (0, 0), img if img.mode == "RGBA" else None)
                            final_img = canvas
                        elif bg_image_fit == "center":
                            # Original size centered
                            canvas = Image.new("RGB", (pw, ph), (255, 255, 255))
                            iw, ih = img.size
                            x = (pw - iw) // 2
                            y = (ph - ih) // 2
                            canvas.paste(img, (x, y), img if img.mode == "RGBA" else None)
                            final_img = canvas
                        else:
                            # Stretch: fill entire page
                            final_img = img.convert("RGB").resize((pw, ph), Image.LANCZOS)

                        buf = _io.BytesIO()
                        final_img.save(buf, format="PNG")
                        buf.seek(0)
                        bg_rect = fitz.Rect(0, 0, page_width, page_height)
                        page.insert_image(bg_rect, stream=buf.read(), overlay=False)
                    except Exception as e:
                        logger.warning(f"Errore inserimento sfondo PDF: {e}")

            # Header
            if include_header and header_text:
                header_x = calc_text_x(header_text, 8, font_body, "center")
                page.insert_text(
                    (header_x, margin_top - 15),
                    header_text,
                    fontsize=8,
                    fontname=font_body,
                    color=(0.5, 0.5, 0.5)
                )
                # Linea sotto header
                page.draw_line(
                    fitz.Point(margin_left, margin_top - 8),
                    fitz.Point(page_width - margin_right, margin_top - 8),
                    color=(0.85, 0.85, 0.85),
                    width=0.5
                )

            # Footer text
            if include_footer and footer_text:
                footer_y = page_height - margin_bottom + 20
                footer_x = calc_text_x(footer_text, 8, font_body, "center")
                page.insert_text(
                    (footer_x, footer_y),
                    footer_text,
                    fontsize=8,
                    fontname=font_body,
                    color=(0.5, 0.5, 0.5)
                )

            # Numeri di pagina
            if include_page_numbers:
                page_num_text = str(page_idx + 1)
                pn_fontsize = 9

                if page_number_position == "bottom_center":
                    pn_x = calc_text_x(page_num_text, pn_fontsize, font_body, "center")
                    pn_y = page_height - margin_bottom + 10 + (15 if include_footer and footer_text else 0)
                elif page_number_position == "bottom_right":
                    pn_x = page_width - margin_right - fitz.get_text_length(page_num_text, fontname=font_body, fontsize=pn_fontsize)
                    pn_y = page_height - margin_bottom + 10 + (15 if include_footer and footer_text else 0)
                elif page_number_position == "top_center":
                    pn_x = calc_text_x(page_num_text, pn_fontsize, font_body, "center")
                    pn_y = margin_top - 25
                elif page_number_position == "top_right":
                    pn_x = page_width - margin_right - fitz.get_text_length(page_num_text, fontname=font_body, fontsize=pn_fontsize)
                    pn_y = margin_top - 25
                else:
                    pn_x = calc_text_x(page_num_text, pn_fontsize, font_body, "center")
                    pn_y = page_height - margin_bottom + 10

                page.insert_text(
                    (pn_x, pn_y),
                    page_num_text,
                    fontsize=pn_fontsize,
                    fontname=font_body,
                    color=(0.5, 0.5, 0.5)
                )

        # deflate: comprime gli stream (i PNG dei grafici altrimenti restano raw)
        pdf_doc.save(file_path, deflate=True)
        pdf_doc.close()

        return FileResponse(
            path=file_path,
            filename=f"tesi_{safe_title}.pdf",
            media_type="application/pdf"
        )
